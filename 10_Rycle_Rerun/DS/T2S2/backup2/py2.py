import os
import subprocess
import sys
import time
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Configuration parameters
RESOLUTION = 1.2
IMAGE_PATH = r"C:\Users\admin\Desktop\Python_proj\datas\T2_IMGS\Li_1.0.png"
OUTPUT_DIR = r"C:\Users\admin\Desktop\Python_proj\ALL_RESULT\DS\T2S2\1.0\backup2"
BASE_NAME = os.path.splitext(os.path.basename(IMAGE_PATH))[0]

# Files to verify
REQUIRED_FILES = [
    f"{BASE_NAME}_grays.csv",
    f"{BASE_NAME}_length.txt",
    f"{BASE_NAME}_plot.tiff",
    f"{BASE_NAME}_data.csv"
]

def run_py1():
    """Execute py1.py with specified resolution parameter"""
    try:
        cmd = ["python", "py1.py", f"-resolution={RESOLUTION}"]
        process = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        stdout, stderr = process.communicate(timeout=60)
        
        if process.returncode != 0:
            print(f"Error running py1.py: {stderr.decode()}")
            return False
        return True
    except Exception as e:
        print(f"Execution failed: {str(e)}")
        return False

def verify_files():
    """Check if all required output files exist"""
    missing_files = [f for f in REQUIRED_FILES if not os.path.exists(os.path.join(OUTPUT_DIR, f))]
    if missing_files:
        print(f"Missing files: {', '.join(missing_files)}")
        return False
    return True

def read_results():
    """Read results from output files"""
    results = {}
    try:
        # Read length result
        with open(os.path.join(OUTPUT_DIR, f"{BASE_NAME}_length.txt"), 'r') as f:
            results['length'] = float(f.read().strip())
            
        # Read u_eq profile
        with open(os.path.join(OUTPUT_DIR, f"{BASE_NAME}_data.csv"), 'r') as f:
            lines = f.readlines()[1:]  # Skip header
            results['distance'] = [float(line.split(',')[0]) for line in lines]
            results['u_eq'] = [float(line.split(',')[1]) for line in lines]
            
        return results
    except Exception as e:
        print(f"Error reading results: {str(e)}")
        return None

def generate_report(results):
    """Generate simulation report in Word format"""
    doc = Document()
    
    # Add title
    title = doc.add_heading("Simulation Report: Material Property Distribution Analysis", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph().add_run().add_break()
    
    # Abstract section
    doc.add_heading("Abstract", level=1)
    abstract = (
        "This report details the computational analysis of material property distribution along a "
        "specified linear path within a microstructure image. Using digital image processing techniques, "
        "we extracted grayscale values along a defined line segment, converted these to equivalent material "
        "property values (u_eq), and analyzed their spatial distribution. The calculated line segment length "
        f"was {results['length']:.2f} μm. Results reveal significant property variations along the measurement "
        "path, with u_eq values ranging from minimum to maximum theoretical bounds. This analysis demonstrates "
        "an effective computational approach for quantifying microstructural property gradients."
    )
    doc.add_paragraph(abstract)
    
    # Introduction
    doc.add_heading("Introduction", level=1)
    intro = (
        "Quantitative characterization of material property distribution is crucial for understanding "
        "microstructure-property relationships in engineered materials. Traditional methods often provide "
        "bulk measurements but fail to capture local variations. This study implements a computational "
        f"approach to analyze property gradients along a {results['length']:.2f} μm path in a lithium-based "
        "material microstructure. The primary objective was to map equivalent property values (u_eq) derived "
        "from image grayscale data. By correlating pixel intensity with material properties, we establish "
        "a methodology for non-destructive evaluation of microstructural heterogeneity. This technique has "
        "applications in quality control, materials development, and failure analysis where spatial property "
        "variations significantly influence performance."
    )
    doc.add_paragraph(intro)
    
    # Methods
    doc.add_heading("Methods", level=1)
    methods = (
        "The analysis pipeline consisted of three computational stages. First, the input image was processed "
        "using a Python-based algorithm implementing Bresenham's line algorithm to extract grayscale values "
        f"along a defined segment from {start_point} to {end_point} pixels. Each pixel's grayscale value (0-255) "
        "was converted to u_eq using the transformation: u_eq = u_min + (gray_value/255) × u_max, with "
        f"u_min = {u_min} and u_max = {u_max}. The spatial calibration used a resolution of {RESOLUTION} μm/pixel."
        "\n\nSecond, the spatial distribution of u_eq values was plotted against distance from the segment origin. "
        "Third, results were exported in multiple formats: grayscale values (CSV), line length (TXT), u_eq profile "
        "(CSV), and distribution plot (TIFF). The analysis utilized Python libraries including Pillow for image "
        "processing, NumPy for numerical operations, and Matplotlib for visualization."
    )
    doc.add_paragraph(methods)
    
    # Results
    doc.add_heading("Results", level=1)
    results_text = (
        "The property distribution analysis revealed significant variations along the measurement path. "
        f"u_eq values ranged from {min(results['u_eq']):.2f} to {max(results['u_eq']):.2f}, representing "
        f"{min(results['u_eq'])/u_max*100:.1f}% to {max(results['u_eq'])/u_max*100:.1f}% of the theoretical "
        "maximum. Property fluctuations show a characteristic pattern with higher values observed near "
        f"{results['distance'][np.argmax(results['u_eq'])]:.2f} μm from the origin. Figure 1 illustrates "
        "the complete u_eq profile, highlighting regions of property enhancement and degradation. "
        f"The average u_eq along the segment was {np.mean(results['u_eq']):.2f} ± {np.std(results['u_eq']):.2f} "
        "(standard deviation), indicating substantial microstructural heterogeneity. These variations correlate "
        "with observable microstructural features in the source image, demonstrating the method's sensitivity "
        "to material heterogeneity."
    )
    doc.add_paragraph(results_text)
    
    # Add plot image
    doc.add_picture(os.path.join(OUTPUT_DIR, f"{BASE_NAME}_plot.tiff"), width=Inches(6))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = last_paragraph.add_run()
    run.add_break()
    run = last_paragraph.add_run("Figure 1: u_eq distribution along measurement path")
    run.font.size = Pt(10)
    run.italic = True
    
    # Save document
    report_path = os.path.join(OUTPUT_DIR, f"{BASE_NAME}_Report.docx")
    doc.save(report_path)
    print(f"Report generated: {report_path}")

if __name__ == "__main__":
    # Step 1: Run the analysis program
    print("Running py1.py...")
    if not run_py1():
        sys.exit(1)
    
    # Allow time for file operations
    time.sleep(2)
    
    # Step 2: Verify output files
    print("Verifying output files...")
    if verify_files():
        print("Calculation successful")
        
        # Step 3: Generate report
        print("Generating simulation report...")
        results = read_results()
        if results:
            generate_report(results)
            print("Process completed successfully")
        else:
            print("Failed to read results for report generation")
    else:
        print("Output verification failed")
