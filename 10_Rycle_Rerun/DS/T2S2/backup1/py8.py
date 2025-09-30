import os
import csv
import numpy as np
from PIL import Image  # Import added to fix the verification error
from docx import Document

# Configuration parameters
INPUT_IMAGE = r"C:\Users\admin\Desktop\Python_proj\datas\T2_IMGS\Li_1.0.png"
OUTPUT_DIR = r"C:\Users\admin\Desktop\Python_proj\ALL_RESULT\DS\T2S2\1.0\backup1"
BASE_NAME = "Li_1.0"

# Expected output files
REQUIRED_FILES = [
    f"{BASE_NAME}_grayscale.csv",
    f"{BASE_NAME}_length.txt",
    f"{BASE_NAME}_data.csv",
    f"{BASE_NAME}_plot.tiff",
    f"{BASE_NAME}_Simulation_Report.docx"
]

def verify_file_existence():
    """Check if all required output files exist"""
    missing = []
    for filename in REQUIRED_FILES:
        path = os.path.join(OUTPUT_DIR, filename)
        if not os.path.exists(path):
            missing.append(path)
    return missing

def verify_report_content():
    """Validate the content and structure of the simulation report"""
    report_path = os.path.join(OUTPUT_DIR, f"{BASE_NAME}_Simulation_Report.docx")
    if not os.path.exists(report_path):
        return ["Report file does not exist"]
    
    doc = Document(report_path)
    issues = []
    
    # Verify section structure
    required_sections = ['Abstract', 'Introduction', 'Methods', 'Results']
    section_headings = [para.text for para in doc.paragraphs if para.style.name.startswith('Heading')]
    
    for section in required_sections:
        if section not in section_headings:
            issues.append(f"Missing section: {section}")
    
    # Verify word count
    total_words = 0
    for para in doc.paragraphs:
        total_words += len(para.text.split())
    
    if total_words < 450:
        issues.append(f"Report too short: {total_words} words (minimum 450 required)")
    
    # Verify figure inclusion
    fig_found = False
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            fig_found = True
            break
            
    if not fig_found:
        issues.append("Figure 1 not found in report")
    
    return issues

def verify_data_consistency():
    """Check consistency between data files"""
    issues = []
    
    try:
        # Verify segment length
        length_path = os.path.join(OUTPUT_DIR, f"{BASE_NAME}_length.txt")
        with open(length_path, 'r') as f:
            length = float(f.read().strip())
        
        # Verify physical dimensions
        if not (40 <= length <= 50):  # Expected range for this segment
            issues.append(f"Unexpected segment length: {length:.2f} µm (expected 40-50 µm)")
        
        # Verify grayscale data
        grayscale_path = os.path.join(OUTPUT_DIR, f"{BASE_NAME}_grayscale.csv")
        grayscale = []
        with open(grayscale_path, 'r') as f:
            reader = csv.reader(f)
            for row in reader:
                if row:  # Skip empty rows
                    try:
                        value = int(row[0])
                        if not (0 <= value <= 255):
                            issues.append(f"Invalid grayscale value: {value} (should be 0-255)")
                        grayscale.append(value)
                    except ValueError:
                        issues.append(f"Non-integer grayscale value: {row[0]}")
        
        # Verify u_eq data
        data_path = os.path.join(OUTPUT_DIR, f"{BASE_NAME}_data.csv")
        distances, u_equ = [], []
        with open(data_path, 'r') as f:
            reader = csv.reader(f)
            header = next(reader)  # Skip header
            if header != ['distance', 'u_eq'] and header != ['distance,u_eq']:
                issues.append(f"Unexpected CSV header: {header}")
            
            for i, row in enumerate(reader):
                if not row:  # Skip empty rows
                    continue
                if len(row) < 2:
                    issues.append(f"Row {i+1} has insufficient columns: {row}")
                    continue
                    
                try:
                    dist_val = float(row[0])
                    ueq_val = float(row[1])
                    
                    if dist_val < 0 or dist_val > length * 1.1:  # Allow 10% tolerance
                        issues.append(f"Distance out of range: {dist_val} at row {i+1}")
                    if ueq_val < 0 or ueq_val > 65000 * 1.05:  # Allow 5% tolerance
                        issues.append(f"u_eq out of range: {ueq_val} at row {i+1}")
                        
                    distances.append(dist_val)
                    u_equ.append(ueq_val)
                except ValueError:
                    issues.append(f"Invalid numeric value in row {i+1}: {row}")
        
        # Verify data lengths
        if len(grayscale) != len(u_equ):
            issues.append(f"Data length mismatch: grayscale({len(grayscale)}) vs u_equ({len(u_equ)})")
        elif len(grayscale) == 0:
            issues.append("No data points found in output files")
        
        # Verify u_eq calculation
        u_min = 0
        u_max = 65000
        if grayscale and u_equ:
            calculated_u_equ = [u_min + (g/255) * u_max for g in grayscale]
            max_diff = max(abs(a - b) for a, b in zip(u_equ, calculated_u_equ))
            if max_diff > 0.1:  # Allow small floating point differences
                issues.append(f"u_eq calculation mismatch (max difference: {max_diff:.6f})")
        
        # Verify plot file
        plot_path = os.path.join(OUTPUT_DIR, f"{BASE_NAME}_plot.tiff")
        try:
            with Image.open(plot_path) as img:
                if img.format != 'TIFF':
                    issues.append(f"Plot is not TIFF format: {img.format}")
                width, height = img.size
                if width < 800 or height < 600:
                    issues.append(f"Plot dimensions too small: {width}x{height} (expected min 800x600)")
        except Exception as e:
            issues.append(f"TIFF image verification failed: {str(e)}")
        
    except Exception as e:
        issues.append(f"Data verification error: {str(e)}")
    
    return issues

def main():
    print("Starting comprehensive verification...")
    print(f"Output directory: {OUTPUT_DIR}")
    
    # Step 1: Verify all files exist
    missing_files = verify_file_existence()
    if missing_files:
        print("Error: Missing output files:")
        for f in missing_files:
            print(f" - {f}")
        print("\nVerification failed. Ensure all previous steps executed successfully.")
        return
    
    print("✓ All required files exist")
    
    # Step 2: Verify data consistency
    data_issues = verify_data_consistency()
    if data_issues:
        print("Error: Data consistency issues found:")
        for issue in data_issues:
            print(f" - {issue}")
        
        # Provide troubleshooting suggestions
        print("\nTroubleshooting recommendations:")
        if any("u_eq" in issue for issue in data_issues):
            print("1. Check u_eq calculation formula in py1.py: u_min + (gray_value/255) * u_max")
        if any("TIFF" in issue for issue in data_issues):
            print("2. Verify matplotlib TIFF export capability: install Pillow library if missing")
            print("3. Check plot generation code in py1.py uses correct DPI settings")
        if any("length" in issue for issue in data_issues):
            print("4. Validate distance calculation in py1.py: distance = resolution * pixel_distance")
        
        print("\nVerification failed. Data files contain inconsistencies.")
        return
    
    print("✓ Data files verified and consistent")
    
    # Step 3: Verify report content
    report_issues = verify_report_content()
    if report_issues:
        print("Error: Report content issues found:")
        for issue in report_issues:
            print(f" - {issue}")
        
        print("\nTroubleshooting recommendations:")
        if "Missing section" in "".join(report_issues):
            print("1. Verify report generator includes all required sections: Abstract, Introduction, Methods, Results")
        if "word" in "".join(report_issues):
            print("2. Expand report content with more technical details and analysis")
        if "Figure" in "".join(report_issues):
            print("3. Check report generator code adds the plot image correctly")
            
        print("\nVerification failed. Report does not meet requirements.")
        return
    
    print("✓ Report content verified")
    
    # Final success message
    print("\n" + "="*60)
    print("✅ ALL TASKS COMPLETED SUCCESSFULLY")
    print("="*60)
    print("Output summary:")
    for fname in REQUIRED_FILES:
        file_path = os.path.join(OUTPUT_DIR, fname)
        size = os.path.getsize(file_path) / 1024
        print(f" - {fname}: {size:.2f} KB")
    
    print("\nWorkflow verification complete. All outputs validated.")

if __name__ == "__main__":
    main()
