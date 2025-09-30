import os
import csv
import numpy as np
from docx import Document

# Configuration parameters
INPUT_IMAGE = r"C:\Users\admin\Desktop\Python_proj\datas\T2_IMGS\Li_1.0.png"
OUTPUT_DIR = r"C:\Users\admin\Desktop\Python_proj\ALL_RESULT\DS\T2S2\1.0\backup1"
BASE_NAME = "Li_1.0"

# Expected output files
REQUIRED_FILES = [
    f"{BASE_NAME}_grayscale.csv",
    f"{BASE_NAME}_length.txt",
    f"{BASE_NAME}_data.csv",
    f"{BASE_NAME}_plot.tiff",
    f"{BASE_NAME}_Simulation_Report.docx"
]

def verify_file_existence():
    """Check if all required output files exist"""
    missing = []
    for filename in REQUIRED_FILES:
        path = os.path.join(OUTPUT_DIR, filename)
        if not os.path.exists(path):
            missing.append(path)
    return missing

def verify_report_content():
    """Validate the content and structure of the simulation report"""
    report_path = os.path.join(OUTPUT_DIR, f"{BASE_NAME}_Simulation_Report.docx")
    if not os.path.exists(report_path):
        return ["Report file does not exist"]
    
    doc = Document(report_path)
    issues = []
    
    # Verify section structure
    required_sections = ['Abstract', 'Introduction', 'Methods', 'Results']
    section_headings = [para.text for para in doc.paragraphs if para.style.name.startswith('Heading')]
    
    for section in required_sections:
        if section not in section_headings:
            issues.append(f"Missing section: {section}")
    
    # Verify word count
    total_words = 0
    for para in doc.paragraphs:
        total_words += len(para.text.split())
    
    if total_words < 450:
        issues.append(f"Report too short: {total_words} words (minimum 450 required)")
    
    # Verify figure inclusion
    fig_found = False
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            fig_found = True
            break
            
    if not fig_found:
        issues.append("Figure 1 not found in report")
    
    # Verify data references
    content = "\n".join(para.text for para in doc.paragraphs)
    required_refs = ["Fig. 1", "u_eq", "segment length", "gradient"]
    for ref in required_refs:
        if ref not in content:
            issues.append(f"Missing reference to '{ref}' in report content")
    
    return issues

def verify_data_consistency():
    """Check consistency between data files"""
    issues = []
    
    try:
        # Verify segment length
        with open(os.path.join(OUTPUT_DIR, f"{BASE_NAME}_length.txt"), 'r') as f:
            length = float(f.read().strip())
        
        # Verify grayscale data
        grayscale = []
        with open(os.path.join(OUTPUT_DIR, f"{BASE_NAME}_grayscale.csv"), 'r') as f:
            reader = csv.reader(f)
            for row in reader:
                if row:  # Skip empty rows
                    grayscale.append(int(row[0]))
        
        # Verify u_eq data
        distances, u_equ = [], []
        with open(os.path.join(OUTPUT_DIR, f"{BASE_NAME}_data.csv"), 'r') as f:
            reader = csv.reader(f)
            next(reader)  # Skip header
            for row in reader:
                if row:  # Skip empty rows
                    distances.append(float(row[0]))
                    u_equ.append(float(row[1]))
        
        # Verify data lengths
        if len(grayscale) != len(u_equ):
            issues.append(f"Data length mismatch: grayscale({len(grayscale)}) vs u_equ({len(u_equ)})")
        
        # Verify u_eq calculation
        u_min = 0
        u_max = 65000
        calculated_u_equ = [u_min + (g/255) * u_max for g in grayscale]
        
        if not np.allclose(u_equ, calculated_u_equ, atol=1.0):
            issues.append("u_eq values do not match calculated values from grayscale data")
        
        # Verify plot file
        try:
            Image.open(os.path.join(OUTPUT_DIR, f"{BASE_NAME}_plot.tiff"))
        except Exception as e:
            issues.append(f"TIFF image verification failed: {str(e)}")
        
    except Exception as e:
        issues.append(f"Data verification error: {str(e)}")
    
    return issues

def main():
    print("Starting comprehensive verification...")
    print(f"Output directory: {OUTPUT_DIR}")
    
    # Step 1: Verify all files exist
    missing_files = verify_file_existence()
    if missing_files:
        print("Error: Missing output files:")
        for f in missing_files:
            print(f" - {f}")
        print("\nVerification failed. Ensure py1.py and report generator executed successfully.")
        return
    
    print("All required files exist")
    
    # Step 2: Verify data consistency
    data_issues = verify_data_consistency()
    if data_issues:
        print("Error: Data consistency issues found:")
        for issue in data_issues:
            print(f" - {issue}")
        print("\nVerification failed. Data files contain inconsistencies.")
        return
    
    print("Data files verified and consistent")
    
    # Step 3: Verify report content
    report_issues = verify_report_content()
    if report_issues:
        print("Error: Report content issues found:")
        for issue in report_issues:
            print(f" - {issue}")
        print("\nVerification failed. Report does not meet requirements.")
        return
    
    print("Report content verified")
    
    # Final success message
    print("\n" + "="*60)
    print("ALL TASKS COMPLETED SUCCESSFULLY")
    print("="*60)
    print("Output summary:")
    for fname in REQUIRED_FILES:
        file_path = os.path.join(OUTPUT_DIR, fname)
        size = os.path.getsize(file_path) / 1024
        print(f" - {fname}: {size:.2f} KB")
    
    print("\nWorkflow verification complete")

if __name__ == "__main__":
    main()
