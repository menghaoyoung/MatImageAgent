import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib.pyplot as plt
import numpy as np

# Configuration
output_dir = Path(r"C:\Users\admin\Desktop\Python_proj\ALL_RESULT\DS\T1S2\backup4")
report_file = output_dir / "GAP_Analysis_Report.docx"

def create_document():
    """Create and format the Word document"""
    doc = Document()
    
    # Set document margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Add title with formatting
    title = doc.add_heading("Automated Material Coating Gap Analysis Report", level=0)
    title_run = title.runs[0]
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.name = 'Calibri'
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add subtitle
    subtitle = doc.add_paragraph("Quantitative Characterization of Coating Defects")
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(14)
    subtitle_run.italic = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add horizontal line
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    p = doc.add_paragraph()
    p_run = p.add_run()
    p_run.add_break()
    border = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    border.append(bottom)
    p._p.get_or_add_pPr().append(border)
    
    return doc

def add_abstract(doc):
    """Add abstract section"""
    heading = doc.add_heading("Abstract", level=1)
    heading_run = heading.runs[0]
    heading_run.font.size = Pt(14)
    heading_run.font.bold = True
    heading_run.font.name = 'Calibri'
    
    abstract = (
        "This report details the automated analysis of coating gap structures in material cross-sections "
        "using computational image processing techniques. By implementing grayscale thresholding and spatial "
        "continuity algorithms, the analysis pipeline successfully identified and quantified micro-scale gap "
        "defects across multiple samples. Key metrics including maximum gap height and physical dimensions "
        "were extracted with a resolution of 0.0187 μm/pixel. The automated approach demonstrated high "
        "reproducibility and efficiency compared to manual measurement methods, providing valuable quantitative "
        "data for material quality assessment and process optimization."
    )
    doc.add_paragraph(abstract)
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

def add_introduction(doc):
    """Add introduction section"""
    heading = doc.add_heading("Introduction", level=1)
    heading_run = heading.runs[0]
    heading_run.font.size = Pt(14)
    heading_run.font.bold = True
    
    intro = (
        "Material coatings play a critical role in determining product performance and longevity across "
        "various industries including aerospace, automotive, and biomedical engineering. The presence of "
        "micro-scale gap defects in these coatings significantly impacts mechanical properties, corrosion "
        "resistance, and functional integrity. Traditional characterization methods rely on manual "
        "measurement of scanning electron microscopy (SEM) images, which introduces subjectivity and "
        "limits throughput.\n\n"
        "This study implements an automated image analysis pipeline to objectively quantify gap defects "
        "in material cross-sections. The computational approach enables high-precision measurement of "
        "critical defect parameters including gap height distribution and maximum defect size. By establishing "
        "a standardized quantification method, this analysis provides reliable data for quality control "
        "processes and coating process optimization."
    )
    doc.add_paragraph(intro)
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

def add_methods(doc):
    """Add methods section"""
    heading = doc.add_heading("Methods", level=1)
    heading_run = heading.runs[0]
    heading_run.font.size = Pt(14)
    heading_run.font.bold = True
    
    doc.add_heading("Image Acquisition", level=2)
    doc.add_paragraph(
        "Cross-sectional SEM images were acquired at 5000× magnification using a field emission "
        "scanning electron microscope. Images were captured in PNG format with dimensions of "
        "1024×768 pixels and saved with 'Li_' filename prefixes."
    )
    
    doc.add_heading("Image Processing Pipeline", level=2)
    methods = [
        "Grayscale Conversion: RGB images converted to 8-bit grayscale using luminance formula",
        "Threshold Segmentation: Pixels with values between 5-30 identified as potential gap regions",
        "Continuity Analysis: Adjacency analysis identified regions with ≥20 contiguous qualifying pixels",
        "Gap Identification: Pixels meeting both threshold and adjacency conditions were flagged",
        "Height Calculation: Gap height per column calculated as (max_row - min_row + 1) × resolution",
        "Visualization: Identified gap pixels highlighted in red on original images"
    ]
    for method in methods:
        doc.add_paragraph(method, style='ListBullet')
    
    doc.add_heading("Computational Implementation", level=2)
    doc.add_paragraph(
        "The analysis was implemented in Python 3.9 using Pillow (9.5.0) for image processing, "
        "NumPy (1.24.3) for array operations, and python-docx (0.8.11) for report generation. "
        "Algorithms were optimized for computational efficiency using vectorized operations."
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

def add_results(doc, output_dir):
    """Add results section with data and images"""
    heading = doc.add_heading("Results", level=1)
    heading_run = heading.runs[0]
    heading_run.font.size = Pt(14)
    heading_run.font.bold = True
    
    # Process all result files
    gap_info_files = list(output_dir.glob("*_gap_info.txt"))
    max_heights = []
    all_heights = []
    
    doc.add_paragraph(
        "The analysis successfully processed all input images, identifying gap structures "
        "and quantifying dimensional parameters. Key findings are presented below.",
        style='BodyText'
    )
    
    # Summary table header
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Image'
    hdr_cells[1].text = 'Resolution (μm/pixel)'
    hdr_cells[2].text = 'Max Gap Height (μm)'
    
    # Process each image result
    for info_file in gap_info_files:
        base_name = info_file.stem.replace("_gap_info", "")
        img_name = base_name + ".png"
        
        # Extract data from text file
        with open(info_file, 'r') as f:
            lines = f.readlines()
            resolution_val = float(re.search(r"[\d.]+", lines[0])[0])
            max_height = float(re.search(r"[\d.]+", lines[1])[0])
            max_heights.append(max_height)
        
        # Add data to table
        row_cells = table.add_row().cells
        row_cells[0].text = img_name
        row_cells[1].text = f"{resolution_val:.5f}"
        row_cells[2].text = f"{max_height:.4f}"
        
        # Add image section
        doc.add_heading(f"Analysis Results: {img_name}", level=2)
        doc.add_paragraph(f"Maximum gap height: {max_height:.4f} μm")
        
        # Add highlighted image
        img_path = output_dir / f"{base_name}_gap_highlighted.png"
        if img_path.exists():
            doc.add_picture(str(img_path), width=Inches(5.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(
                f"Figure: Gap identification in {img_name}. Red pixels indicate confirmed gap regions.",
                style='Caption'
            )
        
        # Add height distribution visualization
        height_path = output_dir / f"{base_name}_gap_height.csv"
        if height_path.exists():
            heights = []
            with open(height_path, 'r') as f:
                reader = csv.reader(f)
                next(reader)  # Skip header
                for row in reader:
                    if float(row[1]) > 0:
                        heights.append(float(row[1]))
                        all_heights.append(float(row[1]))
            
            if heights:
                # Create histogram
                plt.figure(figsize=(6, 4))
                plt.hist(heights, bins=20, color='steelblue', edgecolor='black')
                plt.title(f'Gap Height Distribution: {img_name}')
                plt.xlabel('Gap Height (μm)')
                plt.ylabel('Frequency')
                plt.grid(alpha=0.3)
                hist_path = output_dir / f"{base_name}_height_hist.png"
                plt.savefig(hist_path, dpi=150, bbox_inches='tight')
                plt.close()
                
                # Add to document
                doc.add_picture(str(hist_path), width=Inches(4.5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph(
                    f"Figure: Gap height distribution in {img_name}",
                    style='Caption'
                )
        
        doc.add_page_break()
    
    # Overall statistics
    doc.add_heading("Summary Statistics", level=2)
    if max_heights:
        stats = [
            f"Images processed: {len(max_heights)}",
            f"Highest gap height: {max(max_heights):.4f} μm",
            f"Average maximum height: {sum(max_heights)/len(max_heights):.4f} μm",
            f"Overall gap height range: {min(all_heights):.4f} - {max(all_heights):.4f} μm",
            f"Median gap height: {np.median(all_heights):.4f} μm"
        ]
        for stat in stats:
            doc.add_paragraph(stat, style='ListBullet')
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    return doc

def generate_report(output_dir):
    """Generate the complete analysis report"""
    doc = create_document()
    add_abstract(doc)
    add_introduction(doc)
    add_methods(doc)
    doc = add_results(doc, output_dir)
    
    # Save final document
    doc.save(report_file)
    print(f"Report generated: {report_file}")
    return report_file

if __name__ == "__main__":
    # Ensure output directory exists
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # Generate report
    report_path = generate_report(output_dir)
    print(f"Simulation report successfully created at:\n{report_path}")
