import subprocess
import os
import glob
import time
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def run_py1_and_verify():
    """Run py1.py with resolution parameter and verify output files"""
    # Step 1: Run py1.py with re=0.0187
    print("Executing py1.py with re=0.0187...")
    process = subprocess.Popen(['python', 'py1.py', '-re=0.0187'])
    process.wait()
    
    # Step 2: Verify output files
    input_dir = r"C:\Users\admin\Desktop\Python_proj\distance_analysis_new\Images"
    output_dir = r"C:\Users\admin\Desktop\Python_proj\ALL_RESULT\DS\T1S2\backup3"
    
    valid_extensions = ('.png', '.jpg', '.jpeg')
    input_images = [f for f in os.listdir(input_dir) 
                  if f.startswith("Li_") and f.lower().endswith(valid_extensions)]
    
    all_files_exist = True
    for img in input_images:
        base = os.path.splitext(img)[0]
        required_files = {
            'CSV (Analysis)': f"{base}_gap_analysis.csv",
            'CSV (Height)': f"{base}_gap_height.csv",
            'TXT': f"{base}_summary.txt",
            'Image': f"{base}_gap_highlight.png"
        }
        
        for file_type, filename in required_files.items():
            path = os.path.join(output_dir, filename)
            if not os.path.exists(path):
                print(f"Missing {file_type} file: {filename}")
                all_files_exist = False
    
    if all_files_exist:
        print("Calculation successful")
        return output_dir, input_images
    else:
        print("Verification failed - some output files are missing")
        exit(1)

def generate_report(output_dir, image_bases):
    """Generate Word document report based on output files"""
    doc = Document()
    
    # Title and meta information
    title = doc.add_heading('Gap Analysis Simulation Report', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Abstract section
    doc.add_heading('Abstract', level=1)
    abstract_text = (
        "This report details the analysis of gap formations in sample images using computational methods. "
        "The analysis focused on identifying gap pixels based on grayscale characteristics and quantifying "
        "physical dimensions of detected gaps. The results provide quantitative measurements of gap heights "
        "across various sample images, which are crucial for material quality assessment."
    )
    doc.add_paragraph(abstract_text)
    
    # Introduction section
    doc.add_heading('Introduction', level=1)
    intro_text = (
        "Gap analysis in material science involves identifying microscopic separations that impact material "
        "integrity. This simulation aims to automate gap detection using image processing techniques. The "
        "primary objectives were to: (1) Develop an algorithm for reliable gap pixel identification, "
        "(2) Quantify physical dimensions of detected gaps, and (3) Generate comprehensive reports for "
        "quality assessment. The analysis was performed on SEM images of material samples."
    )
    doc.add_paragraph(intro_text)
    
    # Methods section
    doc.add_heading('Methods', level=1)
    methods_text = [
        "Image Processing Pipeline:",
        "1. Images were loaded using PIL (Pillow) library and converted to grayscale",
        "2. Each pixel was evaluated against two criteria:",
        "   a) Grayscale value between 5-30 (inclusive)",
        "   b) Adjacency to a continuous 20-pixel segment meeting grayscale criteria",
        "3. Identified gap pixels were flagged and recorded",
        "4. Gap heights were calculated per column using physical resolution parameter",
        "5. Outputs included CSV files, summary statistics, and annotated images",
        "",
        "Technical Implementation:",
        "- Python 3.9 with Pillow, numpy, and docx libraries",
        "- Resolution parameter: 0.0187 μm/pixel",
        f"- Processed {len(image_bases)} sample images"
    ]
    for text in methods_text:
        doc.add_paragraph(text)
    
    # Results section
    doc.add_heading('Results', level=1)
    doc.add_paragraph("The analysis revealed the following key findings across all sample images:")
    
    # Add results table
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Image'
    hdr_cells[1].text = 'Physical Resolution (μm/pixel)'
    hdr_cells[2].text = 'Max Gap Height (μm)'
    
    # Process all summary files
    max_heights = []
    for base in image_bases:
        txt_file = os.path.join(output_dir, f"{base}_summary.txt")
        try:
            with open(txt_file, 'r') as f:
                content = f.readlines()
                resol = content[0].split(':')[-1].strip()
                max_height = content[1].split(':')[-1].strip()
                max_heights.append(float(max_height.split()[0]))
        except Exception as e:
            print(f"Error processing {txt_file}: {str(e)}")
            resol = "N/A"
            max_height = "N/A"
        
        row_cells = table.add_row().cells
        row_cells[0].text = base
        row_cells[1].text = resol
        row_cells[2].text = max_height
    
    # Add statistical summary
    if max_heights:
        overall_max = max(max_heights)
        doc.add_paragraph(
            f"Key Statistic: The maximum gap height observed across all images was {overall_max:.4f} μm. "
            "This measurement represents the largest vertical gap dimension detected in the sample set."
        )
    
    # Add annotated images
    doc.add_heading('Annotated Gap Visualization', level=2)
    doc.add_paragraph(
        "The following images highlight detected gap pixels in red. These visualizations "
        "demonstrate the algorithm's ability to identify continuous gap formations:"
    )
    
    for base in image_bases:
        img_path = os.path.join(output_dir, f"{base}_gap_highlight.png")
        if os.path.exists(img_path):
            doc.add_heading(base, level=3)
            doc.add_picture(img_path, width=Cm(12))
            doc.add_paragraph(f"Figure: Gap visualization for {base} (Red pixels indicate gap regions)")
    
    # Save final document
    report_path = os.path.join(output_dir, "Gap_Analysis_Report.docx")
    doc.save(report_path)
    print(f"Report generated: {report_path}")
    return report_path

if __name__ == "__main__":
    # Step 1: Run and verify py1.py
    output_dir, input_images = run_py1_and_verify()
    
    # Extract base names without extensions
    image_bases = [os.path.splitext(img)[0] for img in input_images]
    
    # Step 2: Generate report
    report_path = generate_report(output_dir, image_bases)
    print(f"Simulation report successfully created at {report_path}")
