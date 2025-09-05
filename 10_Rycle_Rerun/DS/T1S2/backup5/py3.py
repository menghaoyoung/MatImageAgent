import os
import glob
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def generate_report(output_dir):
    """Generate Word document report from analysis results"""
    doc = Document()
    
    # Add title
    title = doc.add_heading('GAP Pixel Analysis Report', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Abstract section
    doc.add_heading('Abstract', level=1)
    abstract = (
        "This report summarizes the analysis of GAP pixels in lithium battery electrode images. "
        "The study employed image processing techniques to identify GAP regions based on grayscale "
        "characteristics and calculated vertical gap heights. Results demonstrate the methodology's "
        "effectiveness in quantifying microstructural features critical to battery performance and durability."
    )
    doc.add_paragraph(abstract)
    
    # Introduction section
    doc.add_heading('Introduction', level=1)
    intro = (
        "Lithium-ion battery performance is heavily influenced by electrode microstructure uniformity. "
        "GAP regions, characterized by specific grayscale properties in SEM images, indicate potential "
        "structural deficiencies affecting ion transport. This analysis aims to quantify these features "
        "through automated image processing, providing objective metrics for quality assessment."
    )
    doc.add_paragraph(intro)
    
    # Methods section
    doc.add_heading('Methods', level=1)
    methods = [
        "1. Image Acquisition: SEM images of lithium battery electrodes were collected",
        "2. Grayscale Conversion: Images converted to 8-bit grayscale using Pillow",
        "3. GAP Pixel Identification: Pixels meeting dual criteria were flagged:",
        "   - Grayscale value between 5-30 (inclusive)",
        "   - Adjacent to regions with ≥20 contiguous qualifying pixels",
        "4. Height Calculation: Vertical gaps measured per column using:",
        "   GAP_height = (max_row - min_row + 1) × resolution (μm/pixel)",
        "5. Visualization: GAP pixels highlighted in red for visual verification",
        "6. Statistical Reporting: Maximum gap height extracted for quality metrics"
    ]
    for item in methods:
        doc.add_paragraph(item, style='ListBullet')
    
    # Results section
    doc.add_heading('Results', level=1)
    results_intro = (
        "Analysis was performed on multiple electrode samples. Key findings include physical dimension "
        "parameters and maximum gap heights as summarized below. Visual representations demonstrate GAP "
        "distribution characteristics across different sample regions."
    )
    doc.add_paragraph(results_intro)
    
    # Process all result files
    txt_files = glob.glob(os.path.join(output_dir, '*_gap_report.txt'))
    img_files = glob.glob(os.path.join(output_dir, '*_gap_highlight.png'))
    
    # Create table for numerical results
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Sample'
    hdr_cells[1].text = 'Resolution (μm/px)'
    hdr_cells[2].text = 'Max Height (μm)'
    
    # Process each result file
    for txt_file in txt_files:
        base_name = os.path.basename(txt_file).replace('_gap_report.txt', '')
        
        # Extract data from text file
        with open(txt_file, 'r') as f:
            data = f.read().splitlines()
        resolution = data[0].split(': ')[1]
        max_height = data[1].split(': ')[1]
        
        # Add to table
        row_cells = table.add_row().cells
        row_cells[0].text = base_name
        row_cells[1].text = resolution
        row_cells[2].text = max_height
        
        # Add corresponding image
        img_path = next((f for f in img_files if base_name in f), None)
        if img_path:
            doc.add_paragraph(f"GAP distribution in {base_name}:")
            doc.add_picture(img_path, width=Inches(4.0))
            last_paragraph = doc.paragraphs[-1] 
            last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Save final document
    report_path = os.path.join(output_dir, "GAP_Analysis_Report.docx")
    doc.save(report_path)
    return report_path

if __name__ == "__main__":
    output_directory = r"C:\Users\admin\Desktop\Python_proj\ALL_RESULT\DS\T1S2\backup5"
    report_path = generate_report(output_directory)
    print(f"Report generated successfully: {report_path}")

# NO-RUN-PY
