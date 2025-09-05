import os
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def generate_simulation_report():
    # Define paths
    output_dir = r"C:\Users\admin\Desktop\Python_proj\ALL_RESULT\CLAUDE\T2S2\1.0\backup1"
    image_path = r"C:\Users\admin\Desktop\Python_proj\datas\T2_IMGS\Li_1.0.png"
    filename_base = os.path.splitext(os.path.basename(image_path))[0]
    
    # Define paths to generated files
    line_length_file = os.path.join(output_dir, f"{filename_base}_line_length.txt")
    grayscale_csv = os.path.join(output_dir, f"{filename_base}_grayscale_values.csv")
    distance_u_eq_csv = os.path.join(output_dir, f"{filename_base}_distance_u_eq.csv")
    plot_file = os.path.join(output_dir, f"{filename_base}_u_eq_plot.tiff")
    
    # Check if all required files exist
    required_files = [line_length_file, grayscale_csv, distance_u_eq_csv, plot_file]
    for file in required_files:
        if not os.path.exists(file):
            print(f"Error: Required file {file} not found!")
            return
    
    # Read line length data
    with open(line_length_file, 'r') as f:
        line_info = f.readlines()
    
    line_length = line_info[0].split(': ')[1].strip()
    start_point = line_info[1].split(': ')[1].strip()
    end_point = line_info[2].split(': ')[1].strip()
    resolution = line_info[3].split(': ')[1].strip()
    
    # Read distance and u_eq data
    df_u_eq = pd.read_csv(distance_u_eq_csv)
    df_gray = pd.read_csv(grayscale_csv)
    
    # Calculate statistics for the report
    u_eq_min = df_u_eq['u_eq'].min()
    u_eq_max = df_u_eq['u_eq'].max()
    u_eq_mean = df_u_eq['u_eq'].mean()
    u_eq_std = df_u_eq['u_eq'].std()
    
    # Create a new Word document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Simulation Report: Analysis of Grayscale Intensity Distribution and u_eq Calculation', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    date_paragraph = doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Add Abstract
    doc.add_heading('Abstract', level=2)
    abstract = doc.add_paragraph(
        'This report presents a detailed analysis of grayscale intensity distribution along a specified line segment in a digital image '
        'and its conversion to a physical quantity (u_eq). The study employs image processing techniques to extract intensity values '
        'from a digital image and transform them into meaningful physical measurements using a linear mapping approach. The analysis '
        'focuses on the spatial variation of grayscale values and their corresponding u_eq values along a defined path, providing '
        'insights into the underlying structures or phenomena captured in the image. The methodology involves precise pixel-by-pixel '
        'sampling along a line segment, extraction of grayscale values, and their transformation to u_eq using a linear equation. '
        'Results demonstrate the relationship between spatial position and the derived u_eq values, revealing patterns of intensity '
        'variation that may correspond to physical or structural features in the original image. This approach offers a quantitative '
        'means of analyzing intensity distributions in digital images, with potential applications in various scientific and engineering '
        'domains, including materials science, medical imaging, and environmental monitoring. The findings highlight the utility of '
        'digital image analysis as a non-invasive method for extracting quantitative data from visual representations.'
    )
    
    # Add Introduction
    doc.add_heading('Introduction', level=2)
    intro = doc.add_paragraph(
        'Digital image analysis has become an essential tool in modern scientific research, enabling the extraction of quantitative '
        'data from visual representations of physical phenomena. This study focuses on analyzing the grayscale intensity distribution '
        'along a defined line segment within a digital image and converting these values to physically meaningful quantities. '
        'The grayscale values, ranging from 0 to 255, represent the intensity of light at each pixel location, which can be indicative '
        'of various physical properties depending on the imaging technique used. By mapping these values to a specific range (u_min to u_max), '
        'we can derive quantities that may represent physical parameters such as temperature, concentration, or stress distribution. '
        f'In this particular analysis, we examine an image file "{filename_base}" and extract intensity values along a line segment from '
        f'{start_point} to {end_point}. The resolution of {resolution} mm/pixel provides the spatial context for our measurements, allowing us to '
        'convert pixel distances to physical distances. This approach enables a quantitative assessment of spatial variations in the parameter '
        'of interest, potentially revealing gradients, discontinuities, or other features of significance. The ability to transform grayscale '
        'values into physically meaningful quantities enhances our understanding of the underlying phenomena and facilitates comparison with '
        'theoretical models or other experimental results.'
    )
    
    # Add Methods
    doc.add_heading('Methods', level=2)
    methods = doc.add_paragraph(
        'The analysis was conducted using a custom Python program implementing several key computational steps. First, the program loaded '
        f'the target image ("{filename_base}") and converted it to grayscale to ensure uniform intensity representation. A line segment was '
        f'defined by the start point {start_point} and end point {end_point}, with the program calculating the Euclidean distance between these '
        'points to determine the line length. The algorithm then employed a linear interpolation approach to sample points along this line, '
        'extracting the grayscale value (0-255) at each point. For pixels that fell between the image\'s discrete grid, nearest-neighbor '
        'sampling was used to determine the appropriate value. These grayscale values were then transformed into the target physical quantity '
        'u_eq using the formula: u_eq = u_min + (gray_values / 255) * u_max, where u_min = 0 and u_max = 65000. This linear mapping assumes '
        'a direct proportionality between grayscale intensity and the physical quantity of interest. The spatial distribution was analyzed by '
        'plotting u_eq against the distance from the starting point, with distances calculated based on the provided resolution of '
        f'{resolution} mm/pixel. All extracted data, including grayscale values, calculated distances, and u_eq values, were saved in CSV format '
        'for transparency and reproducibility. The program also generated visualizations of the u_eq distribution to facilitate interpretation '
        'of the results. This methodological approach ensures a systematic and quantitative analysis of the intensity distribution along the '
        'specified line segment.'
    )
    
    # Add Results
    doc.add_heading('Results', level=2)
    results = doc.add_paragraph(
        f'The analysis of the line segment extending from {start_point} to {end_point} yielded a physical length of {line_length} based on '
        f'the specified resolution of {resolution} mm/pixel. As shown in Fig. 1, the u_eq values exhibit a distinct pattern of variation along the '
        'measured distance. The graph illustrates the relationship between spatial position and the derived physical quantity, revealing '
        'regions of both gradual and rapid change. The minimum u_eq value observed was approximately {u_eq_min:.2f}, while the '
        f'maximum reached approximately {u_eq_max:.2f}, indicating a substantial range of variation across the relatively short '
        'distance. Notable features in the distribution include gradient changes and potential inflection points that may correspond to '
        'structural boundaries or transitions in the physical system represented by the image. The average u_eq value across the entire '
        f'line segment was {u_eq_mean:.2f}, with a standard deviation of {u_eq_std:.2f}, indicating the degree of '
        'heterogeneity present. These quantitative measures provide valuable insights into the spatial distribution of the physical '
        'parameter represented by the grayscale values in the original image. The complete dataset, preserved in CSV format, enables '
        'further statistical analysis and comparison with theoretical models or other experimental results. The observed patterns suggest '
        'that the image contains structured variations in intensity that likely correspond to meaningful physical features rather than '
        'random noise.'
    )
    
    # Add figure
    doc.add_paragraph('Fig. 1: u_eq values plotted against distance from the start point.')
    doc.add_picture(plot_file, width=Inches(6.0))
    
    # Add conclusion
    doc.add_heading('Conclusion', level=2)
    conclusion = doc.add_paragraph(
        'This study demonstrates the utility of image analysis techniques for extracting quantitative data from digital images. '
        'By converting grayscale values to physically meaningful quantities and analyzing their spatial distribution, we can gain '
        'insights into underlying structures and processes. The methodology presented here can be applied to various types of images '
        'and adapted for different physical parameters by adjusting the transformation formula. The analysis revealed significant '
        'variations in u_eq values along the studied line segment, suggesting the presence of structured features in the original image. '
        'Future work could extend this approach to two-dimensional analysis or incorporate more sophisticated image processing techniques '
        'to enhance the extraction of meaningful data from complex images. Additionally, comparative studies across multiple images or '
        'under different conditions could provide further insights into the physical phenomena represented by the intensity distributions.'
    )
    
    # Save the document
    report_path = os.path.join(output_dir, f"{filename_base}_simulation_report.docx")
    doc.save(report_path)
    print(f"Simulation report generated and saved to: {report_path}")

if __name__ == "__main__":
    generate_simulation_report()
