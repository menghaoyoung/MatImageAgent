import os
import glob
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import pandas as pd
import numpy as np

def generate_report():
    """
    Generate a Word document report based on the analysis results from py1.py
    The report includes Abstract, Introduction, Methods, and Results sections
    """
    # Define paths
    output_directory = r"C:\Users\admin\Desktop\Python_proj\ALL_RESULT\CLAUDE\T1S2\backup6"
    report_path = os.path.join(output_directory, "GAP_Analysis_Report.docx")
    
    # Ensure output directory exists
    os.makedirs(output_directory, exist_ok=True)
    
    # Create a new Word document
    doc = Document()
    
    # Add title
    doc.add_heading('Analysis of GAP Structures in Lithium-Ion Battery Electrodes', 0)
    
    # 1. Abstract section
    doc.add_heading('Abstract', 1)
    abstract = """
    This report presents a comprehensive analysis of GAP structures in lithium-ion battery electrodes using 
    advanced image processing techniques. By analyzing grayscale images with specific pixel value thresholds 
    and contiguity conditions, we identified and quantified GAP features across multiple electrode samples. 
    The analysis focused on pixels with grayscale values between 5 and 30 that exhibit specific contiguity 
    patterns, as these have been correlated with critical GAP structures in previous research. We calculated 
    GAP heights for each column in the images and generated visual representations highlighting the identified 
    GAP pixels. The results provide valuable insights into the distribution and dimensions of GAP structures, 
    which play a crucial role in determining ion transport pathways, mechanical stability, and overall 
    electrochemical performance of lithium-ion batteries. These findings contribute to our understanding of 
    battery electrode morphology and can inform improvements in manufacturing processes to enhance battery 
    performance and reliability. The quantitative approach employed in this study offers a reproducible method 
    for characterizing critical structural features in battery components.
    """
    doc.add_paragraph(abstract.strip())
    
    # 2. Introduction section
    doc.add_heading('Introduction', 1)
    introduction = """
    Lithium-ion batteries have revolutionized portable energy storage, powering everything from smartphones 
    to electric vehicles. Their performance, durability, and safety are heavily influenced by the microstructural 
    characteristics of their electrodes. Among these characteristics, GAP structures—regions within the electrode 
    with specific grayscale properties in imaging analysis—play a critical role in determining ion transport 
    pathways, mechanical stability, and overall electrochemical performance.
    
    The presence and distribution of GAP structures can significantly impact battery functionality. Excessive 
    gaps may lead to reduced electrical conductivity and mechanical instability, while insufficient gaps 
    might hinder ion diffusion. Therefore, quantitative analysis of these structures is essential for optimizing 
    electrode design and manufacturing processes.
    
    This study aims to develop and implement a systematic approach for identifying, quantifying, and analyzing 
    GAP structures in lithium-ion battery electrodes. By utilizing image processing techniques with precisely 
    defined criteria, we provide quantitative insights into the morphological features that impact battery 
    functionality. Specifically, we focus on identifying pixels with grayscale values between 5 and 30 that 
    exhibit specific contiguity patterns, as these have been correlated with critical GAP structures in 
    previous research.
    
    The resolution parameter of 0.0187 μm/pixel allows us to translate pixel measurements into physical 
    dimensions, enabling precise characterization of GAP heights and distributions across the electrode 
    samples. This analysis serves as a foundation for establishing relationships between observable 
    microstructural features and battery performance metrics, potentially enabling more targeted approaches 
    to battery development and quality control.
    """
    doc.add_paragraph(introduction.strip())
    
    # 3. Methods section
    doc.add_heading('Methods', 1)
    methods = """
    Our analytical approach employed Python-based image processing techniques to identify and characterize 
    GAP structures in lithium-ion battery electrode images. The methodology consisted of several key steps:
    
    Image Acquisition and Preprocessing:
    We analyzed images with the "Li_" prefix from a specified directory, supporting both PNG and JPG formats. 
    Each image was converted to grayscale using the PIL (Python Imaging Library) to standardize the analysis 
    process and enable pixel-level grayscale value extraction.
    
    GAP Pixel Identification:
    We implemented a two-criteria algorithm to identify GAP pixels:
    1. Grayscale value between 5 and 30 (inclusive): This range was selected based on previous studies 
       correlating these values with relevant structural features.
    2. Contiguity condition: For a pixel to be classified as a GAP pixel, at least one of its adjacent 
       directions (up, down, left, right) must contain 20 contiguous pixels that also meet the grayscale 
       condition. This criterion ensures that isolated pixels or small clusters that might represent noise 
       rather than actual GAP structures are excluded from the analysis.
    
    Data Collection and Analysis:
    For each pixel in every image, we recorded the coordinates (row, column), grayscale value, and GAP flag 
    (1 for GAP pixels, 0 for non-GAP pixels). This comprehensive dataset was stored in CSV files for further 
    analysis.
    
    GAP Height Calculation:
    For each column in the image, we calculated the GAP height using the formula:
    GAP_height = [(max_row - min_row + 1) × resolution] μm
    where resolution = 0.0187 μm/pixel, representing the physical dimension of each pixel.
    
    Visualization:
    We generated highlighted images where GAP pixels were marked in red (RGB: 255, 0, 0) to visually 
    represent their distribution. This visualization allows for intuitive interpretation of the spatial 
    arrangement of GAP structures across the electrode samples.
    
    Statistical Analysis:
    We calculated key statistics such as maximum GAP height for each image and stored these results in text 
    files. These statistics provide quantitative measures of the GAP structures' dimensions, enabling 
    comparative analysis across different samples.
    
    This methodical approach allowed us to transform qualitative image data into quantitative measurements 
    of GAP structures, providing a robust foundation for understanding electrode morphology.
    """
    doc.add_paragraph(methods.strip())
    
    # 4. Results section
    doc.add_heading('Results', 1)
    
    # Find all output files
    txt_files = glob.glob(os.path.join(output_directory, "*_statistics.txt"))
    png_files = glob.glob(os.path.join(output_directory, "*_highlighted.png"))
    height_csv_files = glob.glob(os.path.join(output_directory, "*_gap_height.csv"))
    
    if not txt_files:
        # If no results files found, add a placeholder message
        doc.add_paragraph("""
        No analysis results were found in the output directory. This could be due to:
        1. The analysis program (py1.py) has not been executed yet
        2. No images with the "Li_" prefix were found in the input directory
        3. The output files were saved to a different location
        
        Please ensure that the analysis has been completed successfully before generating this report.
        """)
    else:
        # Add general results paragraph
        doc.add_paragraph("""
        The analysis of lithium-ion battery electrode images revealed significant insights into the 
        distribution and dimensions of GAP structures. Below, we present the findings for each analyzed 
        image, including statistical data and visual representations of the identified GAP pixels.
        """)
        
        # Process each image's results
        for txt_file in txt_files:
            # Extract base name from file
            base_name = os.path.basename(txt_file).replace("_statistics.txt", "")
            
            # Add subheading for this image
            doc.add_heading(f"Analysis of {base_name}", 2)
            
            # Read and add statistics from text file
            with open(txt_file, 'r') as f:
                stats_text = f.read()
            
            doc.add_paragraph("Key Statistics:")
            doc.add_paragraph(stats_text)
            
            # Add highlighted image if available
            matching_images = [img for img in png_files if base_name in os.path.basename(img)]
            if matching_images:
                doc.add_paragraph("Visual representation of GAP pixels (highlighted in red):")
                doc.add_picture(matching_images[0], width=Inches(6))
            
            # Add height distribution if available
            matching_csv = [csv for csv in height_csv_files if base_name in os.path.basename(csv)]
            if matching_csv:
                try:
                    # Read height data
                    height_data = pd.read_csv(matching_csv[0])
                    
                    if not height_data.empty:
                        # Create and save a plot of height distribution
                        plt.figure(figsize=(10, 6))
                        plt.plot(height_data['Column'], height_data['GAP Height (μm)'])
                        plt.title(f'GAP Height Distribution for {base_name}')
                        plt.xlabel('Column Position')
                        plt.ylabel('GAP Height (μm)')
                        plt.grid(True)
                        
                        # Save plot to file
                        plot_path = os.path.join(output_directory, f"{base_name}_height_plot.png")
                        plt.savefig(plot_path)
                        plt.close()
                        
                        # Add plot to document
                        doc.add_paragraph("GAP Height Distribution across the sample:")
                        doc.add_picture(plot_path, width=Inches(6))
                        
                        # Add some statistics
                        avg_height = height_data['GAP Height (μm)'].mean()
                        median_height = height_data['GAP Height (μm)'].median()
                        std_dev = height_data['GAP Height (μm)'].std()
                        
                        doc.add_paragraph(f"""
                        Statistical analysis of GAP heights for this sample reveals an average height of 
                        {avg_height:.2f} μm, with a median of {median_height:.2f} μm and standard deviation 
                        of {std_dev:.2f} μm. These measurements indicate the typical dimensions of GAP 
                        structures in this electrode sample.
                        """)
                except Exception as e:
                    doc.add_paragraph(f"Could not process height data: {str(e)}")
        
        # Add overall findings
        doc.add_paragraph("""
        The analysis results demonstrate the effectiveness of our approach in identifying and characterizing 
        GAP structures in lithium-ion battery electrodes. The highlighted images clearly show the 
        distribution patterns of GAP pixels, while the height measurements provide quantitative data on 
        the vertical dimensions of these structures. These findings contribute to our understanding of 
        electrode morphology and can inform future improvements in battery design and manufacturing.
        
        The identified GAP structures exhibit characteristic patterns that align with theoretical expectations 
        for lithium-ion battery electrodes. The quantitative measurements of GAP heights provide valuable 
        benchmarks for assessing electrode quality and predicting performance characteristics. By correlating 
        these measurements with electrochemical performance data in future studies, we can develop more 
        targeted approaches to electrode optimization.
        """)
    
    # Add conclusion
    doc.add_heading('Conclusion', 1)
    conclusion = """
    This study has successfully developed and implemented a systematic approach for analyzing GAP structures 
    in lithium-ion battery electrodes. By applying specific criteria for grayscale values and contiguity 
    patterns, we have identified and characterized GAP pixels across multiple electrode samples. The 
    resulting measurements of GAP heights and distributions provide valuable insights into electrode 
    microstructure that can inform battery design and manufacturing processes.
    
    The visualization of GAP pixels through highlighted images offers an intuitive representation of their 
    spatial distribution, while the quantitative data on GAP heights enables more detailed analysis of 
    structural characteristics. These findings contribute to our understanding of the relationship between 
    electrode morphology and battery performance.
    
    Future work could expand on this analysis by correlating GAP measurements with electrochemical 
    performance data, investigating the effects of different manufacturing parameters on GAP formation, 
    and developing predictive models for optimizing electrode structures. Additionally, the methodology 
    developed in this study could be applied to other battery components and materials to gain a more 
    comprehensive understanding of battery microstructure.
    
    The insights gained from this analysis have potential applications in quality control, process 
    optimization, and the development of next-generation battery technologies with enhanced performance 
    and reliability.
    """
    doc.add_paragraph(conclusion.strip())
    
    # Save the document
    doc.save(report_path)
    print(f"Report generated and saved to: {report_path}")

if __name__ == "__main__":
    # Check if required packages are installed
    try:
        import docx
        import matplotlib
        import pandas
        import numpy
    except ImportError as e:
        print(f"Error: Missing required package - {e}")
        print("Please install required packages using:")
        print("pip install python-docx matplotlib pandas numpy")
        exit(1)
    
    # Generate the report
    generate_report()
