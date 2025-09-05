import os
import glob
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
import time

def generate_report():
    # Define paths
    output_directory = r"C:\Users\admin\Desktop\Python_proj\ALL_RESULT\CLAUDE\T1S2\backup3"
    report_path = os.path.join(output_directory, "GAP_Analysis_Report.docx")
    
    # Create output directory if it doesn't exist
    os.makedirs(output_directory, exist_ok=True)
    
    # Check if output files exist
    txt_files = glob.glob(os.path.join(output_directory, "*_stats.txt"))
    if not txt_files:
        print("No output files found. Please ensure py1.py has been run successfully.")
        return False
    
    # Create a new Document
    doc = Document()
    
    # Set font for the entire document
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    # Add a title
    title = doc.add_heading('Gap Analysis in Lithium-Ion Battery Electrodes', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add Abstract section
    doc.add_heading('Abstract', 1)
    abstract = """
    This report presents a detailed analysis of electrode gaps (GAPs) in lithium-ion battery components. 
    Using advanced image processing techniques, we identified and quantified GAP structures based on grayscale 
    intensity patterns. The analysis focused on pixels with grayscale values between 5 and 30 that exhibited 
    specific continuity patterns, indicating structural gaps in the electrode material. These gaps can significantly 
    impact battery performance, affecting ionic conductivity, mechanical stability, and overall electrochemical behavior. 
    Our findings provide quantitative measurements of GAP heights across different regions, offering insights into 
    manufacturing quality and potential performance implications. The maximum GAP height observed in the samples 
    provides a critical metric for assessing electrode integrity. This analysis methodology provides a valuable tool 
    for quality control and optimization of battery electrode fabrication processes, potentially contributing to the 
    development of more efficient and reliable energy storage solutions.
    """
    doc.add_paragraph(abstract)
    
    # Add Introduction section
    doc.add_heading('Introduction', 1)
    introduction = """
    Lithium-ion batteries represent the cornerstone of modern energy storage technologies, powering everything from 
    portable electronics to electric vehicles and renewable energy systems. The performance, safety, and longevity of 
    these batteries depend significantly on the microstructural properties of their electrodes. Among these properties, 
    the presence and distribution of gaps within electrode materials can substantially impact battery functionality.
    
    Gaps in electrode materials can affect several aspects of battery performance. They may disrupt the continuous 
    pathways needed for efficient ion transport, potentially increasing internal resistance and reducing power capability. 
    Conversely, some degree of porosity is necessary to accommodate electrolyte infiltration and facilitate ion movement. 
    Understanding the precise nature of these gaps—their size, distribution, and morphology—is therefore essential for 
    optimizing electrode design and manufacturing processes.
    
    The quantification of GAP structures presents significant challenges due to their complex morphology and the 
    limitations of conventional imaging techniques. This study employs advanced image processing algorithms to identify 
    and characterize gaps in lithium-ion battery electrode materials. By analyzing grayscale values and continuity 
    patterns in high-resolution images, we can detect regions that likely represent structural gaps. The quantification 
    of these features provides valuable insights into electrode quality and potential performance implications.
    
    The present analysis focuses specifically on identifying GAP structures based on grayscale intensity thresholds and 
    continuity patterns. By defining GAP pixels as those with grayscale values between 5 and 30 that exhibit specific 
    continuity characteristics, we aim to isolate and quantify these important microstructural features. The results 
    of this analysis can inform manufacturing process optimization and quality control procedures, ultimately contributing 
    to the development of more efficient and reliable lithium-ion batteries.
    """
    doc.add_paragraph(introduction)
    
    # Add Methods section
    doc.add_heading('Methods', 1)
    methods = """
    Our analysis methodology leverages computer vision and image processing techniques to identify and quantify gaps 
    in electrode materials. The approach consists of several key steps:
    
    Image Acquisition and Preprocessing: High-resolution images of electrode samples were collected with filenames 
    prefixed with "Li_". These images were converted to grayscale to simplify analysis and ensure consistent processing.
    
    GAP Identification Algorithm: We developed a two-stage algorithm to identify pixels representing gaps:
    1. Grayscale Thresholding: Pixels with grayscale values between 5 and 30 (inclusive) were identified as potential 
       gap candidates. This range was determined based on the typical intensity profile of gap regions in the electrode material.
    2. Continuity Analysis: For each candidate pixel, we examined adjacent pixels (up, down, left, right) to determine if 
       at least one direction contained 20 contiguous pixels also meeting the grayscale threshold condition. This step helps 
       distinguish actual gaps from noise or isolated artifacts.
    
    Quantification: For each identified gap region, we calculated the GAP height per column as:
    GAP_height = [(max_row - min_row + 1) × resolution] μm
    where resolution is the physical dimension parameter (0.0187 μm/pixel in this analysis).
    
    Visualization: We generated highlighted images where identified gap pixels were marked in red (RGB: 255, 0, 0) to 
    facilitate visual inspection and validation of the algorithm's performance.
    
    Data Output: Comprehensive results were stored in CSV files containing pixel-level data and gap height measurements, 
    along with summary statistics in text files. For each image, we generated:
    1. A pixel-level CSV file containing coordinates, grayscale values, and GAP flags
    2. A column-level CSV file containing GAP height measurements
    3. A text file with summary statistics including the maximum GAP height
    4. A highlighted image showing the identified GAP pixels
    
    This methodology enables objective, quantitative assessment of gap structures in electrode materials, providing 
    insights that can inform manufacturing process optimization and quality control.
    """
    doc.add_paragraph(methods)
    
    # Add Results section
    doc.add_heading('Results', 1)
    
    # Find all stats.txt files
    txt_files = glob.glob(os.path.join(output_directory, "*_stats.txt"))
    
    # Find all highlighted images
    img_files = glob.glob(os.path.join(output_directory, "*_highlighted.png"))
    
    # Find all gap height CSV files
    height_csv_files = glob.glob(os.path.join(output_directory, "*_gap_height.csv"))
    
    results_intro = """
    Our analysis revealed significant insights into the distribution and characteristics of gaps within the electrode materials. 
    The following sections present key findings from our image processing and quantification efforts.
    """
    doc.add_paragraph(results_intro)
    
    # Add data from TXT files
    if txt_files:
        doc.add_heading('Gap Statistics', 2)
        stats_table = doc.add_table(rows=1, cols=3)
        stats_table.style = 'Table Grid'
        hdr_cells = stats_table.rows[0].cells
        hdr_cells[0].text = 'Sample'
        hdr_cells[1].text = 'Resolution (μm/pixel)'
        hdr_cells[2].text = 'Maximum GAP Height (μm)'
        
        for txt_file in txt_files:
            image_name = os.path.basename(txt_file).replace('_stats.txt', '')
            
            # Read statistics from the text file
            resolution = "N/A"
            max_height = "N/A"
            
            try:
                with open(txt_file, 'r') as f:
                    lines = f.readlines()
                    for line in lines:
                        if "Physical dimension parameter:" in line:
                            resolution = line.split(":")[1].strip().replace(" μm/pixel", "")
                        elif "Maximum GAP height:" in line:
                            max_height = line.split(":")[1].strip().replace(" μm", "")
                
                # Add row to table
                row_cells = stats_table.add_row().cells
                row_cells[0].text = image_name
                row_cells[1].text = resolution
                row_cells[2].text = max_height
                
            except Exception as e:
                print(f"Error reading {txt_file}: {e}")
        
        doc.add_paragraph("\nThe table above summarizes the key statistics for each analyzed sample. The resolution parameter represents the physical dimension of each pixel in micrometers, while the maximum GAP height indicates the largest vertical gap detected in each sample. These metrics provide valuable insights into the structural characteristics of the electrode materials.")
    
    # Add highlighted images
    if img_files:
        doc.add_heading('Gap Visualization', 2)
        doc.add_paragraph("The following images show the identified gap regions (highlighted in red) in the electrode samples. These visualizations help illustrate the distribution and morphology of gaps across different samples.")
        
        # Limit to first 4 images to keep report concise
        for i, img_file in enumerate(img_files[:4]):
            image_name = os.path.basename(img_file).replace('_highlighted.png', '')
            
            # Add subheading for each image
            doc.add_heading(f"Sample: {image_name}", 3)
            
            # Add description
            doc.add_paragraph(f"Visualization of identified gaps in sample {image_name}. Red pixels indicate regions classified as gaps based on grayscale value and continuity criteria.")
            
            # Add the image
            try:
                doc.add_picture(img_file, width=Inches(6))
                # Add a paragraph after the image for spacing
                doc.add_paragraph()
            except Exception as e:
                doc.add_paragraph(f"Error adding image: {e}")
    
    # Add analysis of gap heights if CSV files exist
    if height_csv_files:
        doc.add_heading('Gap Height Analysis', 2)
        doc.add_paragraph("The gap height analysis provides quantitative measurements of the vertical extent of gaps in the electrode materials. These measurements are crucial for understanding the potential impact of gaps on battery performance and for identifying areas for manufacturing process improvement.")
        
        # Create a summary table for gap heights
        height_table = doc.add_table(rows=1, cols=5)
        height_table.style = 'Table Grid'
        hdr_cells = height_table.rows[0].cells
        hdr_cells[0].text = 'Sample'
        hdr_cells[1].text = 'Average Height (μm)'
        hdr_cells[2].text = 'Maximum Height (μm)'
        hdr_cells[3].text = 'Minimum Height (μm)'
        hdr_cells[4].text = 'Standard Deviation (μm)'
        
        for csv_file in height_csv_files[:4]:  # Limit to first 4 files
            image_name = os.path.basename(csv_file).replace('_gap_height.csv', '')
            
            try:
                # Read the CSV file
                df = pd.read_csv(csv_file)
                
                if not df.empty and 'GAP Height (μm)' in df.columns:
                    # Calculate statistics
                    avg_height = df['GAP Height (μm)'].mean()
                    max_height = df['GAP Height (μm)'].max()
                    min_height = df['GAP Height (μm)'].min()
                    std_height = df['GAP Height (μm)'].std()
                    
                    # Add row to table
                    row_cells = height_table.add_row().cells
                    row_cells[0].text = image_name
                    row_cells[1].text = f"{avg_height:.2f}"
                    row_cells[2].text = f"{max_height:.2f}"
                    row_cells[3].text = f"{min_height:.2f}"
                    row_cells[4].text = f"{std_height:.2f}"
            except Exception as e:
                print(f"Error analyzing {csv_file}: {e}")
    
    # Add conclusion
    doc.add_heading('Conclusion', 1)
    conclusion = """
    This analysis has provided valuable insights into the microstructural characteristics of gaps in lithium-ion battery 
    electrode materials. The identified gaps, characterized by specific grayscale intensity patterns and continuity 
    properties, represent important structural features that can influence battery performance.
    
    The quantification of gap heights across different regions offers a metric for assessing manufacturing consistency 
    and potential impact on electrochemical behavior. Regions with excessive gap heights may indicate areas of concern 
    for electrode integrity and ion transport efficiency. The maximum gap heights observed in our samples provide a 
    critical reference point for quality control and process optimization.
    
    The visualization of gap distributions through highlighted images enables qualitative assessment of gap morphology 
    and spatial arrangement. These visualizations complement the quantitative data, offering a more comprehensive 
    understanding of electrode microstructure.
    
    The methodology developed in this study provides a robust framework for automated gap analysis in electrode materials, 
    offering a valuable tool for quality control and process optimization. Future work could extend this approach to 
    include three-dimensional analysis and correlation with electrochemical performance metrics.
    
    Understanding and controlling these microstructural features is essential for advancing lithium-ion battery technology 
    and addressing the growing demands for high-performance energy storage solutions. By optimizing electrode microstructure, 
    we can potentially improve battery capacity, power capability, cycle life, and safety, contributing to the broader 
    goals of sustainable energy transition.
    """
    doc.add_paragraph(conclusion)
    
    # Save the document
    try:
        doc.save(report_path)
        print(f"Report generated and saved to {report_path}")
        return True
    except Exception as e:
        print(f"Error saving report: {e}")
        return False

if __name__ == "__main__":
    print("Generating GAP analysis report...")
    start_time = time.time()
    success = generate_report()
    end_time = time.time()
    
    if success:
        print(f"Report generation completed in {end_time - start_time:.2f} seconds")
    else:
        print("Report generation failed")
