import os
import glob
from docx import Document
from docx.shared import Inches
import pandas as pd
from datetime import datetime
import matplotlib.pyplot as plt
import numpy as np

def generate_report():
    # Define output directory where the files were saved
    output_dir = r"C:\Users\admin\Desktop\Python_proj\distance_analysis_new\ALL_RESULT\CLAUDE\T1S1\backup2"
    
    # Create a new Document
    doc = Document()
    
    # Add title
    doc.add_heading('Analysis of GAP Pixels in Lithium-Ion Battery Images', 0)
    
    # Add date
    doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    
    # Abstract section
    doc.add_heading('Abstract', 1)
    abstract = """
    This report presents a comprehensive analysis of GAP pixels identified in lithium-ion battery images. 
    GAP pixels are defined by specific criteria: grayscale values between 5 and 30 (inclusive) and adjacency 
    to at least one pixel that has 20 contiguous pixels meeting the same grayscale condition. Using image 
    processing techniques implemented in Python with the PIL library, we analyzed multiple battery images 
    with the "Li_" prefix. The analysis generated per-pixel data stored in CSV files and visual representations 
    highlighting GAP pixels in red. These pixels may correspond to critical structural features or potential 
    defects in battery materials that could influence performance characteristics. The findings provide 
    valuable insights into the microstructural properties of lithium-ion batteries, potentially informing 
    future research in battery design, manufacturing processes, and quality control methodologies. This 
    automated approach demonstrates the utility of computational image analysis in materials science, 
    offering a scalable method for identifying specific features across large datasets of battery images.
    """
    doc.add_paragraph(abstract)
    
    # Introduction section
    doc.add_heading('Introduction', 1)
    introduction = """
    Lithium-ion batteries represent a cornerstone technology in modern energy storage solutions, powering 
    everything from consumer electronics to electric vehicles and grid-scale storage systems. The microstructural 
    characteristics of these batteries play a crucial role in determining their performance, efficiency, and 
    longevity. Understanding these characteristics at the pixel level can provide valuable insights into 
    battery behavior and potential failure mechanisms.
    
    This study focuses on identifying and analyzing GAP pixels within lithium-ion battery images. GAP pixels 
    are defined by two specific criteria: (1) they must have grayscale values between 5 and 30 (inclusive), 
    indicating particular material properties or structural features, and (2) they must be adjacent to at 
    least one pixel that has 20 contiguous pixels also meeting the first condition, suggesting the presence 
    of extended structural elements rather than isolated anomalies.
    
    The identification of these GAP pixels can reveal patterns related to electrode morphology, separator 
    characteristics, or potential defects that might impact battery performance. By automating this analysis 
    process, we can efficiently process multiple images and generate consistent, comparable results that 
    might be difficult to achieve through manual inspection.
    
    The purpose of this report is to present the findings from our GAP pixel analysis across multiple 
    lithium-ion battery images, highlighting the distribution and characteristics of these features and 
    discussing their potential implications for battery research and development.
    """
    doc.add_paragraph(introduction)
    
    # Methods section
    doc.add_heading('Methods', 1)
    methods = """
    Our analysis employed a computational approach using Python and the PIL (Python Imaging Library) to 
    process lithium-ion battery images. The methodology consisted of several key steps:
    
    Image Collection and Preprocessing:
    We focused on images with the "Li_" prefix stored in a specified directory, supporting both PNG and JPG 
    formats. Each image was converted to grayscale to enable pixel-level analysis based on intensity values, 
    which correspond to different material properties or structural features in the battery.
    
    GAP Pixel Identification:
    For each pixel in each image, we applied two specific criteria to identify GAP pixels:
    1. The pixel must have a grayscale value between 5 and 30 (inclusive).
    2. At least one adjacent pixel (up, down, left, or right) must have 20 contiguous pixels that also meet 
       the first criterion.
    
    This definition was designed to identify regions with specific optical properties that extend beyond 
    isolated points, potentially corresponding to meaningful structural features in the battery materials.
    
    Data Storage and Visualization:
    For each image, we generated two output files:
    1. A CSV file containing coordinates (row, column), grayscale value, and GAP flag (0 or 1) for each pixel.
    2. A PNG image highlighting GAP pixels in red (RGB: 255, 0, 0) against the original image background.
    
    These outputs provide both quantitative data for statistical analysis and visual representations for 
    qualitative assessment of GAP pixel distribution.
    
    The entire process was automated to ensure consistency across all images and to enable efficient 
    processing of large datasets. The implementation leveraged numpy for array operations, enhancing 
    computational efficiency when dealing with high-resolution images.
    """
    doc.add_paragraph(methods)
    
    # Results section
    doc.add_heading('Results', 1)
    results_intro = """
    Our analysis of the lithium-ion battery images revealed interesting patterns in the distribution of GAP 
    pixels across different samples. Below, we present the highlighted images showing the identified GAP 
    pixels in red, along with quantitative data extracted from the corresponding CSV files.
    """
    doc.add_paragraph(results_intro)
    
    # Get all highlighted images
    highlighted_images = glob.glob(os.path.join(output_dir, "*_gap_highlighted.png"))
    highlighted_images.sort()  # Sort to ensure consistent order
    
    # Create a summary table for all images
    doc.add_heading('Summary Statistics', 2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # Add header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Image Name'
    header_cells[1].text = 'Total Pixels'
    header_cells[2].text = 'GAP Pixels'
    header_cells[3].text = 'GAP Percentage'
    
    # Collect data for summary statistics
    gap_percentages = []
    image_names = []
    
    # Add each image and its analysis
    for image_path in highlighted_images:
        image_name = os.path.basename(image_path)
        base_name = image_name.replace("_gap_highlighted.png", "")
        
        # Add to document
        doc.add_heading(f"Image: {base_name}", 2)
        doc.add_picture(image_path, width=Inches(6))
        
        # Get corresponding CSV file
        csv_path = os.path.join(output_dir, f"{base_name}_gap_analysis.csv")
        
        if os.path.exists(csv_path):
            # Calculate statistics
            df = pd.read_csv(csv_path)
            total_pixels = len(df)
            gap_pixels = df[df['GAP Flag'] == 1].shape[0]
            gap_percentage = (gap_pixels / total_pixels) * 100 if total_pixels > 0 else 0
            
            # Add to summary data
            gap_percentages.append(gap_percentage)
            image_names.append(base_name)
            
            # Add row to table
            row_cells = table.add_row().cells
            row_cells[0].text = base_name
            row_cells[1].text = str(total_pixels)
            row_cells[2].text = str(gap_pixels)
            row_cells[3].text = f"{gap_percentage:.2f}%"
            
            # Add detailed analysis for this image
            analysis_text = f"""
            Analysis of {base_name}:
            - Total pixels analyzed: {total_pixels}
            - Number of GAP pixels identified: {gap_pixels}
            - Percentage of GAP pixels: {gap_percentage:.2f}%
            
            The distribution of GAP pixels in this image reveals potential structural features in the battery 
            material. The highlighted regions may correspond to specific microstructural elements that could 
            influence battery performance characteristics such as charge capacity, discharge rate, and cycle life.
            """
            doc.add_paragraph(analysis_text)
    
    # Create a bar chart of GAP percentages
    if gap_percentages:
        plt.figure(figsize=(10, 6))
        plt.bar(image_names, gap_percentages)
        plt.xlabel('Image')
        plt.ylabel('GAP Pixel Percentage')
        plt.title('Percentage of GAP Pixels Across Images')
        plt.xticks(rotation=45, ha='right')
        plt.tight_layout()
        
        # Save the chart
        chart_path = os.path.join(output_dir, "gap_percentage_chart.png")
        plt.savefig(chart_path)
        plt.close()
        
        # Add chart to document
        doc.add_heading('Comparative Analysis', 2)
        doc.add_paragraph('The chart below shows the percentage of GAP pixels across all analyzed images:')
        doc.add_picture(chart_path, width=Inches(6))
        
        # Add interpretation
        avg_percentage = np.mean(gap_percentages)
        max_percentage = np.max(gap_percentages)
        min_percentage = np.min(gap_percentages)
        
        interpretation = f"""
        The comparative analysis reveals:
        - Average GAP pixel percentage across all images: {avg_percentage:.2f}%
        - Maximum GAP pixel percentage: {max_percentage:.2f}%
        - Minimum GAP pixel percentage: {min_percentage:.2f}%
        
        This variation suggests differences in microstructural properties across the battery samples, 
        which may correlate with manufacturing conditions, material composition, or battery age. Further 
        investigation could explore the relationship between these GAP pixel distributions and battery 
        performance metrics.
        """
        doc.add_paragraph(interpretation)
    
    # Conclusion
    doc.add_heading('Conclusion', 1)
    conclusion = """
    This study has successfully identified and analyzed GAP pixels across multiple lithium-ion battery images, 
    providing insights into the microstructural characteristics of these critical energy storage devices. The 
    automated image processing approach developed for this analysis offers an efficient and consistent method 
    for identifying specific features that may be relevant to battery performance and reliability.
    
    The variation in GAP pixel distribution across different images suggests that this analysis method may be 
    sensitive to meaningful differences in battery microstructure. These differences could potentially be 
    correlated with performance parameters in future studies, offering a new approach to battery characterization 
    and quality control.
    
    The methodology developed here can be extended to larger datasets and adapted for different types of 
    battery images or other materials. Future work could explore correlations between GAP pixel patterns and 
    battery performance metrics, investigate the physical meaning of these pixels in terms of material properties, 
    and develop more sophisticated analysis techniques to extract additional information from battery images.
    
    In conclusion, this GAP pixel analysis represents a promising approach to battery image characterization 
    that could contribute to advances in battery design, manufacturing, and performance optimization.
    """
    doc.add_paragraph(conclusion)
    
    # Save the document
    report_path = os.path.join(output_dir, "Lithium_Battery_GAP_Analysis_Report.docx")
    doc.save(report_path)
    print(f"Report generated and saved to: {report_path}")

if __name__ == "__main__":
    print("Starting report generation...")
    generate_report()
    print("Report generation completed.")
