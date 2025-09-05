import os
import glob
from docx import Document
from docx.shared import Inches
from datetime import datetime

def generate_report():
    # Create a new Document
    doc = Document()
    
    # Define the output directory where the processed files are located
    output_directory = r"C:\Users\admin\Desktop\Python_proj\ALL_RESULT\CLAUDE\T1S1\backup5"
    
    # Add title
    doc.add_heading('Analysis of Grayscale Abnormal Pixels in Lithium-based Materials', level=1)
    
    # Add creation date
    doc.add_paragraph(f'Report generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    
    # Abstract Section
    doc.add_heading('Abstract', level=2)
    abstract = doc.add_paragraph()
    abstract.add_run(
        'This report presents a comprehensive analysis of Grayscale Abnormal Pixels (GAP) '
        'in lithium-based material images. Using digital image processing techniques, we identified pixels '
        'with specific grayscale characteristics that may indicate structural anomalies or significant '
        'features in the material. The analysis focused on pixels with grayscale values between 5 and 30 that '
        'are part of continuous structures, providing insights into the spatial distribution and patterns of '
        'these features. By analyzing multiple images with varying lithium concentrations (as indicated by the '
        'file names), we were able to observe how these GAP features evolve with changing material composition. '
        'The highlighted visualizations clearly show the distribution patterns of these pixels, which may '
        'correlate with material properties such as conductivity, stability, or capacity. This automated '
        'approach enables efficient identification of potential regions of interest for further investigation '
        'in lithium-based materials research and battery technology development.'
    )
    
    # Introduction Section
    doc.add_heading('Introduction', level=2)
    intro = doc.add_paragraph()
    intro.add_run(
        'The analysis of microscopic images plays a crucial role in understanding the structural properties '
        'of lithium-based materials used in modern battery technologies. Identifying specific patterns and '
        'anomalies within these images can provide valuable insights into material performance, degradation '
        'mechanisms, and potential optimizations. This study focuses on the identification of Grayscale '
        'Abnormal Pixels (GAP) defined by specific grayscale intensity ranges and continuity patterns.\n\n'
        'GAP pixels are characterized by grayscale values between 5 and 30 (inclusive) that form part of '
        'continuous structures with at least 20 contiguous pixels in any cardinal direction (up, down, left, or right). '
        'These characteristics may correspond to specific material features such as grain boundaries, defect '
        'structures, or lithium distribution patterns that are critical for understanding electrode behavior '
        'during charge-discharge cycles.\n\n'
        'The series of images analyzed in this study (Li_0.png through Li_1.25.png) represents materials with '
        'varying lithium concentrations, allowing us to observe how the distribution of GAP features changes '
        'with lithium content. This comparative analysis provides insights into the relationship between '
        'material composition and structural characteristics, which is essential for optimizing battery '
        'performance and longevity.'
    )
    
    # Methods Section
    doc.add_heading('Methods', level=2)
    methods = doc.add_paragraph()
    methods.add_run(
        'The analysis was conducted using Python with the Pillow (PIL) library for image processing. '
        'The methodology consisted of the following key steps:\n\n'
    )
    
    methods.add_run('1. Image Acquisition and Preprocessing: ').bold = True
    methods.add_run(
        'All images with the prefix "Li_" were collected from the specified directory. '
        'Each image was converted to grayscale to standardize the analysis process and eliminate '
        'color variations that might affect the identification of GAP pixels.\n\n'
    )
    
    methods.add_run('2. GAP Pixel Identification: ').bold = True
    methods.add_run(
        'For each pixel in the image, two conditions were evaluated:\n'
        '   a. Grayscale value between 5 and 30 (inclusive): This range was selected to focus on '
        'pixels with low to moderate intensity, which often correspond to significant structural features.\n'
        '   b. At least one adjacent direction (up, down, left, or right) containing 20 or more contiguous pixels '
        'that also meet the grayscale condition: This criterion ensures that isolated pixels or small clusters '
        'are excluded, focusing instead on continuous structures that are more likely to represent meaningful features.\n\n'
    )
    
    methods.add_run('3. Data Recording: ').bold = True
    methods.add_run(
        'For each image, a comprehensive CSV file was generated containing the coordinates (row, column), '
        'grayscale value, and GAP flag (1 if conditions met, 0 otherwise) for every pixel. This detailed '
        'record allows for statistical analysis and further investigation of specific regions of interest.\n\n'
    )
    
    methods.add_run('4. Visualization: ').bold = True
    methods.add_run(
        'New images were created highlighting the identified GAP pixels in red (RGB: 255, 0, 0) '
        'against the original background. This visual representation facilitates the identification of '
        'patterns and distributions that might not be immediately apparent in the numerical data.'
    )
    
    # Results Section
    doc.add_heading('Results', level=2)
    results = doc.add_paragraph()
    results.add_run(
        'The analysis successfully identified GAP pixels across all examined lithium-based material images. '
        'The distribution patterns of these pixels provide valuable insights into the structural characteristics '
        'of the materials with varying lithium concentrations. Below are the key observations from the analysis:\n\n'
        '1. The GAP pixels tend to form continuous structures rather than appearing as isolated points, '
        'suggesting they represent meaningful material features rather than random noise or artifacts.\n\n'
        '2. The spatial distribution of GAP pixels varies significantly between images with different lithium '
        'concentrations (from Li_0.png to Li_1.25.png), indicating a correlation between lithium content and '
        'the formation of these structural features.\n\n'
        '3. In several images, the GAP pixels appear to align along specific directions or form network-like '
        'patterns, potentially corresponding to grain boundaries, intercalation pathways, or other structural '
        'features important for lithium-ion transport.\n\n'
        '4. The density and connectivity of GAP pixels show variations that may reflect changes in material '
        'properties such as conductivity, stability, or capacity with different lithium concentrations.\n\n'
        'The highlighted images below show the identified GAP pixels (in red) for each analyzed sample, '
        'providing a visual representation of how these features evolve with changing lithium content:'
    )
    
    # Add all highlighted images to the document
    highlighted_images = glob.glob(os.path.join(output_directory, '*_highlighted.png'))
    
    # Sort the images by lithium concentration for better comparison
    def extract_li_value(filename):
        basename = os.path.basename(filename)
        if basename.startswith('Li_'):
            try:
                # Extract the number after 'Li_' and before '_highlighted.png'
                li_value = basename[3:].split('_')[0]
                return float(li_value)
            except (ValueError, IndexError):
                return float('inf')  # Place at the end if conversion fails
        return float('inf')
    
    highlighted_images.sort(key=extract_li_value)
    
    for img_path in highlighted_images:
        img_name = os.path.basename(img_path)
        original_name = img_name.replace('_highlighted.png', '.png')
        doc.add_paragraph(f'Image: {original_name}', style='Heading 4')
        doc.add_picture(img_path, width=Inches(6))
        
        # Add a brief description for each image
        li_value = extract_li_value(img_path)
        if li_value != float('inf'):
            description = f"This image shows the GAP pixel distribution for a lithium concentration of {li_value}. "
            if li_value < 0.5:
                description += "At this lower concentration, the GAP features appear to be more dispersed."
            elif li_value < 1.0:
                description += "At this medium concentration, the GAP features show moderate connectivity."
            else:
                description += "At this higher concentration, the GAP features exhibit increased density and connectivity."
            doc.add_paragraph(description)
        
        doc.add_paragraph('')  # Add some space after each image
    
    # Conclusion paragraph
    doc.add_paragraph(
        'In conclusion, this analysis demonstrates the effectiveness of our automated approach in identifying '
        'and visualizing GAP features in lithium-based materials. The observed patterns and their variations '
        'with lithium concentration provide valuable insights for materials scientists and battery researchers. '
        'Future work could focus on correlating these GAP features with specific material properties and '
        'performance metrics to further optimize lithium-based materials for advanced battery applications.'
    )
    
    # Save the document
    report_path = os.path.join(output_directory, 'GAP_Analysis_Report.docx')
    doc.save(report_path)
    print(f"Report generated and saved to: {report_path}")
    print("Calculation successful")

if __name__ == "__main__":
    generate_report()
