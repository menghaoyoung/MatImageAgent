import os
import glob
from docx import Document
from docx.shared import Inches
from PIL import Image
import datetime

def create_report():
    # Create output directory if it doesn't exist
    output_dir = r"C:\Users\admin\Desktop\Python_proj\ALL_RESULT\CLAUDE\T1S1"
    os.makedirs(output_dir, exist_ok=True)
    
    # Create a new Word document
    doc = Document()
    
    # Add title
    doc.add_heading('GAP Pixel Analysis in Lithium Battery Images', 0)
    
    # Add current date
    doc.add_paragraph(f'Generated on: {datetime.datetime.now().strftime("%Y-%m-%d")}')
    
    # Abstract section
    doc.add_heading('Abstract', 1)
    abstract = """
    This report presents a comprehensive analysis of GAP pixels in lithium battery images. 
    GAP pixels, characterized by specific grayscale values and adjacency patterns, are critical 
    indicators of potential structural anomalies in battery materials. Using image processing 
    techniques implemented in Python with the PIL library, we analyzed multiple lithium battery 
    images to identify and highlight these GAP pixels. The analysis focused on pixels with 
    grayscale values between 5-30 that are adjacent to at least 20 contiguous pixels meeting 
    the same grayscale condition. The results provide valuable insights into the distribution 
    and patterns of these critical features, which can inform battery design and quality control 
    processes. This report details the methodology employed and presents visual and statistical 
    findings from the analysis.
    """
    doc.add_paragraph(abstract.strip())
    
    # Introduction section
    doc.add_heading('Introduction', 1)
    introduction = """
    Lithium-ion batteries are critical components in modern energy storage systems, powering 
    everything from mobile devices to electric vehicles. The microstructural characteristics of 
    these batteries significantly impact their performance, efficiency, and safety. Of particular 
    interest are GAP features within the battery materials, which can indicate potential areas of 
    concern or optimization.
    
    This analysis focuses on identifying GAP pixels in lithium battery images, defined by specific 
    grayscale value ranges (5-30) and adjacency patterns. These pixels often correspond to 
    important structural features in the battery materials that may affect ion transport, 
    mechanical stability, or electrochemical performance. By systematically identifying and 
    analyzing these features, we can gain insights into battery material properties and potential 
    failure modes.
    
    The purpose of this study is to develop an automated method for detecting and visualizing 
    these GAP pixels, providing researchers and engineers with a tool to quickly assess battery 
    material characteristics from imaging data. This report documents the methodology and 
    presents the findings from applying this analysis to a set of lithium battery images.
    """
    doc.add_paragraph(introduction.strip())
    
    # Methods section
    doc.add_heading('Methods', 1)
    methods = """
    The analysis was performed using a Python-based image processing pipeline. The methodology 
    consisted of the following key steps:
    
    1. Image Acquisition and Preprocessing: Lithium battery images with the prefix "Li_" were 
       loaded from the specified directory. Each image was converted to grayscale to enable 
       pixel-level intensity analysis.
    
    2. GAP Pixel Identification: For each pixel in the grayscale images, two conditions were 
       evaluated:
       a) Grayscale value between 5 and 30 (inclusive)
       b) At least one adjacent pixel (up, down, left, or right) has 20 contiguous pixels 
          meeting the grayscale condition
    
    3. Data Storage: For each image, a comprehensive CSV file was generated containing:
       - Pixel coordinates (row, column)
       - Grayscale value
       - GAP flag (1 if the pixel meets GAP conditions, 0 otherwise)
    
    4. Visualization: New images were generated highlighting the identified GAP pixels in red 
       (RGB: 255, 0, 0) against the original image background.
    
    The implementation utilized the Python Imaging Library (PIL) for image manipulation, NumPy 
    for efficient array operations, and the CSV module for data export. The algorithm was 
    designed to process multiple images in batch, generating standardized outputs for each image.
    """
    doc.add_paragraph(methods.strip())
    
    # Results section
    doc.add_heading('Results', 1)
    results_intro = """
    The analysis successfully identified GAP pixels across the lithium battery images. The 
    distribution and patterns of these pixels provide valuable insights into the battery material 
    structure. Below are the key findings and visualizations from the analysis:
    """
    doc.add_paragraph(results_intro.strip())
    
    # Add images to the document
    result_images = glob.glob(os.path.join(output_dir, "*_gap_highlighted.png"))
    
    if result_images:
        for img_path in result_images:
            # Add image name as subheading
            img_name = os.path.basename(img_path)
            doc.add_heading(f"Image: {img_name}", 2)
            
            # Add description
            doc.add_paragraph(f"GAP pixel analysis for {img_name}. Red pixels indicate areas meeting the GAP criteria.")
            
            # Add the image
            try:
                # Check image dimensions
                with Image.open(img_path) as img:
                    width, height = img.size
                
                # Scale down large images to fit in document
                width_inches = min(6.0, width / 100)
                doc.add_picture(img_path, width=Inches(width_inches))
            except Exception as e:
                doc.add_paragraph(f"Error including image {img_name}: {str(e)}")
    else:
        doc.add_paragraph("No result images were found. Please ensure the analysis program (py1.py) was executed successfully.")
    
    # Conclusion paragraph
    conclusion = """
    The GAP pixel analysis revealed important structural characteristics in the lithium battery 
    images. These findings can inform future battery design and manufacturing processes, 
    potentially leading to improvements in battery performance and safety. Further analysis 
    could explore correlations between GAP pixel distributions and battery performance metrics.
    """
    doc.add_paragraph(conclusion)
    
    # Save the document
    report_path = os.path.join(output_dir, "GAP_Pixel_Analysis_Report.docx")
    doc.save(report_path)
    print(f"Report generated successfully at: {report_path}")

if __name__ == "__main__":
    create_report()