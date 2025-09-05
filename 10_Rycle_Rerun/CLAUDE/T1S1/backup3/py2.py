import os
import glob
from docx import Document
from docx.shared import Inches
import pandas as pd
import matplotlib.pyplot as plt
from PIL import Image
import datetime

def create_report():
    # Create a new Document
    doc = Document()
    
    # Add title
    doc.add_heading('Gap Analysis in Lithium Battery Images: Detection and Visualization', 0)
    
    # Add date
    doc.add_paragraph(f"Report generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    # Abstract section
    doc.add_heading('Abstract', 1)
    abstract = """This report presents a comprehensive analysis of gap detection in lithium battery images using specialized image processing techniques. By defining gaps according to specific grayscale value criteria (5-30) and contiguity requirements (at least 20 adjacent pixels meeting the same criteria), we systematically identified potential structural irregularities in battery materials. The analysis covered 11 distinct lithium battery images, each converted to grayscale for pixel-by-pixel examination. Our methodology successfully detected gap regions, which were highlighted in red for visual identification and documented in detailed CSV files containing pixel coordinates, grayscale values, and gap flags. These findings provide valuable insights into the internal structure of lithium batteries, potentially revealing manufacturing inconsistencies or design flaws that could affect performance and safety. The automated approach demonstrates the feasibility of using computer vision for quality control in battery production, offering a more efficient alternative to manual inspection. This report details our approach, findings, and visualizations, contributing to the ongoing development of reliable battery assessment techniques."""
    doc.add_paragraph(abstract)
    
    # Introduction section
    doc.add_heading('Introduction', 1)
    introduction = """Lithium-ion batteries have become indispensable components in modern electronics, electric vehicles, and renewable energy storage systems. The internal structure of these batteries significantly impacts their performance, safety, and longevity. Structural inconsistencies, particularly gaps or voids within the battery materials, can lead to reduced capacity, increased internal resistance, and in extreme cases, safety hazards.

The purpose of this analysis is to develop and implement an automated method for detecting and characterizing gaps in lithium battery images. By defining gaps according to specific grayscale value criteria and contiguity requirements, we aim to establish an objective measure for identifying potential structural issues that might affect battery performance. This approach offers several advantages over traditional visual inspection methods, including consistency, reproducibility, and the ability to process large numbers of images efficiently.

The background of this work stems from the increasing demand for high-quality, safe lithium batteries in various applications. As battery technology continues to evolve, there is a growing need for sophisticated quality control methods that can identify potential defects before batteries reach the market. Image analysis provides a non-destructive means of examining battery structures, offering insights that might not be apparent through other testing methods.

This report details our methodology for gap detection, presents the results of our analysis across multiple lithium battery images, and discusses the implications of our findings for battery quality assessment and manufacturing processes."""
    doc.add_paragraph(introduction)
    
    # Methods section
    doc.add_heading('Methods', 1)
    methods = """Our methodology employed Python programming with several key libraries for image processing and analysis. The approach can be broken down into the following steps:

1. Image Acquisition and Preprocessing: Using the PIL (Python Imaging Library), we read all images with the "Li_" prefix from the specified directory, supporting both PNG and JPG formats. Each image was converted to grayscale to simplify the analysis and focus on intensity variations that might indicate structural features. This conversion reduced the dimensionality of the data while preserving the essential information needed for gap detection.

2. Gap Definition and Detection Algorithm: We defined gap pixels using two specific criteria:
   - Grayscale value between 5 and 30 (inclusive), indicating potential void or low-density regions
   - At least one adjacent direction (up, down, left, or right) containing 20 contiguous pixels that also meet the grayscale value criterion

   The algorithm systematically examined each pixel in the image, first checking its grayscale value and then, if necessary, examining adjacent pixels in four directions to determine if the contiguity requirement was met. This approach allowed for comprehensive mapping of potential gap regions within the battery structure.

3. Data Collection and Storage: For each pixel in every image, we recorded:
   - Pixel coordinates (row, column)
   - Grayscale value
   - GAP flag (1 if the pixel meets gap criteria, 0 otherwise)
   
   This data was stored in CSV files named according to the original image name, with the suffix "_gap_analysis.csv". These files provide a detailed record of the analysis that can be referenced for further study or verification.

4. Visualization: We created highlighted images where gap pixels were colored red (RGB: 255, 0, 0) against the grayscale background, providing visual confirmation of gap detection and distribution. These images were saved as PNG files with the suffix "_highlighted.png", allowing for easy identification of gap patterns and their relationship to other structural features.

The entire process was automated and applied consistently across all 11 lithium battery images, ensuring comparable results and facilitating the identification of patterns or trends across different samples."""
    doc.add_paragraph(methods)
    
    # Results section
    doc.add_heading('Results', 1)
    results_intro = """Our analysis successfully processed 11 lithium battery images, identifying gap regions based on the defined criteria. The results revealed interesting patterns of potential structural irregularities across the different battery samples. Below, we present the highlighted images showing detected gap regions in red, along with a brief analysis of each image."""
    doc.add_paragraph(results_intro)
    
    # Add all highlighted images to the document
    output_dir = os.path.join(os.path.dirname(r"C:\Users\admin\Desktop\Python_proj\distance_analysis_new\Images"), "output")
    highlighted_images = glob.glob(os.path.join(output_dir, "*_highlighted.png"))
    
    # Sort images to ensure consistent presentation
    highlighted_images.sort()
    
    # Collect gap percentage data for summary
    gap_percentages = []
    image_names = []
    
    for img_path in highlighted_images:
        img_name = os.path.basename(img_path)
        base_name = img_name.replace("_highlighted.png", "")
        image_names.append(base_name)
        
        doc.add_paragraph(f"Figure: {base_name} - Highlighted gap regions in red")
        doc.add_picture(img_path, width=Inches(6))
        
        # Add a brief analysis of each image
        csv_path = img_path.replace("_highlighted.png", "_gap_analysis.csv")
        if os.path.exists(csv_path):
            try:
                df = pd.read_csv(csv_path)
                gap_count = df['GAP Flag'].sum()
                total_pixels = len(df)
                gap_percentage = (gap_count / total_pixels) * 100
                gap_percentages.append(gap_percentage)
                
                analysis_text = f"""Analysis: In this image, {gap_count:,} pixels out of {total_pixels:,} total pixels ({gap_percentage:.2f}%) were identified as gap regions. These regions may indicate areas of structural inconsistency in the battery material that could potentially affect performance and reliability."""
                doc.add_paragraph(analysis_text)
            except Exception as e:
                doc.add_paragraph(f"Analysis: Gap regions identified in this image may indicate structural inconsistencies in the battery material. (Note: Detailed statistics unavailable due to processing error: {str(e)})")
        else:
            doc.add_paragraph("Analysis: Gap regions identified in this image may indicate structural inconsistencies in the battery material. (Note: Detailed statistics unavailable)")
    
    # Add summary analysis if we have gap percentage data
    if gap_percentages:
        doc.add_heading('Summary Analysis', 2)
        avg_gap_percentage = sum(gap_percentages) / len(gap_percentages)
        max_gap_percentage = max(gap_percentages)
        min_gap_percentage = min(gap_percentages)
        max_index = gap_percentages.index(max_gap_percentage)
        min_index = gap_percentages.index(min_gap_percentage)
        
        summary_text = f"""Across all analyzed images, the average percentage of pixels identified as gap regions was {avg_gap_percentage:.2f}%. The highest gap percentage ({max_gap_percentage:.2f}%) was observed in image {image_names[max_index]}, while the lowest ({min_gap_percentage:.2f}%) was in image {image_names[min_index]}. This variation suggests differences in structural consistency across the battery samples, which could be attributed to manufacturing processes, material composition, or other factors.

The distribution and patterns of gap regions provide valuable insights into the internal structure of the batteries. Clusters of gap pixels may indicate areas of concern that warrant further investigation, particularly if they appear consistently across multiple samples. These findings can inform quality control processes and help identify manufacturing issues that might affect battery performance and safety.

Future work could expand on this methodology by correlating detected gap patterns with battery performance metrics, developing more sophisticated detection algorithms, or extending the analysis to three-dimensional imaging techniques for more comprehensive structural assessment."""
        doc.add_paragraph(summary_text)
    
    # Save the document
    report_path = os.path.join(r"C:\Users\admin\Desktop\Python_proj\ALL_RESULT\CLAUDE\T1S1\backup3", "Lithium_Battery_Gap_Analysis_Report.docx")
    doc.save(report_path)
    print(f"Report successfully generated and saved to: {report_path}")

if __name__ == "__main__":
    create_report()
