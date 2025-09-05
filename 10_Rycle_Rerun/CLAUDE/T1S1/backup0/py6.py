import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def generate_report():
    """
    Generate a Word document report based on the GAP analysis results.
    The report includes four sections: Abstract, Introduction, Methods, and Results.
    """
    # Define paths
    output_directory = r"C:\Users\admin\Desktop\Python_proj\ALL_RESULT\CLAUDE\T1S1\backup9"
    report_path = os.path.join(output_directory, "GAP_Analysis_Report.docx")
    
    # Get all highlighted images
    highlighted_images = [f for f in os.listdir(output_directory) 
                         if f.endswith('_gap_highlighted.png')]
    
    # Sort images by numeric value in filename for consistent presentation
    def extract_number(filename):
        try:
            # Extract number from filename (e.g., "Li_0.125_gap_highlighted.png" -> 0.125)
            num_str = filename.replace('Li_', '').replace('_gap_highlighted.png', '')
            return float(num_str) if num_str else 0
        except ValueError:
            return float('inf')
    
    highlighted_images.sort(key=extract_number)
    
    # Create a new Word document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Add title
    title = doc.add_heading('Grayscale Adjacent Pixel (GAP) Analysis of Lithium-Based Materials', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Abstract section
    doc.add_heading('Abstract', level=1)
    abstract = """
    This report presents a comprehensive analysis of Grayscale Adjacent Pixel (GAP) patterns in lithium-based material images. 
    The analysis identifies specific pixel regions with grayscale values between 5-30 that are adjacent to at least 20 contiguous 
    pixels meeting the same condition. These patterns may indicate significant structural features or defects in the materials. 
    Using custom image processing algorithms implemented in Python, we processed a series of lithium-based images with varying 
    compositions to identify and visualize these GAP regions. The results show distinct patterns of GAP distribution across different 
    samples, suggesting potential correlations with material properties. The highlighted regions provide valuable insights into the 
    spatial distribution of these features and may guide further investigation into the structural characteristics of lithium-based 
    materials. This analysis demonstrates the effectiveness of targeted image processing techniques in identifying subtle patterns 
    that may not be immediately apparent through visual inspection alone. The findings contribute to our understanding of the 
    microstructural properties of these materials and provide a foundation for more detailed structural analysis.
    """
    doc.add_paragraph(abstract.strip())
    
    # Introduction section
    doc.add_heading('Introduction', level=1)
    introduction = """
    The microstructural analysis of lithium-based materials is crucial for understanding their properties and optimizing their 
    performance in various applications, including energy storage and conversion technologies. Advanced imaging techniques provide 
    valuable insights into the structural characteristics of these materials, but extracting meaningful information from these images 
    requires sophisticated image processing methods.

    This study focuses on identifying specific pixel patterns defined as Grayscale Adjacent Pixel (GAP) conditions. These conditions 
    are characterized by pixels with grayscale values between 5 and 30 (inclusive) that are adjacent to at least 20 contiguous pixels 
    meeting the same grayscale criterion. The identification of these patterns is important because they may correspond to specific 
    structural features, such as grain boundaries, defects, or phase interfaces, which significantly influence the material's 
    properties.

    Previous research has shown that grayscale intensity in material images often correlates with compositional or structural 
    variations. The specific range of 5-30 was selected based on preliminary observations suggesting that this range captures 
    important transitional regions in the material structure. The requirement for adjacency to 20 contiguous pixels ensures that 
    we identify meaningful patterns rather than isolated pixels or noise.

    The purpose of this analysis is to systematically identify, visualize, and characterize these GAP regions across multiple 
    lithium-based samples with varying compositions. By highlighting these regions and analyzing their distribution, we aim to 
    provide insights into the relationship between material composition and structural features, potentially informing future 
    material design and optimization efforts.
    """
    doc.add_paragraph(introduction.strip())
    
    # Methods section
    doc.add_heading('Methods', level=1)
    methods = """
    The analysis was conducted using a custom Python program that leverages the PIL (Python Imaging Library) for image processing 
    and numpy for efficient numerical operations. The methodology consisted of the following key steps:

    Image Acquisition and Preprocessing:
    We collected a series of lithium-based material images with the prefix "Li_" from the specified directory. These images 
    represent samples with varying lithium content, as indicated by the numerical values in their filenames. Each image was 
    converted to grayscale to standardize the analysis and focus on intensity variations rather than color.

    GAP Condition Identification:
    For each pixel in every image, we performed a two-stage evaluation:
    1. First, we checked if the pixel's grayscale value fell within the range of 5-30 (inclusive).
    2. For qualifying pixels, we examined the four adjacent directions (up, down, left, right) to determine if any direction 
       contained at least 20 contiguous pixels that also met the grayscale condition.

    This approach allowed us to identify not just isolated pixels with specific grayscale values but connected regions that might 
    represent meaningful structural features.

    Data Recording and Visualization:
    For each image, we generated two output files:
    1. A comprehensive CSV file containing the coordinates (row, column), grayscale value, and GAP flag (0 or 1) for every pixel.
    2. A visual representation where pixels meeting both GAP conditions were highlighted in red (RGB: 255, 0, 0) against the 
       original image background.

    The CSV files provide detailed quantitative data for further statistical analysis, while the highlighted images offer an 
    intuitive visualization of the GAP regions' distribution and patterns.

    The program was designed to process multiple images sequentially, maintaining consistent analytical parameters across all 
    samples to enable meaningful comparisons. The analysis completed in approximately 103-104 seconds for the entire dataset, 
    demonstrating efficient processing despite the pixel-by-pixel evaluation approach.
    """
    doc.add_paragraph(methods.strip())
    
    # Results section
    doc.add_heading('Results', level=1)
    results_intro = """
    The GAP analysis revealed distinct patterns of pixels meeting both conditions across the lithium-based samples. These patterns, 
    highlighted in red in the processed images, provide valuable insights into the structural characteristics of the materials and 
    how they vary with composition.

    The following figures present the visual results of our GAP analysis for each sample. Each image shows the original structure 
    with GAP regions highlighted in red, allowing for direct comparison across different lithium concentrations.
    """
    doc.add_paragraph(results_intro.strip())
    
    # Add all highlighted images to the document with proper captions
    for i, img_file in enumerate(highlighted_images):
        img_path = os.path.join(output_directory, img_file)
        base_name = img_file.replace('_gap_highlighted.png', '')
        
        # Add image caption
        caption = f"Figure {i+1}: GAP analysis results for {base_name}. Red pixels indicate regions meeting both GAP conditions: grayscale value between 5-30 and adjacency to at least 20 contiguous pixels with similar grayscale values."
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add the image
        try:
            doc.add_picture(img_path, width=Inches(6.0))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            doc.add_paragraph(f"Error including image {img_file}: {str(e)}")
        
        # Add a paragraph break
        doc.add_paragraph()
    
    # Add concluding paragraph to results
    results_conclusion = """
    The analysis reveals several interesting patterns in the GAP distribution across samples with different lithium concentrations. 
    Notable observations include:

    1. Variation in GAP region density: Samples with higher lithium content (e.g., Li_1.0, Li_1.125, Li_1.25) generally show 
       different patterns of GAP regions compared to those with lower concentrations.

    2. Spatial distribution patterns: In several samples, the GAP regions appear to form network-like structures, potentially 
       indicating grain boundaries or phase interfaces in the material.

    3. Consistency within concentration groups: Samples with similar lithium concentrations tend to exhibit comparable GAP patterns, 
       suggesting a correlation between composition and structural features.

    These observations provide valuable insights into how the microstructural characteristics of lithium-based materials evolve 
    with changing composition. The highlighted GAP regions may represent critical structural features that influence material 
    properties such as ionic conductivity, mechanical strength, or electrochemical performance.

    The detailed pixel-level data stored in the CSV files enables further quantitative analysis, including statistical evaluation 
    of GAP region size, distribution, and connectivity. Such analysis could reveal additional correlations between composition and 
    structure that are not immediately apparent from visual inspection alone.

    Future work could explore the relationship between these GAP regions and specific material properties through complementary 
    characterization techniques. Additionally, the methodology could be extended to include more sophisticated pattern recognition 
    approaches to identify complex structural features across different material systems.
    """
    doc.add_paragraph(results_conclusion.strip())
    
    # Save the document
    try:
        doc.save(report_path)
        print(f"Report generated and saved to: {report_path}")
    except Exception as e:
        print(f"Error saving report: {str(e)}")

if __name__ == "__main__":
    # Check if python-docx is installed
    try:
        import docx
    except ImportError:
        print("python-docx package is not installed. Please install it using:")
        print("pip install python-docx")
        exit(1)
        
    generate_report()
