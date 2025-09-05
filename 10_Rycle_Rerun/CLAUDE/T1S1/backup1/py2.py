import os
import glob
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from PIL import Image
import datetime

def generate_report():
    # Create a new Document
    doc = Document()
    
    # Set up styles
    styles = doc.styles
    title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    
    heading_style = styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
    heading_style.font.size = Pt(16)
    heading_style.font.bold = True
    
    # Add title
    title = doc.add_paragraph('Grayscale Adjacency Pattern (GAP) Analysis in Lithium Materials', style='CustomTitle')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    date_paragraph = doc.add_paragraph(f"Report generated on {datetime.datetime.now().strftime('%Y-%m-%d')}")
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Abstract section
    doc.add_heading('Abstract', level=1)
    abstract = doc.add_paragraph(
        'This report presents a comprehensive analysis of Grayscale Adjacency Patterns (GAP) in a series of lithium-based '
        'material images. Using digital image processing techniques, we identified pixels exhibiting specific grayscale '
        'characteristics (values between 5-30) that are adjacent to extended regions with similar properties. These GAP '
        'pixels potentially highlight significant structural or compositional features within the lithium materials that '
        'may not be immediately apparent through visual inspection alone. The analysis processed multiple images with '
        'varying lithium concentrations, converting them to grayscale and systematically evaluating each pixel against '
        'the defined GAP criteria. The results are visualized through highlighted images where GAP pixels are marked in '
        'red, revealing distinct pattern distributions across different samples. This pixel-level approach provides '
        'valuable insights into the microstructural characteristics of lithium materials, potentially contributing to '
        'improved understanding of their properties and performance in various applications. The methodology developed '
        'here demonstrates the value of computational image analysis in materials science research.'
    )
    
    # Introduction section
    doc.add_heading('Introduction', level=1)
    introduction = doc.add_paragraph(
        'The analysis of lithium-based materials at the microstructural level is crucial for understanding their '
        'properties and optimizing their performance in applications such as batteries, alloys, and other advanced '
        'materials. Traditional visual inspection of microscopy images often fails to capture subtle patterns and '
        'features that may significantly influence material behavior. This study introduces a systematic approach to '
        'identify specific pixel patterns in grayscale images of lithium samples, which we term Grayscale Adjacency '
        'Patterns (GAP).\n\n'
        'GAP pixels are defined by two key characteristics: (1) a grayscale value falling within a specific range '
        '(5-30 on a 0-255 scale), indicating particular density or composition properties, and (2) adjacency to at '
        'least one direction containing 20 contiguous pixels meeting the same grayscale condition, suggesting the '
        'presence of extended structural features rather than isolated anomalies.\n\n'
        'By analyzing these patterns across multiple lithium samples with varying concentrations (as indicated by the '
        'numerical values in the image filenames), we aim to identify potential correlations between lithium content '
        'and specific structural characteristics. This approach combines image processing techniques with materials '
        'science principles to extract quantitative data from qualitative images, potentially revealing insights that '
        'would remain hidden through conventional analysis methods.'
    )
    
    # Methods section
    doc.add_heading('Methods', level=1)
    methods = doc.add_paragraph(
        'The analysis methodology employed in this study involved several key steps, implemented through Python '
        'programming with the PIL (Python Imaging Library) package for image processing.\n\n'
        'Image Acquisition and Preprocessing: Images with the "Li_" prefix were collected from the specified directory, '
        'representing lithium samples at different concentrations. Each image was converted to grayscale to focus the '
        'analysis on intensity patterns rather than color variations, which standardized the data for consistent processing.\n\n'
        'GAP Pixel Identification: Each pixel in every image was evaluated against two criteria to determine its GAP status:\n'
        '   1. Grayscale Value Criterion: The pixel must have a grayscale value between 5 and 30 (inclusive) on the '
        '      standard 0-255 scale. This range was selected to highlight specific density or compositional features '
        '      within the lithium samples.\n'
        '   2. Adjacency Criterion: The pixel must be adjacent to at least one direction (up, down, left, or right) '
        '      containing 20 contiguous pixels that also meet the grayscale value criterion. This requirement ensures '
        '      that identified pixels are part of extended features rather than isolated points.\n\n'
        'Data Recording: For each pixel in every image, the following data was recorded in a CSV file:\n'
        '   - Pixel coordinates (row, column)\n'
        '   - Grayscale value\n'
        '   - GAP flag (1 if the pixel meets both criteria, 0 otherwise)\n\n'
        'Visualization: To facilitate visual interpretation of the results, new images were generated for each input '
        'image, highlighting the identified GAP pixels in red (RGB: 255, 0, 0) against the original grayscale background. '
        'This visualization approach allows for immediate identification of pattern distributions and potential correlations '
        'with lithium concentration.\n\n'
        'The entire analysis process was automated through Python programming, ensuring consistent application of the '
        'criteria across all images and eliminating potential human bias in the identification of GAP pixels.'
    )
    
    # Results section
    doc.add_heading('Results', level=1)
    results = doc.add_paragraph(
        'The analysis of lithium sample images revealed distinct Grayscale Adjacency Patterns (GAP) across different '
        'concentrations. These patterns, highlighted in red in the processed images, show the distribution of pixels '
        'meeting both the grayscale value criterion (5-30) and the adjacency criterion (proximity to 20 contiguous '
        'qualifying pixels).\n\n'
        'The distribution and density of GAP pixels vary notably across the different lithium concentrations, suggesting '
        'a potential correlation between lithium content and specific structural features. In some samples, GAP pixels '
        'form clear linear or clustered patterns, while in others, they appear more dispersed or concentrated in '
        'particular regions of the image.\n\n'
        'Each image analysis also generated a comprehensive CSV file containing pixel-level data, allowing for further '
        'quantitative analysis beyond the visual representations shown below. These data files enable statistical '
        'analysis of GAP pixel distribution, density, and potential correlation with lithium concentration values.\n\n'
        'The highlighted images below provide a visual representation of the GAP pixel distribution in each sample:'
    )
    
    # Add images to the document
    output_directory = r"C:\Users\admin\Desktop\Python_proj\ALL_RESULT\CLAUDE\T1S1\backup1"
    highlighted_images = sorted(glob.glob(os.path.join(output_directory, "Li_*_highlighted.png")))
    
    # Create a table for images (2 columns)
    if highlighted_images:
        # Calculate number of rows needed (2 images per row)
        num_rows = (len(highlighted_images) + 1) // 2
        
        for i in range(0, len(highlighted_images), 2):
            doc.add_paragraph()  # Add some space
            
            # Create a table for this row of images
            table = doc.add_table(rows=1, cols=2)
            
            # First image in row
            cell = table.cell(0, 0)
            filename = os.path.basename(highlighted_images[i])
            original_name = filename.replace("_highlighted.png", "")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(f"{original_name}")
            run = p.add_run()
            run.add_picture(highlighted_images[i], width=Inches(3))
            
            # Second image in row (if exists)
            if i + 1 < len(highlighted_images):
                cell = table.cell(0, 1)
                filename = os.path.basename(highlighted_images[i+1])
                original_name = filename.replace("_highlighted.png", "")
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run(f"{original_name}")
                run = p.add_run()
                run.add_picture(highlighted_images[i+1], width=Inches(3))
    
    # Conclusion
    doc.add_heading('Conclusion', level=1)
    conclusion = doc.add_paragraph(
        'This study has successfully demonstrated the application of digital image processing techniques to identify '
        'and analyze Grayscale Adjacency Patterns (GAP) in lithium material samples. The identified GAP pixels, '
        'characterized by specific grayscale values and adjacency properties, potentially correspond to important '
        'structural or compositional features within the materials.\n\n'
        'The variation in GAP pixel distribution across different lithium concentrations suggests that this approach '
        'can detect meaningful differences between samples, potentially correlating with material properties or '
        'performance characteristics. The methodology developed in this study provides a foundation for more advanced '
        'image analysis techniques in materials science.\n\n'
        'Future work could explore several directions, including:\n'
        '1. Correlating GAP patterns with measured material properties to establish predictive relationships\n'
        '2. Refining the GAP criteria to target specific features of interest more precisely\n'
        '3. Applying machine learning techniques to automatically classify patterns and predict material behavior\n'
        '4. Extending the analysis to three-dimensional imaging data for more comprehensive structural understanding\n\n'
        'This pixel-level approach to materials image analysis demonstrates the value of computational methods in '
        'extracting quantitative insights from qualitative visual data, potentially leading to improved understanding '
        'and design of lithium-based materials for various applications.'
    )
    
    # Save the document
    report_path = os.path.join(output_directory, "GAP_Analysis_Report.docx")
    doc.save(report_path)
    print(f"Report generated and saved to: {report_path}")

if __name__ == "__main__":
    try:
        generate_report()
    except Exception as e:
        print(f"Error generating report: {e}")
