# py3_generate_report.py
import os
import docx
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def get_image_files(input_directory):
    return [
        f for f in os.listdir(input_directory)
        if f.startswith("Li_") and f.lower().endswith(('.png', '.jpg', '.jpeg'))
    ]

def read_txt_report(txt_path):
    """
    Read the txt report to get re and max GAP height.
    Returns: (re, max_gap_height)
    """
    re = None
    max_height = None
    if not os.path.exists(txt_path):
        return None, None
    with open(txt_path, 'r') as f:
        for line in f:
            if "Physical dimension" in line:
                try:
                    re = float(line.strip().split(":")[1].split()[0])
                except:
                    re = None
            if "Max GAP height" in line:
                try:
                    max_height = float(line.strip().split(":")[1].split()[0])
                except:
                    max_height = None
    return re, max_height

def add_section_heading(doc, text, level=1):
    doc.add_heading(text, level=level)

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold

def add_image(doc, image_path, width_inch=4.5, caption=None):
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(width_inch))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        if caption:
            p = doc.add_paragraph(caption)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.runs[0].italic = True

def generate_report(input_dir, output_dir, output_docx):
    doc = docx.Document()
    
    # Title
    doc.add_heading('Simulation Report: Automated GAP Detection in Lithium Images', 0)

    # Abstract
    add_section_heading(doc, 'Abstract', 1)
    add_paragraph(doc,
        "This report presents the results of an automated simulation for detecting GAP regions in lithium images using pixel-wise grayscale analysis. "
        "Using a Python-based pipeline, images with the 'Li_' prefix were analyzed to extract regions of interest according to strict grayscale and adjacency criteria. "
        "The study quantifies the height of detected GAPs per column, providing insights into the physical dimensions of these features. "
        "Highlighted images and tabulated results are used to support the findings."
    )
    
    # Introduction
    add_section_heading(doc, 'Introduction', 1)
    add_paragraph(doc,
        "Lithium imaging is crucial for understanding electrochemical processes in battery research. Detecting and quantifying microstructural gaps, or 'GAP' regions, "
        "is important for evaluating electrode performance and failure mechanisms. Manual assessment of such features is time-consuming and subjective. "
        "This simulation leverages automated image processing to detect GAP pixels based on grayscale thresholds and spatial continuity, "
        "enabling reproducible and efficient quantification of key microstructural features."
    )
    
    # Methods
    add_section_heading(doc, 'Methods', 1)
    add_paragraph(doc,
        "The implemented Python workflow utilizes the Pillow and NumPy libraries to process all images from the specified input directory whose filenames begin with 'Li_'. "
        "Each image is converted to grayscale, and all pixels with values in the 5–30 range are considered candidates for GAP regions. "
        "A second condition is checked: at least one adjacent direction (up, down, left, or right) must contain 20 contiguous pixels also within the 5–30 grayscale range. "
        "Pixels meeting both criteria are flagged as GAP pixels. "
        "For each image, a CSV file is generated containing the coordinates, grayscale value, and GAP flag for every pixel. "
        "Another CSV file records the GAP height per column, calculated as (max_row - min_row + 1) multiplied by the physical resolution parameter (μm/pixel). "
        "A TXT file summarizes the physical dimension used and the maximum detected GAP height. "
        "Additionally, a PNG image is produced for each input image, highlighting all GAP pixels in red for visual verification. "
        "The output files are used to support the results section of this report."
    )
    
    # Results
    add_section_heading(doc, 'Results', 1)
    image_files = get_image_files(input_dir)
    if not image_files:
        add_paragraph(doc, "No 'Li_' images found in the specified directory.")
    else:
        for fname in image_files:
            base = os.path.splitext(fname)[0]
            txt_report_path = os.path.join(output_dir, f"{base}_result.txt")
            highlighted_img_path = os.path.join(output_dir, f"{base}_gap_highlighted.png")
            gap_height_csv = os.path.join(output_dir, f"{base}_gap_height.csv")
            re, max_gap_height = read_txt_report(txt_report_path)

            # Section per image
            add_section_heading(doc, f"Results for {base}", 2)
            if re is None or max_gap_height is None:
                add_paragraph(doc, "Result TXT file not found or incomplete.")
            else:
                add_paragraph(doc, f"Physical resolution used: {re} μm/pixel")
                add_paragraph(doc, f"Maximum GAP height detected: {max_gap_height} μm")
            if os.path.exists(gap_height_csv):
                add_paragraph(doc,
                    f"Detailed per-column GAP height data is available in the CSV file: {os.path.basename(gap_height_csv)}."
                )
            else:
                add_paragraph(doc, "GAP height CSV not found.")
            # Insert highlighted image
            if os.path.exists(highlighted_img_path):
                add_image(doc, highlighted_img_path, width_inch=4.5,
                          caption=f"GAP pixels (red) highlighted for {base}.")
            else:
                add_paragraph(doc, "Highlighted GAP image not found.")

    # Save report
    doc.save(output_docx)
    print(f"Word report generated: {output_docx}")

if __name__ == "__main__":
    # These paths must match the prior outputs
    input_images_dir = r"C:\Users\admin\Desktop\Python_proj\distance_analysis_new\Images"
    output_files_dir = r"C:\Users\admin\Desktop\Python_proj\ALL_RESULT\GPT\T1S2\backup5"
    output_docx = os.path.join(output_files_dir, "GAP_Simulation_Report.docx")
    generate_report(input_images_dir, output_files_dir, output_docx)
