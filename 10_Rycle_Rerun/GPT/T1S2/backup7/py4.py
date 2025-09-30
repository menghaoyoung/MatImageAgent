import os
import glob
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def find_result_files(output_dir, prefix='Li_'):
    """
    Returns a dict mapping base image name to its result files:
      - gap_analysis_csv
      - gap_height_csv
      - gap_stats_txt
      - gap_highlight_png
    """
    files = os.listdir(output_dir)
    images = {}
    for f in files:
        if f.startswith(prefix) and (f.endswith('_gap_analysis.csv') or
                                     f.endswith('_gap_height.csv') or
                                     f.endswith('_gap_stats.txt') or
                                     f.endswith('_GAP_highlight.png')):
            base = f.split('_gap_')[0]
            if base not in images:
                images[base] = {}
            if f.endswith('_gap_analysis.csv'):
                images[base]['analysis_csv'] = os.path.join(output_dir, f)
            elif f.endswith('_gap_height.csv'):
                images[base]['height_csv'] = os.path.join(output_dir, f)
            elif f.endswith('_gap_stats.txt'):
                images[base]['stats_txt'] = os.path.join(output_dir, f)
            elif f.endswith('_GAP_highlight.png'):
                images[base]['highlight_png'] = os.path.join(output_dir, f)
    return images

def read_stats_txt(stats_txt_path):
    """
    Reads the stats TXT file and returns (resolution, max_height)
    """
    with open(stats_txt_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    re = None
    max_height = None
    for l in lines:
        if 'μm/pixel' in l:
            re = l.strip().split(':')[-1].strip()
        if 'Max GAP height' in l:
            max_height = l.strip().split(':')[-1].strip()
    return re, max_height

def generate_report(output_dir, report_path):
    images = find_result_files(output_dir)
    if not images:
        print("No result files found in output directory.")
        return

    doc = Document()

    # Title
    doc.add_heading('Simulation Report on GAP Pixel Analysis in Li_ Images', 0)

    # Abstract
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        "This report documents the automated analysis of grayscale images with the 'Li_' prefix using a custom Python workflow. The analysis identifies GAP pixels based on grayscale thresholds and spatial continuity, quantifies GAP column heights, and generates annotated result images. The report provides an overview of the computational approach, summarizes the input parameters, and presents the resulting statistics and visuals."
    )

    # Introduction
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph(
        "In high-precision image analysis, the identification of regions with specific grayscale characteristics is critical for material science and biomedical imaging. GAP pixels, defined by their grayscale value and spatial arrangement, can indicate important structural features. This simulation aims to automate the detection and quantification of such regions in a set of images prefixed with 'Li_'."
    )

    # Methods
    doc.add_heading('Methods', level=1)
    doc.add_paragraph(
        "The procedure employs a Python program leveraging the Pillow (PIL) library for image processing, NumPy for pixel manipulation, and CSV for structured output. Each input image is converted to grayscale, after which pixels are evaluated for the GAP condition: grayscale value between 5 and 30 (inclusive) and the presence of at least one adjacent direction with 20 contiguous qualifying pixels. The program outputs per-pixel analysis CSV files, per-column GAP height data, summary statistics, and annotated images with GAP pixels highlighted in red. All results are organized per image in the specified output directory."
    )

    # Results
    doc.add_heading('Results', level=1)
    for base, files in images.items():
        doc.add_heading(f"Results for {base}", level=2)

        # Add stats from TXT
        stats_txt = files.get('stats_txt')
        re_val, max_height = read_stats_txt(stats_txt) if stats_txt else (None, None)
        p = doc.add_paragraph()
        if re_val and max_height:
            p.add_run(f"Physical dimension parameter: {re_val} μm/pixel\n")
            p.add_run(f"Max GAP height: {max_height} μm\n")
        else:
            p.add_run("Statistics not available.\n")

        # Insert highlight image if exists
        highlight_img = files.get('highlight_png')
        if highlight_img and os.path.exists(highlight_img):
            doc.add_paragraph("GAP regions highlighted in red:")
            try:
                doc.add_picture(highlight_img, width=Inches(3.5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            except Exception as e:
                doc.add_paragraph(f"[Unable to insert image: {e}]")
        else:
            doc.add_paragraph("[No highlight image available]")

        doc.add_paragraph()  # Space

    # Save report
    doc.save(report_path)
    print(f"Simulation report generated: {report_path}")

if __name__ == "__main__":
    output_dir = r"C:\Users\admin\Desktop\Python_proj\ALL_RESULT\GPT\T1S2\backup6"
    report_path = os.path.join(output_dir, "GAP_Pixel_Analysis_Report.docx")
    generate_report(output_dir, report_path)
