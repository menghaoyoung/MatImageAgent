import os
import docx
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Directories and file patterns
output_dir = r"C:\Users\admin\Desktop\Python_proj\ALL_RESULT\GPT\T1S2\backup8"
img_input_dir = r"C:\Users\admin\Desktop\Python_proj\distance_analysis_new\Images"

def get_gap_files():
    # Collect all analysis, gap_height, txt, and GAP image files for "Li_" images
    file_dict = {}
    for fname in os.listdir(output_dir):
        if fname.startswith("Li_") and fname.endswith("_gap_analysis.csv"):
            img_base = fname[:-len("_gap_analysis.csv")]
            file_dict[img_base] = {
                "analysis_csv": os.path.join(output_dir, f"{img_base}_gap_analysis.csv"),
                "gap_height_csv": os.path.join(output_dir, f"{img_base}_gap_height.csv"),
                "gap_txt": os.path.join(output_dir, f"{img_base}_gap_height.txt"),
                "gap_png": os.path.join(output_dir, f"{img_base}_GAP.png"),
            }
    return file_dict

def extract_stats_from_txt(txt_path):
    # Reads the TXT for dimension and max height
    stats = {}
    with open(txt_path, 'r', encoding='utf-8') as f:
        for line in f:
            if 'Physical dimension parameter' in line:
                stats['resolution'] = line.strip().split(':')[-1].strip()
            if 'Max GAP height' in line:
                stats['max_height'] = line.strip().split(':')[-1].strip()
    return stats

def create_report(file_dict, doc_path):
    doc = docx.Document()
    # Title
    doc.add_heading("Simulation Report: Automated GAP Height Analysis Based on Grayscale Imaging", 0)

    # Abstract
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "This report presents a fully automated pipeline for the identification and quantification of GAP regions in microscopy images using grayscale intensity analysis. "
        "By leveraging Python and the Pillow library, the program detects GAP pixels and computes physical height distributions per image column. The results facilitate rapid, repeatable measurement of microstructural features, as demonstrated on sample images. "
        "Statistical results and visual highlights are provided to support analytical conclusions."
    )

    # Introduction
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(
        "Precise measurement of microstructural gaps is crucial in materials science and biological imaging. Manual measurement is tedious and error-prone, motivating the need for automated, reproducible approaches. "
        "This simulation aims to extract accurate GAP region heights based on grayscale thresholds and spatial continuity criteria, outputting comprehensive pixel-level data, per-column statistics, and visual overlays. "
        "The approach accelerates the quantitative analysis of experimental images and ensures consistent results."
    )

    # Methods
    doc.add_heading("Methods", level=1)
    doc.add_paragraph(
        "All input images with prefix 'Li_' in PNG or JPG format are loaded from the specified directory. Images are converted to grayscale using the Pillow library. "
        "For each pixel, the program checks if its grayscale value is between 5 and 30 (inclusive), and whether there exists at least one adjacent pixel (up, down, left, or right) with 20 contiguous pixels also meeting the grayscale threshold. "
        "Pixels satisfying these conditions are flagged as GAP. The program records pixel coordinates, grayscale value, and GAP flag into a CSV file. "
        "For each image column, the minimum and maximum row indices of GAP pixels are used to calculate the GAP height as (max_row-min_row+1)×resolution (μm). "
        "All results are saved as CSV (pixel analysis and per-column GAP heights), TXT (resolution and max height), and PNG (original image with GAP pixels highlighted in red). "
        "Document generation and results compilation are performed using the python-docx library."
    )

    # Results
    doc.add_heading("Results", level=1)
    for img_base, files in file_dict.items():
        doc.add_heading(f"Results for {img_base}", level=2)
        # Add TXT stats
        stats = extract_stats_from_txt(files['gap_txt'])
        p = doc.add_paragraph()
        p.add_run(f"Physical resolution: ").bold = True
        p.add_run(f"{stats.get('resolution', 'N/A')}\n")
        p.add_run("Maximum GAP height: ").bold = True
        p.add_run(f"{stats.get('max_height', 'N/A')}\n")
        # Add the GAP PNG image
        if os.path.exists(files['gap_png']):
            doc.add_paragraph("Detected GAP pixels are highlighted in red as shown below:")
            # Resize for document: scale to width 4" (maintain aspect ratio)
            doc.add_picture(files['gap_png'], width=Inches(4.0))
        else:
            doc.add_paragraph("GAP visualization image not found.")
        doc.add_paragraph("")  # Blank line

    doc.add_page_break()
    doc.save(doc_path)
    print(f"Simulation report generated: {doc_path}")

if __name__ == "__main__":
    # Collect all outputs for "Li_" images
    file_dict = get_gap_files()
    if not file_dict:
        print("No output files found. Please run GAP analysis first.")
    else:
        report_path = os.path.join(output_dir, "GAP_Simulation_Report.docx")
        create_report(file_dict, report_path)
