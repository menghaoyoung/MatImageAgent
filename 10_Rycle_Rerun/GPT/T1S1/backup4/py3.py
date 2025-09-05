import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def get_gap_outputs(output_dir):
    # Get base names of all images processed (from missing output list or from actual files)
    files = os.listdir(output_dir)
    csvs = [f for f in files if f.endswith('_gap_analysis.csv')]
    pngs = [f for f in files if f.endswith('_gap_highlight.png')]

    base_names = set()
    for f in csvs:
        base = f.replace('_gap_analysis.csv', '')
        base_names.add(base)
    for f in pngs:
        base = f.replace('_gap_highlight.png', '')
        base_names.add(base)

    # Only include those with both csv and png
    result = []
    for base in sorted(base_names):
        csv_path = os.path.join(output_dir, base + '_gap_analysis.csv')
        png_path = os.path.join(output_dir, base + '_gap_highlight.png')
        if os.path.isfile(csv_path) and os.path.isfile(png_path):
            result.append({'base': base, 'csv': csv_path, 'png': png_path})
    return result

def count_gap_pixels(csv_path):
    count = 0
    total = 0
    with open(csv_path, 'r') as f:
        next(f)
        for line in f:
            parts = line.strip().split(',')
            if len(parts) == 4:
                total += 1
                if parts[3] == '1':
                    count += 1
    return count, total

def add_image_with_caption(doc, img_path, caption):
    doc.add_picture(img_path, width=Inches(3.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p = doc.add_paragraph(caption)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def generate_report(output_dir, report_path):
    doc = Document()
    doc.add_heading('Automated GAP Pixel Detection in Grayscale Images', 0)

    # Abstract
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        "This report documents the automated analysis of a series of grayscale images to identify and highlight GAP pixels. "
        "GAP pixels are defined as pixels whose grayscale value lies within the range 5–30 (inclusive) and which are adjacent "
        "to a run of at least 20 contiguous pixels, in one of four directions, also satisfying the grayscale criterion. "
        "The analysis includes per-pixel data output and visualization of GAP pixels, supporting further research in image-based "
        "feature detection or quality control."
    )

    # Introduction
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph(
        "Image analysis is a crucial component in scientific and industrial workflows, enabling automated detection of patterns, "
        "defects, or regions of interest. In this task, the goal was to systematically process a batch of images with filenames "
        "starting with 'Li_', extract pixel-level grayscale information, and identify GAP pixels based on strict adjacency and "
        "grayscale continuity criteria. The processed results facilitate visual inspection and further quantitative analysis."
    )

    # Methods
    doc.add_heading('Methods', level=1)
    doc.add_paragraph(
        "All input images were read from the specified folder, filtered by the 'Li_' prefix and PNG/JPG format. Each image "
        "was converted to grayscale using the Pillow library. For each pixel, its grayscale value was checked to determine "
        "if it lay within the 5 to 30 range. For such pixels, the program assessed their immediate neighbors in four cardinal "
        "directions (up, down, left, right) to identify if in any direction, a contiguous segment of 20 pixels (including the neighbor) "
        "also met the grayscale condition. Pixels satisfying both criteria were flagged as GAP pixels. For each image, a CSV file "
        "was generated recording every pixel's coordinates, grayscale value, and GAP flag, and a new PNG image was saved, highlighting "
        "GAP pixels in red. Python automation ensured reproducibility and scalability of the workflow."
    )

    # Results
    doc.add_heading('Results', level=1)
    outputs = get_gap_outputs(output_dir)
    if not outputs:
        doc.add_paragraph("No output images or CSV results were found in the specified output directory. Please ensure the analysis has been run successfully.")
    else:
        for out in outputs:
            gap_count, total_count = count_gap_pixels(out['csv'])
            percent = (gap_count/total_count*100) if total_count else 0
            doc.add_paragraph(
                f"Image: {out['base']}\n"
                f"  - Total pixels: {total_count}\n"
                f"  - GAP pixels: {gap_count} ({percent:.2f}%)"
            )
            add_image_with_caption(doc, out['png'], f"GAP Highlighted Result: {out['base']}")

        doc.add_paragraph(
            "The figures above visually highlight GAP pixels (in red) on the processed grayscale images. "
            "The quantitative results show the frequency and distribution of GAP pixels, which may correlate with image features or artifacts "
            "relevant to downstream analysis."
        )

    doc.save(report_path)
    print(f"Report generated: {report_path}")

if __name__ == "__main__":
    output_dir = r"C:\Users\admin\Desktop\Python_proj\ALL_RESULT\GPT\T1S1\backup4"
    report_path = os.path.join(output_dir, "GAP_Pixel_Analysis_Report.docx")
    generate_report(output_dir, report_path)
