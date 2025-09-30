import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def get_image_and_csv_files(output_dir):
    # Collect all CSV and highlight image pairs
    files = os.listdir(output_dir)
    csv_files = [f for f in files if f.endswith('_gap_analysis.csv')]
    img_files = [f for f in files if f.endswith('_gap_highlight.png')]
    base_names = [os.path.splitext(f)[0].replace('_gap_analysis', '') for f in csv_files]
    pairs = []
    for base in base_names:
        csv_path = os.path.join(output_dir, f"{base}_gap_analysis.csv")
        img_path = os.path.join(output_dir, f"{base}_gap_highlight.png")
        if os.path.isfile(csv_path) and os.path.isfile(img_path):
            pairs.append((csv_path, img_path, base))
    return pairs

def count_gap_pixels_from_csv(csv_path):
    count = 0
    total = 0
    with open(csv_path, encoding="utf-8") as f:
        for idx, line in enumerate(f):
            if idx == 0:
                continue  # skip header
            total += 1
            if line.strip().endswith(",1"):
                count += 1
    return count, total

def write_simulation_report(output_dir, doc_path):
    pairs = get_image_and_csv_files(output_dir)
    if not pairs:
        print("No analysis results found in output directory.")
        return

    doc = Document()
    doc.add_heading("Simulation Report: GAP Detection in Grayscale Images", 0)

    # Abstract
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph((
        "This simulation report presents the analysis and detection of GAP pixels in a series of grayscale images. "
        "Using a custom Python program, input images were evaluated for specific grayscale value criteria and spatial continuity conditions. "
        "The results, including per-pixel data and highlighted images, provide insight into the distribution and prevalence of GAP regions. "
        "The simulation aims to assist researchers in understanding the structural properties of the analyzed images."
    ))

    # Introduction
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph((
        "Image processing and pattern recognition are fundamental in fields such as materials science, biomedical imaging, and computer vision. "
        "In this simulation, we focus on identifying 'GAP' pixels in grayscale images, where a GAP pixel is defined as having a grayscale value between 5 and 30, "
        "and being adjacent to a direction with at least 20 contiguous pixels also within this grayscale range. "
        "By automating this process, we aim to efficiently locate critical regions of interest, which could be indicative of gaps, defects, or relevant structural features."
    ))

    # Methods
    doc.add_heading("Methods", level=1)
    doc.add_paragraph((
        "The analysis was performed using a Python script leveraging the Pillow and NumPy libraries. "
        "Each image file prefixed with 'Li_' from the specified directory was converted to grayscale. "
        "For every pixel, the program checked whether its grayscale intensity was between 5 and 30 (inclusive). "
        "Additionally, for each pixel, the script evaluated if at least one of its four main neighboring directions "
        "(up, down, left, right) contained 20 consecutive pixels (including the neighbor) within the same grayscale threshold. "
        "Pixels meeting both criteria were flagged as GAP pixels. "
        "The script produced a CSV file per image, documenting the row, column, grayscale value, and GAP flag per pixel, "
        "and also generated a new image highlighting all GAP pixels in red (RGB: 255, 0, 0)."
    ))

    # Results
    doc.add_heading("Results", level=1)
    doc.add_paragraph((
        f"Analysis was performed on {len(pairs)} images with results summarized below. "
        "For each image, the number of detected GAP pixels and a visual representation are provided."
    ))

    for csv_path, img_path, base in pairs:
        gap_count, total_pixels = count_gap_pixels_from_csv(csv_path)
        percent = (gap_count / total_pixels * 100) if total_pixels else 0
        doc.add_heading(f"Image: {base}", level=2)
        doc.add_paragraph(f"GAP pixels: {gap_count} / {total_pixels} ({percent:.2f}%)")
        # Insert the highlight image
        try:
            doc.add_picture(img_path, width=Inches(3.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        except Exception as e:
            doc.add_paragraph(f"[Unable to insert image: {e}]")

    doc.add_page_break()
    doc.save(doc_path)
    print(f"Simulation report generated: {doc_path}")

if __name__ == "__main__":
    # Output directory and report path as specified
    output_dir = r"C:\Users\admin\Desktop\Python_proj\ALL_RESULT\GPT\T1S1\backup9"
    report_docx = os.path.join(output_dir, "GAP_Simulation_Report.docx")
    write_simulation_report(output_dir, report_docx)
