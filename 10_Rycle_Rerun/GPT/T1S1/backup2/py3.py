import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_image(doc, image_path, width_inch=4):
    try:
        doc.add_picture(image_path, width=Inches(width_inch))
        last_paragraph = doc.paragraphs[-1] 
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    except Exception as e:
        doc.add_paragraph(f"[Image not shown: {os.path.basename(image_path)}] ({e})")

def main():
    # Set the output directory (from py1.py outputs)
    output_dir = r"C:\Users\admin\Desktop\Python_proj\ALL_RESULT\GPT\T1S1\backup3"
    # Find all highlighted images
    img_files = [f for f in os.listdir(output_dir) if f.startswith("Li_") and f.endswith("_gap_highlight.png")]
    csv_files = [f for f in os.listdir(output_dir) if f.startswith("Li_") and f.endswith("_gap_analysis.csv")]

    # Create the document
    doc = Document()
    doc.add_heading("Simulation Report: Automated Detection and Highlighting of GAP Pixels in Images", 0)

    # Abstract
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph((
        "This simulation report presents the application and outcomes of a Python-based algorithm for automatic detection of 'GAP' pixels in a set of scientific images. "
        "The methodology leverages pixel-level grayscale analysis and adjacency logic to identify and visually highlight critical regions within each image. "
        "The results, as demonstrated by the generated images and data, provide a systematic approach for rapid and reproducible GAP pixel analysis."
    ))

    # Introduction
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph((
        "Image analysis is a crucial tool in scientific research, allowing for quantitative and qualitative assessment of experimental results. "
        "Manual pixel inspection is not only tedious but prone to error, especially in high-resolution datasets. "
        "To address this, we implemented an automated workflow to detect 'GAP' pixels, defined by specific grayscale intensity ranges and spatial continuity among neighbors. "
        "This report summarizes the approach and presents key findings from the image data processed in this simulation."
    ))

    # Methods
    doc.add_heading("Methods", level=1)
    doc.add_paragraph((
        "The Python program processes all images in the specified folder starting with the prefix 'Li_', supporting both PNG and JPEG formats. "
        "Each image is first converted to grayscale using the Pillow library. For each pixel, the grayscale intensity is extracted and checked against two conditions: "
        "(1) the value lies within the range 5–30 (inclusive), and (2) at least one adjacent direction (up, down, left, right) contains 20 contiguous pixels also meeting this grayscale criterion. "
        "Pixels meeting both conditions are flagged as 'GAP'. "
        "For each image, a CSV file is generated documenting the row, column, grayscale value, and GAP flag for every pixel. "
        "Additionally, a new PNG image is generated, where all GAP pixels are marked in red (RGB: 255, 0, 0) for rapid visualization. "
        "All output files are stored in the specified output directory."
    ))

    # Results
    doc.add_heading("Results", level=1)
    if not img_files:
        doc.add_paragraph("No highlighted images were generated. Please ensure py1.py was executed and images were available for analysis.")
    else:
        doc.add_paragraph(
            "The automated analysis successfully processed the available images. "
            "For each input image, a corresponding CSV file of pixel data and a highlighted PNG image were generated. "
            "The following figures show the images with GAP pixels marked in red, enabling immediate identification of regions of interest."
        )
        for img_file in img_files:
            doc.add_heading(f"Highlighted GAPs: {img_file}", level=2)
            add_image(doc, os.path.join(output_dir, img_file))
            # Optionally, add a reference to the CSV file
            csv_match = img_file.replace('_gap_highlight.png', '_gap_analysis.csv')
            if csv_match in csv_files:
                doc.add_paragraph(f"See data: {csv_match}")
    # Save the document
    report_path = os.path.join(output_dir, "Simulation_Report_GAP_Pixel_Analysis.docx")
    doc.save(report_path)
    print(f"Word simulation report generated at: {report_path}")

if __name__ == "__main__":
    main()
