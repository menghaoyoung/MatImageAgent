import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def get_gap_images_and_csvs(output_dir):
    """
    Return a list of (basename, csv_path, img_path) tuples for each image with GAP analysis results.
    """
    files = os.listdir(output_dir)
    basenames = set()
    for f in files:
        if f.startswith('Li_') and f.endswith('_gap_analysis.csv'):
            base = f[:-len('_gap_analysis.csv')]
            basenames.add(base)
    results = []
    for base in sorted(basenames):
        csv_path = os.path.join(output_dir, f"{base}_gap_analysis.csv")
        img_path = os.path.join(output_dir, f"{base}_gap_highlight.png")
        if os.path.exists(csv_path) and os.path.exists(img_path):
            results.append((base, csv_path, img_path))
    return results

def count_gap_pixels(csv_path):
    """
    Count total pixels and GAP=1 pixels from the CSV.
    """
    import csv
    total = 0
    gap = 0
    with open(csv_path, newline='', encoding='utf-8') as f:
        reader = csv.DictReader(f)
        for row in reader:
            total += 1
            if row['GAP_flag'] == '1':
                gap += 1
    return total, gap

def write_simulation_report(output_dir, report_path):
    """
    Generate a simulation report as a Word document according to the instructions.
    """
    gap_results = get_gap_images_and_csvs(output_dir)
    document = Document()

    # Title
    document.add_heading("Simulation Analysis of GAP Pixels in Li_ Images", 0)

    # Abstract
    document.add_heading("Abstract", level=1)
    document.add_paragraph(
        "This report presents a detailed analysis of GAP pixels in a series of images prefixed with 'Li_'. "
        "Using automated Python scripts, each image was processed to identify pixels whose grayscale values "
        "fall within the range of 5 to 30 and which are adjacent to a contiguous sequence of 20 similar pixels. "
        "The results, including pixel-level CSV data and highlighted images, are summarized herein."
    )

    # Introduction
    document.add_heading("Introduction", level=1)
    document.add_paragraph(
        "The identification of specific pixel patterns in scientific images is crucial for quantitative analysis, "
        "especially in material science and related fields. GAP pixels, as defined by a strict grayscale criterion "
        "and local spatial continuity, indicate regions of interest that may correspond to microstructural features "
        "or defects. This analysis automates the detection of such pixels in a batch of images, facilitating large-scale "
        "and objective quantification."
    )

    # Methods
    document.add_heading("Methods", level=1)
    document.add_paragraph(
        "All images with filenames beginning with 'Li_' and in PNG or JPG format were processed from the specified directory. "
        "Each image was converted to grayscale using the Pillow library. For every pixel, its grayscale value was extracted. "
        "A pixel was marked as a GAP pixel if its grayscale value was between 5 and 30 (inclusive) and if at least one "
        "adjacent direction (up, down, left, or right) contained a contiguous line of 20 pixels also meeting the grayscale condition. "
        "For each image, a CSV file was generated, recording the row and column, grayscale value, and GAP flag for all pixels. "
        "Additionally, a new PNG image was produced where all GAP pixels were highlighted in red (RGB: 255, 0, 0), while other pixels "
        "retained their original grayscale color. The processing was performed using custom Python scripts employing PIL and numpy for "
        "image manipulation and csv for data output."
    )

    # Results
    document.add_heading("Results", level=1)
    if gap_results:
        for base, csv_path, img_path in gap_results:
            total_pixels, gap_pixels = count_gap_pixels(csv_path)
            percent = 100.0 * gap_pixels / total_pixels if total_pixels > 0 else 0
            document.add_paragraph(f"Image: {base}", style='List Bullet')
            document.add_paragraph(
                f"Total pixels: {total_pixels:,}\n"
                f"GAP pixels: {gap_pixels:,} ({percent:.3f}%)", style='List Continue'
            )
            # Insert image
            if os.path.exists(img_path):
                # Try to insert at a reasonable size
                document.add_picture(img_path, width=Inches(3.0))
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        document.add_paragraph("No GAP analysis results found in the specified output directory.")

    # Save
    document.save(report_path)
    print(f"Simulation report generated: {report_path}")

if __name__ == "__main__":
    output_dir = r"C:\Users\admin\Desktop\Python_proj\ALL_RESULT\GPT\T1S1\backup7"
    report_path = os.path.join(output_dir, "GAP_Pixel_Simulation_Report.docx")
    write_simulation_report(output_dir, report_path)
