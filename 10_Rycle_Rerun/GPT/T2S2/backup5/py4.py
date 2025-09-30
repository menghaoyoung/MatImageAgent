import os
import sys
import io
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Ensure utf-8 output for Windows console
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

# File paths and parameters (as per previous steps)
image_path = r"C:\Users\admin\Desktop\Python_proj\datas\T2_IMGS\Li_1.0.png"
output_dir = r"C:\Users\admin\Desktop\Python_proj\ALL_RESULT\GPT\T2S2\1.0\backup5"
base_filename = os.path.splitext(os.path.basename(image_path))[0]
gray_csv_path = os.path.join(output_dir, f"{base_filename}_line_gray.csv")
length_txt_path = os.path.join(output_dir, f"{base_filename}_line_length.txt")
ueq_csv_path = os.path.join(output_dir, f"{base_filename}_line_ueq.csv")
tiff_path = os.path.join(output_dir, f"{base_filename}_ueq_curve.tiff")
report_path = os.path.join(output_dir, f"{base_filename}_simulation_report.docx")

# Read the measured line segment length
with open(length_txt_path, 'r', encoding='utf-8') as f:
    line_length_text = f.read().strip()

# Read the grayscale and u_eq data
import pandas as pd
gray_df = pd.read_csv(gray_csv_path)
ueq_df = pd.read_csv(ueq_csv_path)

# Calculate summary stats for results section
gray_min = gray_df['GrayValue'].min()
gray_max = gray_df['GrayValue'].max()
u_min = ueq_df['u_eq'].astype(float).min()
u_max = ueq_df['u_eq'].astype(float).max()
total_points = len(gray_df)

# Start writing Word document
doc = Document()

# Title
title = "Simulation Report: Quantitative Line Profile Analysis of Li_1.0.png"
doc.add_heading(title, 0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# --- Abstract ---
abstract = (
    "This simulation report presents a detailed quantitative analysis of the grayscale "
    "intensity profile along a specified line segment in the image 'Li_1.0.png'. "
    "Utilizing a custom Python-based approach, the grayscale values were extracted along the "
    "user-defined line, converted to equivalent u_eq values, and comprehensively analyzed. The "
    "methodology, results, and scientific interpretation are discussed in detail, providing "
    "insight into the spatial variation of grayscale and u_eq along the segment. The data and "
    "resulting plots serve as a foundation for further material or image-based investigations."
)
doc.add_heading('Abstract', level=1)
doc.add_paragraph(abstract)

# --- Introduction ---
introduction = (
    "Accurate analysis of grayscale intensity profiles is critical in a variety of scientific "
    "and engineering fields, including materials science, microscopy, and image-based diagnostics. "
    "In this study, we focus on the image 'Li_1.0.png', performing a quantitative profile analysis "
    "along a carefully selected line segment. The purpose of this simulation is to extract, "
    "visualize, and quantify the grayscale variations and their corresponding u_eq values along "
    "the segment. These results facilitate a deeper understanding of local inhomogeneities, "
    "textural features, or compositional gradients within the image. The approach enables "
    "researchers to connect image data to underlying physical or chemical properties, enhancing "
    "the interpretability of imaging datasets."
)
doc.add_heading('Introduction', level=1)
doc.add_paragraph(introduction)

# --- Methods ---
methods = (
    "The analysis was conducted using a custom Python script (py1.py), which utilizes image "
    "processing libraries such as Pillow, NumPy, and matplotlib. The process began by specifying "
    "the start and end points of the line segment as (152, 29) and (135, 92), respectively, in "
    "pixel coordinates. The image resolution was set to 1.08 μm/pixel. Using Bresenham's algorithm, "
    "all pixel coordinates along the line were determined. The grayscale values (0-255) were "
    "extracted from the image at these coordinates and saved into a CSV file. The physical length "
    "of the line segment was calculated and stored in a text file. Subsequently, u_eq values were "
    "computed using the formula: u_eq = u_min + (gray_value / 255) * u_max, with u_min=0 and "
    "u_max=65535. The u_eq profile was plotted against the distance from the start point and saved "
    "as a TIFF image. All relevant numerical data were exported for further analysis. This "
    "programmatic approach ensures reproducibility, precision, and efficient data handling."
)
doc.add_heading('Methods', level=1)
doc.add_paragraph(methods)

# --- Results ---
results = (
    f"The analysis extracted a total of {total_points} grayscale values along the defined line "
    f"segment. The measured length of the segment was: {line_length_text}. Grayscale values "
    f"ranged from {gray_min} to {gray_max}, corresponding to u_eq values between {int(u_min)} and {int(u_max)}. "
    "The u_eq profile curve (as shown in Fig. 1) demonstrates the spatial variation of the signal "
    "along the line, highlighting regions of notable increase or decrease in intensity. These "
    "variations may relate to underlying material characteristics or imaging artifacts. The "
    "exported CSV files provide a detailed record of both the grayscale and u_eq values for each "
    "position along the line, supporting further statistical or physical analysis. The TIFF-format "
    "plot ensures high-quality visualization suitable for publication or presentations."
)
doc.add_heading('Results', level=1)
doc.add_paragraph(results)

# Insert the figure (TIFF plot)
if os.path.exists(tiff_path):
    doc.add_paragraph('Fig. 1. u_eq vs Distance from Start Point for Li_1.0.png.')
    doc.add_picture(tiff_path, width=Inches(5.5))

# Save the document
doc.save(report_path)
print(f"Simulation report has been generated and saved to: {report_path}")
