import os
import sys
import csv
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    # Define file paths
    image_path = r"C:\Users\admin\Desktop\Python_proj\datas\T2_IMGS\Li_1.0.png"
    output_dir = r"C:\Users\admin\Desktop\Python_proj\ALL_RESULT\GPT\T2S2\1.0\backup9"
    img_filename = os.path.splitext(os.path.basename(image_path))[0]
    csv_gray_path = os.path.join(output_dir, f"{img_filename}_line_gray.csv")
    txt_len_path = os.path.join(output_dir, f"{img_filename}_line_length.txt")
    csv_u_eq_path = os.path.join(output_dir, f"{img_filename}_distance_u_eq.csv")
    tiff_plot_path = os.path.join(output_dir, f"{img_filename}_u_eq_curve.tiff")
    docx_path = os.path.join(output_dir, f"{img_filename}_simulation_report.docx")

    # Read length file robustly
    length_lines = []
    for enc in ['utf-8', 'utf-8-sig', 'gbk', 'gb2312', 'latin1']:
        try:
            with open(txt_len_path, "r", encoding=enc) as f:
                length_lines = f.readlines()
            break
        except Exception:
            continue
    if not length_lines:
        print("Failed to read line length file with known encodings.")
        return

    line_length_str = ""
    for line in length_lines:
        if "μm" in line or "um" in line or "u" in line:
            line_length_str = line.strip()

    # Read u_eq and distances
    distances = []
    u_eqs = []
    with open(csv_u_eq_path, "r", encoding='utf-8', errors='ignore') as f:
        reader = csv.reader(f)
        headers = next(reader)
        for row in reader:
            try:
                distances.append(float(row[0]))
                u_eqs.append(float(row[1]))
            except Exception:
                continue
    n_points = len(distances)
    u_eq_min = min(u_eqs) if u_eqs else 0
    u_eq_max = max(u_eqs) if u_eqs else 0
    u_eq_avg = sum(u_eqs) / n_points if n_points > 0 else 0

    # Create Word document
    document = Document()
    document.add_heading('Simulation and Analysis of Line Grayscale and Equivalent Potential (μ_eq) Extraction from Lithium Image', 0)

    # Abstract
    abstract = (
        "This report presents a simulation study involving the extraction of grayscale values along a specified line segment "
        "within a lithium sample image and the computation of equivalent potential (μ_eq) based on these grayscale values. "
        "Utilizing Python-based image analysis and data processing, the approach facilitates precise measurement of physical "
        "properties at the microscale. The workflow integrates automated extraction of pixel intensity profiles, conversion of "
        "grayscale to equivalent physical parameters, and visualization of resulting spatial distributions. Results are "
        "systematically documented, ensuring reproducibility and clarity for further scientific analysis."
    )
    document.add_heading('Abstract', level=1)
    document.add_paragraph(abstract)

    # Introduction
    introduction = (
        "Quantitative analysis of grayscale variations and corresponding physical property gradients is crucial in materials "
        "science, especially when investigating electrochemical phenomena in lithium-based systems. By mapping grayscale "
        "intensity along arbitrary line segments within high-resolution images, researchers can infer local changes in "
        "composition, potential, or other relevant parameters. This simulation aims to demonstrate an automated approach "
        "for extracting and analyzing such line profiles using a specified resolution (1.08 μm/pixel). The method enables "
        "conversion of image pixel data into real-world spatial measurements and derived properties like μ_eq, providing "
        "insights into microstructural or electrochemical heterogeneities within the sample."
    )
    document.add_heading('Introduction', level=1)
    document.add_paragraph(introduction)

    # Methods
    methods = (
        "The simulation leverages Python's imaging and scientific libraries to automate data extraction and analysis. "
        "First, the program identifies all pixel coordinates along the user-defined line segment using Bresenham's algorithm. "
        "Grayscale values (0-255) are then sampled from the image at each point along this line and saved as a CSV file. "
        "The physical length of the line is calculated by multiplying pixel distance by the given resolution (1.08 μm/pixel) "
        f"({line_length_str}). Equivalent potential (μ_eq) values are computed for each pixel using the formula μ_eq = u_min + "
        "(gray_value / 255) * u_max, where u_min = 0 and u_max = 65535. The μ_eq values and their corresponding distances "
        "from the start point are saved for further analysis. A plot of μ_eq versus distance is generated and exported as a "
        "TIFF image. All intermediate and final results are named according to the source image, ensuring traceability "
        "and data management integrity."
    )
    document.add_heading('Methods', level=1)
    document.add_paragraph(methods)

    # Results
    results = (
        f"The analysis produced a total of {n_points} data points along the line segment from the image '{img_filename}.png'. "
        "The grayscale data captured local variations in pixel intensity, which were translated into physical μ_eq values "
        f"ranging from {u_eq_min:.2f} to {u_eq_max:.2f}, with an average of {u_eq_avg:.2f}. The computed line length is "
        f"{line_length_str.split(':')[-1].strip() if line_length_str else 'unknown'}. The μ_eq versus distance curve, "
        "shown in Fig. 1, reveals the spatial distribution of equivalent potential along the sampled segment. These results "
        "demonstrate the effectiveness of the workflow for detailed microscale analysis, offering a reproducible and "
        "automated path from raw image data to interpretable physical profiles. The methodology can be readily extended "
        "to similar studies involving other materials or imaging modalities."
    )
    document.add_heading('Results', level=1)
    document.add_paragraph(results)

    # Insert Figure
    document.add_paragraph("Figure 1. μ_eq vs. distance for the Li_1.0 sample line segment.")
    if os.path.exists(tiff_plot_path):
        try:
            document.add_picture(tiff_plot_path, width=Inches(5))
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            document.add_paragraph(f"(Figure could not be inserted due to: {e})")

    # Save the document
    document.save(docx_path)
    print(f"Word simulation report generated: {docx_path}")

if __name__ == '__main__':
    main()
