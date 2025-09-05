import os
import csv
import numpy as np
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def read_segment_length(txt_path):
    with open(txt_path, 'r', encoding='utf-8') as f:
        for line in f:
            if "Line segment length" in line:
                # Extract the length and unit
                parts = line.split(":")
                if len(parts) > 1:
                    value_unit = parts[1].strip()
                    value_str = value_unit.split()[0]
                    try:
                        value = float(value_str)
                        unit = value_unit.split()[1] if len(value_unit.split()) > 1 else ""
                        return value, unit
                    except Exception:
                        continue
    return None, None

def stats_from_csv(csv_path):
    vals = []
    with open(csv_path, 'r', encoding='utf-8') as f:
        reader = csv.reader(f)
        next(reader)  # skip header
        for row in reader:
            vals.append(float(row[-1]))
    vals = np.array(vals)
    return {
        'min': float(np.min(vals)),
        'max': float(np.max(vals)),
        'mean': float(np.mean(vals)),
        'std': float(np.std(vals)),
        'n': int(len(vals))
    }

def insert_figure(doc, img_path, caption):
    doc.add_picture(img_path, width=Inches(5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(caption).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def main():
    # Paths and filenames
    image_path = r"C:\Users\admin\Desktop\Python_proj\datas\T2_IMGS\Li_1.0.png"
    output_dir = r"C:\Users\admin\Desktop\Python_proj\ALL_RESULT\GPT\T2S2\1.0\backup"
    base_name = os.path.splitext(os.path.basename(image_path))[0]

    gray_csv = os.path.join(output_dir, f"{base_name}_line_gray.csv")
    length_txt = os.path.join(output_dir, f"{base_name}_line_length.txt")
    ueq_csv = os.path.join(output_dir, f"{base_name}_line_u_eq.csv")
    tiff_path = os.path.join(output_dir, f"{base_name}_u_eq_vs_distance.tiff")

    # Read stats for report
    segment_length, length_unit = read_segment_length(length_txt)
    gray_stats = stats_from_csv(gray_csv)
    ueq_stats = stats_from_csv(ueq_csv)

    # Create Word doc
    doc = Document()
    doc.add_heading('Simulation and Analysis of Line Segment Grayscale and $u_{eq}$ Distribution', 0)

    # Abstract (≈200 words)
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        "This report presents a comprehensive simulation and analysis of grayscale intensity values "
        "and the corresponding equivalent potential ($u_{eq}$) along a defined line segment within a high-resolution image. "
        "Utilizing Python-based image processing tools, grayscale data were extracted for all pixels along the segment, "
        "with physical length calibrated according to pixel resolution. The grayscale data were then mapped to $u_{eq}$ values "
        "using a linear transformation. Results, including the $u_{eq}$ profile and descriptive statistics, are visualized and "
        "discussed. The workflow demonstrates automated extraction, transformation, and documentation, offering a robust approach "
        "for quantitative line profile analysis in scientific images. The methodology and results are relevant for applications "
        "requiring local intensity or potential mapping, such as material science and imaging-based diagnostics. "
        "The generated report includes tabular data, summary statistics, and graphical illustrations based on the simulation, "
        "providing a detailed and reproducible account of the process and outcomes."
    )

    # Introduction (≈200 words)
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph(
        "Accurate analysis of line profiles within images is essential for understanding spatial variations in properties "
        "such as intensity, potential, or material composition. In scientific imaging, quantifying the grayscale values along "
        "a specific path can yield critical information about gradients, interfaces, and local heterogeneities. This simulation "
        "focuses on a segment drawn between two points in a high-resolution image, utilizing the pixel values to infer an "
        "equivalent potential ($u_{eq}$) profile. The conversion from grayscale to $u_{eq}$ is based on a linear mapping between "
        "the minimum and maximum possible values, reflecting the underlying physical model. The approach allows for detailed, "
        "spatially resolved analysis, which is applicable to a variety of image-based scientific studies, including materials "
        "characterization and electronic structure mapping. This report details the automation of the data extraction, "
        "transformation, and visualization processes, thereby providing a reproducible template for similar analyses."
    )

    # Methods (≈200 words)
    doc.add_heading('Methods', level=1)
    doc.add_paragraph(
        "The analysis was performed using a custom Python workflow. First, a PNG image file was loaded and converted to grayscale "
        "to standardize intensity measurements. The Bresenham algorithm was employed to trace integer pixel coordinates along the "
        "line segment defined by the specified start and end points. Grayscale values for each pixel along the segment were "
        "extracted and saved to a CSV file. The Euclidean distance between the endpoints was computed in pixel units and converted "
        "to micrometers using the provided resolution (0.9 μm/pixel). Next, grayscale values were linearly mapped to $u_{eq}$ "
        "using the equation $u_{eq} = u_{min} + (gray/255) \cdot (u_{max} - u_{min})$, with $u_{min}=0$ and $u_{max}=65000$. "
        "The $u_{eq}$ profile was saved as a CSV file, and its values were plotted against distance using matplotlib, generating "
        "a TIFF image for documentation. All data files and visualizations were named according to the source image, and the "
        "pipeline was fully automated, including the creation of this Word report using the python-docx library."
    )

    # Results (≈200 words)
    doc.add_heading('Results', level=1)
    results_para = (
        f"The analyzed line segment spanned {segment_length:.2f} {length_unit}, "
        f"covering {gray_stats['n']} pixels between the specified start and end points. "
        f"The grayscale values along the segment ranged from {gray_stats['min']:.1f} to {gray_stats['max']:.1f}, "
        f"with a mean of {gray_stats['mean']:.1f} and standard deviation of {gray_stats['std']:.1f}. "
        f"After transformation, $u_{{eq}}$ values ranged from {ueq_stats['min']:.1f} to {ueq_stats['max']:.1f}, "
        f"with a mean of {ueq_stats['mean']:.1f} and standard deviation of {ueq_stats['std']:.1f}. "
        "The $u_{eq}$ profile is plotted in Fig. 1, showing spatial variations along the line, which may correspond to material or intensity inhomogeneities. "
        "The automated workflow successfully generated all supporting files, as verified in the previous step. "
        "This process demonstrates an efficient, reproducible means of extracting and analyzing quantitative line profiles in scientific images."
    )
    doc.add_paragraph(results_para)

    # Insert Figure 1
    if os.path.exists(tiff_path):
        insert_figure(doc, tiff_path, "Fig. 1. $u_{eq}$ as a function of distance along the specified line segment.")

    # Save the document
    report_path = os.path.join(output_dir, f"{base_name}_simulation_report.docx")
    doc.save(report_path)
    print(f"Simulation report saved to: {report_path}")

if __name__ == "__main__":
    main()
