import os
import sys
import csv
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# File paths and names
output_dir = r"C:\Users\admin\Desktop\Python_proj\ALL_RESULT\GPT\T2S2\1.0\backup3"
base_filename = "Li_1.0"
gray_csv = os.path.join(output_dir, f"{base_filename}_gray_values.csv")
length_txt = os.path.join(output_dir, f"{base_filename}_length.txt")
ueq_csv = os.path.join(output_dir, f"{base_filename}_μeq_values.csv")
ueq_tiff = os.path.join(output_dir, f"{base_filename}_μeq_curve.tiff")
docx_path = os.path.join(output_dir, f"{base_filename}_simulation_report.docx")

# Helper: Read line segment length
def read_length_txt(length_txt):
    # Try encoding detection and handle possible BOM or odd characters
    try:
        with open(length_txt, 'r', encoding='utf-8') as f:
            line = f.readline()
    except UnicodeDecodeError:
        with open(length_txt, 'r', encoding='latin1') as f:
            line = f.readline()
    # Remove any non-numeric chars except dot and first minus
    import re
    match = re.search(r'([0-9]*\.?[0-9]+)', line)
    if match:
        return float(match.group(1))
    raise ValueError("Could not parse length from txt file.")

# Helper: Read μeq csv
def read_ueq_csv(ueq_csv):
    distances = []
    ueqs = []
    with open(ueq_csv, 'r', encoding='utf-8-sig') as f:
        reader = csv.reader(f)
        for i,row in enumerate(reader):
            if i == 0:
                continue
            if len(row) >= 2:
                try:
                    distances.append(float(row[0]))
                    ueqs.append(float(row[1]))
                except Exception:
                    continue
    return distances, ueqs

# Helper: Read grayscale values
def read_gray_csv(gray_csv):
    values = []
    with open(gray_csv, 'r', encoding='utf-8-sig') as f:
        reader = csv.reader(f)
        for i, row in enumerate(reader):
            if i == 0:
                continue
            if len(row) >= 3:
                try:
                    values.append(int(row[2]))
                except Exception:
                    continue
    return values

def add_heading_style(document):
    styles = document.styles
    try:
        h1 = styles['Heading 1']
        h1.font.size = Pt(16)
        h1.font.bold = True
        h1.font.name = 'Arial'
    except Exception:
        pass

def add_figure_caption(paragraph, text):
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def main():
    # Read values from previous outputs
    length_um = read_length_txt(length_txt)
    distances, ueqs = read_ueq_csv(ueq_csv)
    gray_values = read_gray_csv(gray_csv)

    # Create document
    document = Document()
    add_heading_style(document)

    # Title
    document.add_heading('Simulation Report: Grayscale Analysis and u_eq Mapping Along Line Segment in Li_1.0 Image', level=1)

    # Abstract (~200 words)
    abstract = (
        "This simulation report presents a detailed analysis of the grayscale intensity and equivalent voltage (u_eq) along a specific line segment in the micrograph 'Li_1.0.png'. "
        "Utilizing automated Python scripting, the grayscale values along the defined segment were extracted, mapped to u_eq using a linear transformation, and visualized. "
        "The analysis aims to bridge image-based measurements with physical property mapping, facilitating deeper insight into microstructural characteristics. "
        "The workflow comprises automated extraction of grayscale values, accurate length measurement based on image resolution, and computation of u_eq values using provided bounds. "
        "Results are presented through tabular data and plotted curves, providing a robust foundation for subsequent physical interpretation and simulation. "
        "The methodology and results outlined here demonstrate a reproducible approach for converting digital image data to quantitative physical metrics, potentially applicable to a range of materials characterization studies."
    )
    document.add_heading('Abstract', level=2)
    document.add_paragraph(abstract)

    # Introduction (~200 words)
    introduction = (
        "Digital image analysis is a pivotal tool in material science, allowing for quantitative extraction of microstructural features from high-resolution images. "
        "In this report, a line segment is defined between (152, 29) and (135, 92) pixels in a grayscale image of a lithium-based material. "
        "The primary objective is to map grayscale values along this segment to physical quantities—specifically, an equivalent voltage (u_eq)—using a user-defined linear relationship. "
        "Such mapping is vital for correlating image features to underlying physical phenomena, such as local electrical properties, phase distribution, or compositional gradients. "
        "The approach leverages Python-based automation, ensuring accuracy, repeatability, and ease of data handling. "
        "By integrating image processing with data transformation, this workflow establishes a template for future analyses where high-throughput, reproducible quantification is essential for research and development in advanced materials."
    )
    document.add_heading('Introduction', level=2)
    document.add_paragraph(introduction)

    # Methods (~200 words)
    methods = (
        "The workflow centers around a Python program (py1.py) that automates the extraction and analysis of grayscale pixel values along a user-specified line segment within the image. "
        "First, the program reads the image ('Li_1.0.png') and employs Bresenham’s line algorithm to enumerate all pixel coordinates between the defined start and end points. "
        "Grayscale intensity values (0–255) for each pixel along the segment are extracted and saved as a CSV array. "
        "The physical length of the line is computed using the supplied image resolution (1.08 μm/pixel), with the result saved to a text file for reference. "
        "To map grayscale intensity to equivalent voltage (u_eq), a linear transformation is applied: u_eq = u_min + (gray_value / 255) * u_max, where u_min=0 and u_max=65535. "
        "Distances from the start point and corresponding u_eq values are recorded in a CSV file. "
        "Finally, a plot of u_eq versus distance is generated and saved as a TIFF image (see Fig. 1). "
        "All outputs are named according to the original image filename and stored in the designated results directory."
    )
    document.add_heading('Methods', level=2)
    document.add_paragraph(methods)

    # Results (~200 words)
    results = (
        f"The automated analysis successfully extracted grayscale and u_eq data along the specified line segment. "
        f"The measured length of the segment is {length_um:.2f} μm, as recorded in the output text file. "
        "A total of {} points were analyzed, with grayscale values ranging from {} to {}. "
        "The transformation to u_eq yielded values between {:.0f} and {:.0f}. "
        "Figure 1 illustrates the variation of u_eq along the line, clearly revealing spatial variation in local image intensity and, by extension, the inferred physical property. "
        "This methodology enabled the rapid, reproducible quantification of spatially resolved data from digital images. "
        "The results, including tabulated grayscale and u_eq values, provide a robust framework for further analysis, comparison, and interpretation within the broader context of materials research."
    ).format(len(gray_values), min(gray_values), max(gray_values), min(ueqs), max(ueqs))
    document.add_heading('Results', level=2)
    document.add_paragraph(results)

    # Insert figure
    document.add_paragraph()
    if os.path.exists(ueq_tiff):
        p = document.add_paragraph()
        run = p.add_run()
        run.add_picture(ueq_tiff, width=Inches(4.5))
        add_figure_caption(document.add_paragraph(), "Fig. 1. u_eq as a function of distance along the line segment in 'Li_1.0.png'.")
    else:
        document.add_paragraph("[Figure file missing: {}]".format(ueq_tiff))

    # Save document
    document.save(docx_path)
    print(f"Word report generated: {docx_path}")

if __name__ == '__main__':
    main()
