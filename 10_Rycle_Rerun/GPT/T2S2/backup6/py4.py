import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import matplotlib.pyplot as plt
import pandas as pd

# File paths and names
output_dir = r"C:\Users\admin\Desktop\Python_proj\ALL_RESULT\GPT\T2S2\1.0\backup6"
image_filename = "Li_1.0"
gray_csv = os.path.join(output_dir, f"{image_filename}_line_grayscale.csv")
length_txt = os.path.join(output_dir, f"{image_filename}_line_length.txt")
dist_ueq_csv = os.path.join(output_dir, f"{image_filename}_distance_u_eq.csv")
tiff_img = os.path.join(output_dir, f"{image_filename}_u_eq_curve.tiff")
docx_path = os.path.join(output_dir, f"{image_filename}_simulation_report.docx")

def get_fig1_jpg_from_tiff(tiff_path, out_jpg_path):
    # Convert TIFF figure to JPG for embedding (Word may not show TIFF directly)
    img = plt.imread(tiff_path)
    plt.imsave(out_jpg_path, img)
    return out_jpg_path

def try_read_csv_with_encodings(csv_path, encodings=('utf-8', 'gbk', 'latin-1', 'utf-16')):
    """Try reading CSV with several encodings until one works."""
    for enc in encodings:
        try:
            df = pd.read_csv(csv_path, encoding=enc)
            return df
        except Exception:
            continue
    raise RuntimeError(f"Failed to read CSV {csv_path} with tried encodings.")

def main():
    # Read calculated data
    distances = []
    u_eqs = []
    try:
        df = try_read_csv_with_encodings(dist_ueq_csv)
        distances = df.iloc[:, 0].tolist()
        u_eqs = df.iloc[:, 1].tolist()
    except Exception as e:
        print("Error reading CSV:", e)
        return

    try:
        with open(length_txt, 'r', encoding='utf-8', errors='replace') as f:
            line_length = float(f.read().strip())
    except Exception as e:
        print("Error reading line length:", e)
        line_length = None

    # Prepare figure for Word (convert TIFF to JPG)
    fig1_jpg = os.path.join(output_dir, f"{image_filename}_u_eq_curve.jpg")
    if not os.path.exists(fig1_jpg):
        get_fig1_jpg_from_tiff(tiff_img, fig1_jpg)

    # Create the Word document
    doc = Document()
    doc.styles['Normal'].font.name = u'Times New Roman'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'Times New Roman')
    doc.styles['Normal'].font.size = Pt(11)

    # Title (centered, bold, 16pt)
    title = doc.add_heading("Simulation Report: Analysis of Grayscale and Equivalent Potential Along a Line Segment", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Abstract
    doc.add_heading("Abstract", level=1)
    abstract = (
        "This simulation report presents a quantitative analysis of grayscale intensities and their transformation "
        "into equivalent potential values (u_eq) along a defined line segment in a high-resolution image. By extracting "
        "grayscale values at each point along the segment and converting them using a linear scaling formula, the study "
        "visualizes the variation of potential across the sample. The results, including the calculated line segment "
        "length, grayscale distributions, and u_eq profiles, are documented for scientific interpretation. The analysis "
        "demonstrates the utility of automated image processing in extracting precise line profiles and translating "
        "them into physically meaningful quantities, as shown in Fig. 1. The workflow is fully automated, ensuring "
        "reproducibility and facilitating further research on image-based quantitative characterization."
    )
    doc.add_paragraph(abstract)

    # Introduction
    doc.add_heading("Introduction", level=1)
    introduction = (
        "The quantitative analysis of grayscale values along specific paths in scientific images is essential for "
        "understanding the spatial variation of material properties, chemical concentrations, or physical potentials. "
        "In materials science and related fields, digital image processing allows researchers to extract precise "
        "profiles from high-resolution images. This report focuses on a line segment analysis across an image (Li_1.0.png) "
        "using a resolution of 1.08 μm per pixel. The primary objective is to convert grayscale values (ranging from "
        "0 to 255) along the defined segment into equivalent potential (u_eq) values, based on the known minimum and "
        "maximum potential range. The resulting u_eq profile provides insight into the physical variations present in the "
        "sample and serves as a basis for further scientific interpretation. The automated approach adopted in this "
        "work ensures accuracy, repeatability, and rapid analysis, meeting the demands of modern quantitative research."
    )
    doc.add_paragraph(introduction)

    # Methods
    doc.add_heading("Methods", level=1)
    methods = (
        "The analysis was conducted using a Python script (py1.py) designed to automate the extraction of image "
        "intensities and subsequent calculations. The script reads a specified image, defines a line segment using "
        "given start and end pixel coordinates ((152, 29) to (135, 92)), and employs Bresenham's algorithm to map out "
        "all pixel points along the segment. Grayscale values at these points are extracted from the image and saved "
        "to a CSV file. The physical length of the segment is calculated using the pixel resolution (1.08 μm/pixel) and "
        "saved to a text file. Each grayscale value is then transformed into an equivalent potential (u_eq) using the "
        "formula: u_eq = u_min + (gray_value / 255) * (u_max - u_min), where u_max = 65535 and u_min = 0. The distance "
        "from the starting point and the corresponding u_eq values are both recorded and visualized by plotting u_eq "
        "against distance, with the resulting graph saved as a TIFF image. This workflow provides a reproducible "
        "pipeline for quantitative line profile analysis."
    )
    doc.add_paragraph(methods)

    # Results
    doc.add_heading("Results", level=1)
    results = (
        f"The analysis of the image segment produced a measured line length of {line_length:.3f} μm, as determined "
        "by the pixel resolution and the number of points along the segment. Grayscale values ranged from low to high "
        "intensities, reflecting material or potential heterogeneity along the path. The calculated u_eq values spanned "
        f"from {min(u_eqs):.0f} to {max(u_eqs):.0f}, corresponding directly to the grayscale intensity distribution. "
        "The resulting plot (see Fig. 1) depicts the u_eq profile as a function of distance, revealing distinct regions "
        "of potential variation. This quantitative profile can be used to correlate image features with underlying physical "
        "properties. The output files—CSV files containing grayscale and u_eq data, the length text file, and the plot image—"
        "are all saved in the specified output directory, ensuring traceability. Overall, the automated approach enabled "
        "rapid and accurate extraction of physically meaningful information from the original image data."
    )
    doc.add_paragraph(results)

    # Insert Figure 1 (u_eq vs distance)
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_picture(fig1_jpg, width=Inches(4.5))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Fig. 1. Plot of u_eq (equivalent potential) as a function of distance along the specified line segment.")

    # Save Word document
    doc.save(docx_path)
    print(f"Word report saved to: {docx_path}")

if __name__ == "__main__":
    main()
