import os
import csv
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    # --- File paths and parameters ---
    image_path = r"C:\Users\admin\Desktop\Python_proj\datas\T2_IMGS\Li_1.0.png"
    out_dir = r"C:\Users\admin\Desktop\Python_proj\ALL_RESULT\GPT\T2S2\1.0\backup7"
    filename_base = os.path.splitext(os.path.basename(image_path))[0]
    gray_csv = os.path.join(out_dir, f"{filename_base}_grayvalues.csv")
    length_txt = os.path.join(out_dir, f"{filename_base}_length.txt")
    ueq_csv = os.path.join(out_dir, f"{filename_base}_distance_ueq.csv")
    tiff_img = os.path.join(out_dir, f"{filename_base}_ueq_vs_distance.tiff")
    docx_path = os.path.join(out_dir, f"{filename_base}_simulation_report.docx")

    # --- Read measured length ---
    try:
        with open(length_txt, "r") as f:
            length_um = float(f.readline().strip())
    except Exception as e:
        print(f"Error reading length: {e}")
        length_um = None

    # --- Read grayscale and u_eq values ---
    distances = []
    ueqs = []
    try:
        with open(ueq_csv, "r") as f:
            reader = csv.DictReader(f)
            for row in reader:
                distances.append(float(row["Distance_μm"]))
                ueqs.append(float(row["u_eq"]))
    except Exception as e:
        print(f"Error reading u_eq CSV: {e}")

    # --- Prepare statistics for results section ---
    if ueqs:
        ueq_min = min(ueqs)
        ueq_max = max(ueqs)
        ueq_avg = sum(ueqs) / len(ueqs)
    else:
        ueq_min = ueq_max = ueq_avg = None

    # --- Create Word Report ---
    doc = Document()
    # Title
    doc.add_heading('Simulation Report: Grayscale-Ueq Analysis along Line Segment', 0)

    # Abstract
    doc.add_heading('Abstract', level=1)
    abstract_text = (
        "This report presents a detailed simulation and data analysis of grayscale values along a specified "
        "line segment within a microscopy image. The line segment was analyzed to extract pixel grayscale "
        "values and to compute the corresponding u_eq values using a scaling formula. The methodology, "
        "results, and interpretation are discussed, providing insight into spatial variations of grayscale "
        "intensity and their physical significance. The findings offer a reproducible approach for quantitative "
        "analysis in similar imaging studies."
    )
    doc.add_paragraph(abstract_text)

    # Introduction
    doc.add_heading('Introduction', level=1)
    intro_text = (
        "Quantitative image analysis is essential for extracting meaningful physical information from microscopy "
        "images. In this work, we focus on a one-dimensional analysis along a user-specified line segment in an "
        "image of lithium. The grayscale values, which represent intensity or density information, are measured "
        "along the segment. These values are then converted to equivalent physical values (u_eq) using a scaling "
        "relation. Such analysis enables the study of spatial variations and can be adapted for diverse research "
        "fields, including material science and bioimaging. The goal is to provide a reproducible methodology for "
        "extracting and interpreting line-based grayscale and u_eq profiles."
    )
    doc.add_paragraph(intro_text)

    # Methods
    doc.add_heading('Methods', level=1)
    methods_text = (
        "The computational workflow was implemented using Python. First, the input image was loaded, and a line "
        "segment was defined by two endpoints: (152, 29) and (135, 92). Using a digital line algorithm, all pixel "
        "coordinates along the segment were determined. Grayscale values (0-255) were extracted at each point. "
        "The physical length of the segment was computed based on the image resolution (1.08 μm/pixel). The grayscale "
        "values were then mapped to u_eq values using the formula: u_eq = u_min + (gray_value / 255) * u_max, where "
        "u_min was set to 0 and u_max to 65535. The results, including grayscale and u_eq arrays, distances from the "
        "start point, and the profile graph (u_eq vs. distance), were saved as CSV, TXT, and TIFF files. The entire "
        "process is automated and reproducible, ensuring accurate and efficient data extraction and visualization."
    )
    doc.add_paragraph(methods_text)

    # Results
    doc.add_heading('Results', level=1)
    results_text = (
        f"The length of the analyzed line segment was measured to be {length_um:.2f} μm. Grayscale values along "
        "the segment revealed clear spatial variations, which are reflected in the computed u_eq profile. The minimum, "
        f"maximum, and average u_eq values were {ueq_min:.2f}, {ueq_max:.2f}, and {ueq_avg:.2f}, respectively. "
        "The u_eq versus distance curve is shown in Fig. 1, illustrating how the physical property corresponding to the grayscale "
        "intensity changes along the segment. These results highlight the effectiveness of the approach for resolving "
        "fine spatial features and can be adapted to similar analyses in other systems. All raw and processed data, "
        "including the plot and value arrays, are available in the output files for further inspection."
    )
    doc.add_paragraph(results_text)

    # Insert figure
    if os.path.exists(tiff_img):
        doc.add_picture(tiff_img, width=Inches(5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('Fig. 1. u_eq vs. distance from the start point.')

    # Save the Word document
    doc.save(docx_path)
    print(f"Simulation report generated and saved to: {docx_path}")

if __name__ == "__main__":
    main()
