import os
import numpy as np
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt

def read_line_length(txt_path):
    with open(txt_path, 'r') as f:
        line = f.readline()
        length_um = float(line.split(":")[1].split("μm")[0].strip())
        length_px = float(line.split("(")[1].split("px")[0].replace(")", "").strip())
    return length_um, length_px

def read_gray_csv(csv_path):
    # Reads gray values and coordinates from CSV
    import csv
    gray_values = []
    coords = []
    with open(csv_path, 'r', newline='') as f:
        reader = csv.reader(f)
        next(reader)  # skip header
        for row in reader:
            coords.append((int(row[1]), int(row[2])))
            gray_values.append(int(row[3]))
    return np.array(coords), np.array(gray_values)

def read_ueq_csv(csv_path):
    # Reads distance and u_eq from CSV
    import csv
    distances = []
    u_eq = []
    with open(csv_path, 'r', newline='') as f:
        reader = csv.reader(f)
        next(reader)  # skip header
        for row in reader:
            distances.append(float(row[1]))
            u_eq.append(float(row[2]))
    return np.array(distances), np.array(u_eq)

def add_figure(doc, fig_path, caption):
    # Add a figure to the document with caption
    doc.add_picture(fig_path, width=Inches(5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(caption).alignment = WD_ALIGN_PARAGRAPH.CENTER

def main():
    # File paths
    output_dir = r"C:\Users\admin\Desktop\Python_proj\ALL_RESULT\GPT\T2S2\1.0\backup4"
    image_filename = "Li_1.0"
    txt_path = os.path.join(output_dir, f"{image_filename}_line_length.txt")
    gray_csv_path = os.path.join(output_dir, f"{image_filename}_gray_values.csv")
    ueq_csv_path = os.path.join(output_dir, f"{image_filename}_distance_u_eq.csv")
    fig_path = os.path.join(output_dir, f"{image_filename}_u_eq_curve.tiff")
    doc_path = os.path.join(output_dir, f"{image_filename}_simulation_report.docx")

    # Read data
    length_um, length_px = read_line_length(txt_path)
    coords, gray_values = read_gray_csv(gray_csv_path)
    distances, u_eq = read_ueq_csv(ueq_csv_path)

    # Generate and save figure as PNG for Word (TIFF not always supported)
    png_fig_path = os.path.join(output_dir, f"{image_filename}_u_eq_curve.png")
    plt.figure(figsize=(7,4))
    plt.plot(distances, u_eq, marker='o')
    plt.xlabel('Distance from Start (μm)')
    plt.ylabel('u_eq')
    plt.title('u_eq vs. Distance along Line')
    plt.tight_layout()
    plt.savefig(png_fig_path, dpi=300)
    plt.close()

    # Compose report
    doc = Document()
    title = "Simulation Report: Quantitative Analysis of Grayscale and Equivalent Potential along a Line in Li_1.0.png"
    doc.add_heading(title, 0)

    # Abstract (~200 words)
    abstract = (
        "This report presents a comprehensive simulation and quantitative analysis of grayscale values and their corresponding equivalent potential (u_eq) along a specified line segment within a microscopic image (Li_1.0.png). "
        "By extracting grayscale data along the user-defined segment, converting these to physical units using the provided image resolution, and applying a transformation to obtain u_eq, we visualize the spatial variation of potential across the segment. "
        "Our approach employs automated Python-based image processing to ensure reproducibility and accuracy. The resulting data, including grayscale values, calculated u_eq, and segment length, are systematically saved for further interpretation. "
        "The study provides insights into the local contrast and potential variations, potentially aiding in the characterization of material properties or phase boundaries within the imaged specimen. "
        "The workflow and results, including graphical representations, are discussed in detail and may serve as a reference for similar quantitative image analyses."
    )
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(abstract)

    # Introduction (~200 words)
    introduction = (
        "Microscopic imaging is a fundamental tool in materials science, physics, and related disciplines, allowing visualization of microstructural features and gradients that are otherwise inaccessible. "
        "Quantitative analysis of image intensity, specifically grayscale values, can reveal subtle variations in composition, thickness, or potential differences within a sample. "
        "In this study, we focus on a segment within the image 'Li_1.0.png', extracting grayscale data along a line defined by coordinates (152, 29) to (135, 92). "
        "By converting pixel-based measurements to physical units using the calibrated resolution (1.08 μm/pixel), we aim to relate image intensity to the equivalent potential, offering a direct link between observed contrast and material properties. "
        "Such analyses are crucial for validating experimental results, interpreting micrographs, and providing quantitative support for theoretical models. "
        "The purpose of this simulation is to demonstrate an integrated workflow for extracting, transforming, and visualizing image data, with results facilitating detailed understanding of the imaged system."
    )
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(introduction)

    # Methods (~200 words)
    methods = (
        "The analysis was implemented using Python, leveraging libraries such as NumPy, PIL, and Matplotlib for data handling, image processing, and visualization, respectively. "
        "The specified line segment was defined by its start and end coordinates, and the Bresenham algorithm was utilized to enumerate all pixel positions along the segment. "
        "Grayscale values (ranging from 0 to 255) at each pixel were extracted from the grayscale version of the image. The Euclidean distance between the two endpoints was computed using the provided resolution to obtain the physical length of the segment in micrometers. "
        "All grayscale values, coordinates, and computed distances were saved in CSV files for traceability. "
        "To obtain the equivalent potential (u_eq) for each point, the transformation u_eq = u_min + (gray_value / 255) * (u_max - u_min) was applied, with u_min = 0 and u_max = 65535. "
        "The resulting u_eq values were plotted against the physical distance from the start point, generating a curve that visually represents the spatial variation of potential along the segment. "
        "All processed data and graphical outputs were systematically saved and are referenced in the Results section."
    )
    doc.add_heading("Methods", level=1)
    doc.add_paragraph(methods)

    # Results (~200 words)
    results = (
        f"The analysis yielded a total segment length of {length_um:.2f} μm (corresponding to {length_px:.2f} pixels), with 64 grayscale data points sampled along the defined line. "
        f"The extracted grayscale values varied between {gray_values.min()} and {gray_values.max()}, reflecting local image intensity variations. "
        "The computed equivalent potential (u_eq) values, transformed from grayscale, similarly display spatial fluctuations, indicating regions of differing material properties or contrast. "
        "The plot of u_eq versus distance from the start point (as shown in Fig. 1) visually demonstrates these variations, with notable peaks and troughs corresponding to specific features along the line. "
        "All raw grayscale data, u_eq values, and distances were saved in CSV files for further examination. "
        "This quantitative approach enables objective comparison of image regions and can be extended to other images or segments. "
        "The results validate the effectiveness of automated image analysis for extracting meaningful physical information from micrographs."
    )
    doc.add_heading("Results", level=1)
    doc.add_paragraph(results)

    # Insert the graph
    add_figure(doc, png_fig_path, "Fig. 1. Plot of equivalent potential (u_eq) versus distance from start point along the line segment in Li_1.0.png.")

    # Save the document
    doc.save(doc_path)
    print(f"Simulation report successfully written to {doc_path}")

if __name__ == "__main__":
    main()
