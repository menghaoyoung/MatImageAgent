import os
import docx
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def find_result_files(result_dir):
    """
    Locate all *_GAP_map.png images and *_gap_analysis.csv files in the result directory.
    Returns (list_of_png_paths, list_of_csv_paths)
    """
    pngs = []
    csvs = []
    for fname in os.listdir(result_dir):
        if fname.endswith('_GAP_map.png'):
            pngs.append(os.path.join(result_dir, fname))
        elif fname.endswith('_gap_analysis.csv'):
            csvs.append(os.path.join(result_dir, fname))
    return pngs, csvs

def simulate_report_word(result_dir, output_word_path):
    """
    Generate a Word report summarizing simulation results, including images and sectioned text.
    """
    pngs, csvs = find_result_files(result_dir)
    doc = docx.Document()

    # Title
    doc.add_heading('Simulation Report on GAP Pixel Detection in Polymeric Images', 0)

    # Abstract
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        "This report presents the results of a computational analysis pipeline developed to identify and highlight GAP pixels "
        "within a set of polymeric material images. GAP pixels are defined by specific grayscale and spatial adjacency criteria, "
        "reflecting potential areas of interest for subsequent material characterization. The computational approach combines "
        "contrast enhancement using CLAHE, pixelwise analysis, and visualization of the results. The findings offer a robust, "
        "automatable method for supporting material research and quality assessment."
    )

    # Introduction
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph(
        "Polymeric materials often exhibit complex microstructures, where subtle grayscale variations in microscopy images can "
        "indicate important morphological features. Manual identification of such features is challenging and subjective. "
        "To address this, we developed a Python-based image analysis pipeline that automatically processes images, enhances contrast, "
        "and identifies 'GAP' pixels—defined by a grayscale value between 1 and 150 and adjacency to a contiguous segment of similar pixels. "
        "This automation supports high-throughput analysis and reduces human bias in materials research."
    )

    # Methods
    doc.add_heading('Methods', level=1)
    doc.add_paragraph(
        "The workflow begins by scanning a specified directory for images with filenames starting with 'Poly_'. Each image (PNG or JPG) is read "
        "using OpenCV and enhanced using Contrast Limited Adaptive Histogram Equalization (CLAHE) with a clipLimit of 3 and a tileGridSize of 10x10. "
        "The enhanced images are then converted to grayscale using the Pillow library. "
        "For each pixel, the grayscale value is extracted. A pixel is flagged as a GAP pixel if: (1) its grayscale value is between 1 and 150 inclusive, "
        "and (2) at least one of its four cardinal neighbors (up, down, left, right) has a contiguous run of 25 pixels, all satisfying the same grayscale condition. "
        "The results for each image are saved as a CSV file detailing pixel coordinates, grayscale value, and GAP flag, and a new PNG image visualizes GAP pixels in black (flag=1) "
        "and non-GAP pixels in white (flag=0)."
    )

    # Results
    doc.add_heading('Results', level=1)
    if not pngs:
        doc.add_paragraph(
            "No GAP map images were found in the output directory. Please ensure the analysis was successfully executed and output images are present."
        )
    else:
        doc.add_paragraph(
            "The pipeline processed the input images, producing a GAP analysis CSV and a visualization PNG for each. The new images below display detected GAP pixels in black and "
            "other regions in white, providing clear localization of areas with the specified grayscale and adjacency profile. This enables rapid inspection and quantitative assessment "
            "of potential microstructural gaps. The CSV files offer further quantitative detail for downstream analysis."
        )
        for img_path in pngs:
            doc.add_paragraph(os.path.basename(img_path))
            try:
                doc.add_picture(img_path, width=Inches(4.5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                doc.add_paragraph(f"Unable to insert image: {e}")

    doc.add_paragraph(
        "\nOverall, this automated approach efficiently highlights and quantifies regions of interest in polymeric images, "
        "facilitating objective and reproducible material analysis."
    )

    doc.save(output_word_path)
    print(f"Simulation report generated: {output_word_path}")

if __name__ == "__main__":
    # Output files path from previous steps
    result_dir = r"C:\Users\admin\Desktop\Python_proj\ALL_RESULT\GPT\T3\backup4"
    output_word = os.path.join(result_dir, "Simulation_Report_GAP_Analysis.docx")
    simulate_report_word(result_dir, output_word)
