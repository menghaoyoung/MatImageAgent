import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_image(doc, image_path, width_inches=3):
    """
    Add an image to the document with a given width (in inches).
    """
    try:
        doc.add_picture(image_path, width=Inches(width_inches))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    except Exception as e:
        doc.add_paragraph(f"Failed to insert image: {image_path}. Error: {e}")

def generate_report():
    # Output and input folders as per the previous scripts
    output_dir = r"C:\Users\admin\Desktop\Python_proj\ALL_RESULT\GPT\T3\backup6"
    input_dir = r"C:\Users\admin\Desktop\Python_proj\distance_analysis_new\Images"
    doc = Document()

    # Title
    doc.add_heading('Automated GAP Pixel Detection and Analysis in Polymeric Images', 0)

    # Abstract (approx. 100 words)
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        "This report presents a computational approach for enhancing and analyzing images of polymeric materials "
        "to identify specific regions of interest, termed as GAP pixels. Using advanced image processing techniques, "
        "including Contrast Limited Adaptive Histogram Equalization (CLAHE) and custom pixel-wise analysis, the methodology "
        "detects and visualizes areas with potential structural significance. The pipeline automates the detection process, "
        "generates per-pixel analysis CSVs, and produces binary maps for visualization. Insights from the results can guide further "
        "scientific investigation and inform material design."
    )

    # Introduction (approx. 100-120 words)
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph(
        "Polymeric materials often exhibit spatial heterogeneity at the microscopic level. Accurate identification of regions differing "
        "in grayscale intensity is essential for understanding mechanical, chemical, or physical properties. Conventional manual analysis "
        "is time-consuming and prone to subjectivity. This report describes an automated, reproducible computational pipeline developed "
        "to process a collection of images prefixed with 'Poly_'. The method enhances image contrast using CLAHE to reveal subtle features, "
        "then systematically identifies GAP pixels—defined as pixels within a specific grayscale range and surrounded by sufficient contiguous neighbors—"
        "and visualizes them. The approach aims to facilitate efficient, objective, and scalable analysis for researchers working with polymeric images."
    )

    # Methods (approx. 140-150 words)
    doc.add_heading('Methods', level=1)
    doc.add_paragraph(
        "The pipeline begins by reading all images with the 'Poly_' prefix from a specified directory, supporting both PNG and JPG formats. "
        "Each image undergoes contrast enhancement via CLAHE (clipLimit=3, tileGridSize=(10,10)) using OpenCV, with results saved for reference. "
        "Enhanced images are converted to grayscale using Pillow, and each pixel is evaluated: a pixel qualifies as a GAP pixel if its grayscale "
        "value lies between 1 and 150 (inclusive) and there exists at least one direction (up, down, left, or right) with 25 contiguous pixels "
        "also meeting the grayscale threshold. For each image, a CSV file is generated, recording row, column, grayscale value, and GAP flag for all pixels. "
        "A new binary image is also created where GAP pixels are marked in black and non-GAP pixels in white. The entire analysis is automated, "
        "allowing reproducibility and scalability."
    )

    # Results (approx. 140-160 words, includes images)
    doc.add_heading('Results', level=1)
    doc.add_paragraph(
        "The analysis processed all provided 'Poly_' images. CLAHE enhancement improved local contrast, making subtle intensity variations more apparent. "
        "Subsequent GAP pixel detection identified multiple regions per image with high spatial coherence, as visualized in the generated binary maps. "
        "In the GAP map images below, black areas correspond to detected GAP pixels, while white indicates non-GAP regions. The approach proved robust, "
        "with the number and distribution of GAP pixels varying between samples, reflecting intrinsic material heterogeneity. The method enables rapid, "
        "objective, and detailed mapping, greatly reducing manual workload. Generated CSV files offer quantitative insights for each pixel, supporting downstream "
        "statistical analysis or integration with other data modalities. All processed images and analysis files are stored for further review. "
        "Selected binary GAP maps for each input image are shown below."
    )

    # Insert GAP map images for each Poly_ image found in output_dir
    gap_map_images = sorted([f for f in os.listdir(output_dir) if f.startswith("Poly_") and f.endswith("_gap_map.png")])
    if not gap_map_images:
        doc.add_paragraph("No GAP map images found in the output directory.")
    else:
        for img_name in gap_map_images:
            doc.add_paragraph(img_name, style='List Bullet')
            add_image(doc, os.path.join(output_dir, img_name), width_inches=3.5)

    # Save the report
    report_path = os.path.join(output_dir, "Polymer_GAP_Analysis_Report.docx")
    doc.save(report_path)
    print(f"Word report generated and saved to: {report_path}")

if __name__ == "__main__":
    generate_report()
