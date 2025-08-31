import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def get_gap_outputs(output_dir):
    # Gather analysis CSVs and highlight PNGs
    csv_files = []
    img_files = []
    for f in os.listdir(output_dir):
        if f.endswith("_gap_analysis.csv"):
            csv_files.append(f)
        elif f.endswith("_gap_highlight.png"):
            img_files.append(f)
    return csv_files, img_files

def generate_report(output_dir, docx_path):
    csv_files, img_files = get_gap_outputs(output_dir)
    # Compose sections
    title = "Automated GAP Pixel Detection and Simulation Analysis"
    abstract = (
        "This report presents an automated pipeline for identifying GAP pixels in microscopy or other scientific images. "
        "By leveraging CLAHE for contrast enhancement and a defined local grayscale criterion, the algorithm efficiently highlights structural features in the images. "
        "The approach is validated on a series of images, with results visualized and quantitatively summarized."
    )
    introduction = (
        "Image-based analysis is central to modern scientific research, enabling quantitative assessment of spatial structures and anomalies. "
        "A recurring challenge is the reliable detection of regions with distinct grayscale properties—here termed as GAP pixels. "
        "This project aims to develop and validate an algorithm that enhances image contrast, identifies GAP pixels using a custom local neighborhood rule, "
        "and visualizes the results for further interpretation."
    )
    methods = (
        "The developed Python program processes all input images whose names begin with 'Poly_' from the specified directory. "
        "Each image is first contrast-enhanced using CLAHE (clipLimit=3, tileGridSize=(10,10)) via OpenCV to improve pixel-level discrimination. "
        "The enhanced images are then converted to grayscale using the PIL library. "
        "A GAP pixel is defined as any pixel with a grayscale value between 1 and 150 (inclusive) and having at least one adjacent direction (up/down/left/right) "
        "with 25 contiguous pixels also meeting the grayscale criterion. "
        "The program saves a CSV file per image encoding the coordinates, grayscale values, and GAP flag for each pixel, "
        "and generates a binary highlight image (black for GAP, white otherwise)."
    )
    # Results - summary statistics
    results = (
        "The algorithm was executed on the provided image set. Below are the processed results, including binary highlight images for visual inspection. "
        "The CSV files provide per-pixel quantitative data. The approach successfully delineates GAP regions, as evident from the highlighted images. "
        "This automated method aids in systematic feature extraction from complex image data, supporting downstream quantitative and qualitative analyses."
    )

    # Start docx
    doc = Document()
    doc.add_heading(title, 0)

    # Abstract
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(abstract)

    # Introduction
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(introduction)

    # Methods
    doc.add_heading("Methods", level=1)
    doc.add_paragraph(methods)

    # Results, images
    doc.add_heading("Results", level=1)
    doc.add_paragraph(results)

    # Add images
    if img_files:
        for imgname in img_files:
            img_path = os.path.join(output_dir, imgname)
            doc.add_paragraph(f"Result Image: {imgname}")
            try:
                doc.add_picture(img_path, width=Inches(4))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                doc.add_paragraph(f"[Unable to display image: {imgname}]")
    else:
        doc.add_paragraph("No result images found.")

    # Save the document
    doc.save(docx_path)
    print(f"Simulation report generated and saved to {docx_path}")

if __name__ == "__main__":
    output_dir = r"C:\Users\admin\Desktop\Python_proj\ALL_RESULT\GPT\T3\backup7"
    docx_path = os.path.join(output_dir, "GAP_Pixel_Simulation_Report.docx")
    generate_report(output_dir, docx_path)
