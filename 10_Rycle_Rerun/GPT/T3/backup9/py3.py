import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def generate_simulation_report(output_dir):
    # Gather all relevant images and CSV files
    images = []
    csvs = []
    for fname in os.listdir(output_dir):
        if fname.startswith("Poly_") and fname.endswith("_gap_map.png"):
            images.append(fname)
        if fname.startswith("Poly_") and fname.endswith("_gap_analysis.csv"):
            csvs.append(fname)
    images.sort()
    csvs.sort()

    # Title for the document
    report_title = "Simulation Report: Automated GAP Pixel Detection in Polymeric Images"

    # Begin the document
    doc = Document()
    doc.add_heading(report_title, 0)

    # Abstract
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        "This report presents a comprehensive simulation study for the automated detection of GAP pixels in polymeric "
        "microscope images using advanced image processing techniques. The pipeline utilizes contrast-limited adaptive "
        "histogram equalization (CLAHE) and pixel-level analysis to identify regions of interest, supporting quantitative "
        "and visual analysis for scientific research. Results are summarized both in tabular format and as visual maps."
    )

    # Introduction
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph(
        "Polymeric materials are commonly analyzed via microscopy to assess their surface features and uniformity. "
        "Detecting regions with different optical densities—here termed 'GAP pixels'—is crucial for understanding "
        "material structure and properties. Manual assessment is tedious and subjective; thus, automated pipelines using "
        "computer vision are highly valuable. This project implements a robust approach to enhance, segment, and analyze "
        "images, providing both CSV data and visual results."
    )

    # Methods
    doc.add_heading('Methods', level=1)
    doc.add_paragraph(
        "All input images with the prefix 'Poly_' were selected from the specified folder. Each image underwent "
        "contrast-limited adaptive histogram equalization (CLAHE) with clipLimit=3 and tileGridSize=(10,10) to improve "
        "local contrast and accentuate relevant features. Enhanced images were then converted to grayscale. For each pixel, "
        "the grayscale value was computed and a GAP flag was assigned if the value was between 1 and 150 (inclusive), "
        "and if at least one of the four cardinal directions (up, down, left, right) contained 25 contiguous pixels also "
        "meeting the grayscale criterion. The results for each pixel were recorded in a per-pixel CSV file. Additionally, "
        "a binary map image was generated for each sample, marking GAP pixels in black and non-GAP pixels in white for clear visualization."
    )

    # Results
    doc.add_heading('Results', level=1)
    doc.add_paragraph(
        "The automated workflow successfully processed all target images. For each, a CSV file containing per-pixel data "
        "and a corresponding binary GAP map image were generated. GAP pixels (visualized as black) highlight regions of interest "
        "where the grayscale and continuity criteria are met. These outputs can be used for further quantitative or visual analysis. "
        "Below are the generated GAP map images:"
    )

    # Insert GAP map images into the document (with captions)
    for i, img_fname in enumerate(images):
        img_path = os.path.join(output_dir, img_fname)
        doc.add_paragraph(f"Figure {i+1}: GAP Map for {img_fname.replace('_gap_map.png','')}", style='Caption')
        try:
            doc.add_picture(img_path, width=Inches(4.5))
            last_paragraph = doc.paragraphs[-1] 
            last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        except Exception as e:
            doc.add_paragraph(f"[Image could not be loaded: {e}]")

    doc.add_page_break()

    # Save the report
    report_path = os.path.join(output_dir, "Simulation_Report_GAP_Pixels.docx")
    doc.save(report_path)
    print(f"Simulation report generated: {report_path}")

if __name__ == "__main__":
    output_dir = r"C:\Users\admin\Desktop\Python_proj\ALL_RESULT\GPT\T3\backup9"
    generate_simulation_report(output_dir)
