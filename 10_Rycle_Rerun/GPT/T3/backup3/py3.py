import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def get_gap_images_and_csvs(output_directory):
    """
    Get all gap highlight images and csv files in the output directory, grouped by image basename.
    """
    gap_images = []
    csv_files = []
    for fname in os.listdir(output_directory):
        if fname.endswith('_gap_highlight.png'):
            gap_images.append(fname)
        elif fname.endswith('_gap_analysis.csv'):
            csv_files.append(fname)
    # Sort by base name for pairing
    gap_images.sort()
    csv_files.sort()
    return gap_images, csv_files

def add_image_with_caption(document, image_path, caption):
    try:
        document.add_picture(image_path, width=Inches(3))
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_paragraph(caption).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    except Exception as e:
        document.add_paragraph(f"[Image {os.path.basename(image_path)} could not be displayed: {e}]")
        print(f"Error displaying image {image_path}: {e}")

def generate_simulation_report(output_directory, output_doc_path):
    gap_images, csv_files = get_gap_images_and_csvs(output_directory)
    if not gap_images:
        print("No GAP highlight images found in output directory. Report not generated.")
        return
    if not csv_files:
        print("No CSV files found in output directory. Report not generated.")
        return
    
    document = Document()

    # Title
    document.add_heading("Simulation Report: GAP Pixel Analysis in Polymeric Images", 0)

    # Abstract
    document.add_heading("Abstract", level=1)
    document.add_paragraph(
        "This report presents a detailed analysis of GAP pixel distribution in polymeric material images. "
        "We applied image enhancement and pixel analysis techniques to automatically identify and quantify GAP regions. "
        "The approach combines computer vision (CLAHE, grayscale conversion) and custom pixel-wise inspection based on contextual grayscale thresholds. "
        "The results, presented as per-pixel CSV data and visual highlight images, facilitate systematic evaluation and comparison of structural features in polymer images."
    )

    # Introduction
    document.add_heading("Introduction", level=1)
    document.add_paragraph(
        "Polymeric materials often exhibit microstructural heterogeneities, such as gaps or pores, "
        "which can significantly influence their mechanical and functional properties. "
        "Automated detection and quantification of such GAP regions in imaging data is critical for quality control and research. "
        "This work demonstrates a pipeline for GAP pixel extraction from a collection of microscopy or digital images, "
        "enabling objective, reproducible analysis of gap distributions."
    )

    # Methods
    document.add_heading("Methods", level=1)
    document.add_paragraph(
        "All images prefixed with 'Poly_' were collected from the designated input folder and processed using a systematic pipeline. "
        "First, each image underwent Contrast Limited Adaptive Histogram Equalization (CLAHE) enhancement using OpenCV with clipLimit=3 and tileGridSize=(10,10), "
        "which improved local contrast and feature visibility. "
        "The enhanced images were converted to grayscale using the Pillow library, and pixel values were extracted for analysis.\n\n"
        "A custom GAP detection algorithm was implemented: a pixel was labeled as a GAP pixel if its grayscale value was within [1, 150] and "
        "if at least one direction (up, down, left, or right) contained 25 contiguous pixels (including self) within this grayscale range. "
        "For each image, two outputs were generated: (1) a per-pixel CSV file detailing coordinates, grayscale value, and GAP flag; "
        "and (2) a new PNG image where GAP pixels were marked in black and non-GAP pixels in white. "
        "These outputs provide both quantitative and visual representations of the GAP regions."
    )

    # Results
    document.add_heading("Results", level=1)
    document.add_paragraph(
        f"A total of {len(gap_images)} images were processed and analyzed. "
        "Below, each generated highlight image is shown, where black pixels indicate detected GAP regions. "
        "The corresponding CSV files contain pixel-level data, which can be further used for statistical analysis or machine learning purposes."
    )

    for img_fname in gap_images:
        img_path = os.path.join(output_directory, img_fname)
        base_name = img_fname.replace('_gap_highlight.png', '')
        caption = f"GAP Highlight Result: {base_name}"
        add_image_with_caption(document, img_path, caption)
        document.add_paragraph(f"CSV data for this image: {base_name}_gap_analysis.csv")

    document.add_page_break()
    document.add_paragraph("End of report.")

    document.save(output_doc_path)
    print(f"Word report generated: {output_doc_path}")

if __name__ == "__main__":
    output_directory = r"C:\Users\admin\Desktop\Python_proj\ALL_RESULT\GPT\T3\backup2"
    output_doc_path = os.path.join(output_directory, "Simulation_Report_GAP_Pixel_Analysis.docx")
    generate_simulation_report(output_directory, output_doc_path)
