import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_image(doc, img_path, width_inch=4.0):
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(width_inch))
    else:
        doc.add_paragraph(f"Image not found: {img_path}")

def main():
    # Paths used in previous tasks
    input_dir = r"C:\Users\admin\Desktop\Python_proj\distance_analysis_new\Images"
    output_dir = r"C:\Users\admin\Desktop\Python_proj\ALL_RESULT\GPT\T3\backup8"

    # Collect processed image names
    images = [f for f in os.listdir(input_dir) if f.startswith("Poly_") and f.lower().endswith(('.png', '.jpg', '.jpeg'))]
    images.sort()
    
    # Prepare mapping for gap map images
    gap_map_imgs = []
    for img in images:
        base_name = os.path.splitext(img)[0]
        gap_map_img_path = os.path.join(output_dir, f"{base_name}_gap_map.png")
        if os.path.exists(gap_map_img_path):
            gap_map_imgs.append(gap_map_img_path)

    # Start Word document
    doc = Document()
    doc.add_heading('Simulation Report: Automated GAP Pixel Detection in Polymeric Images', 0)

    # Abstract
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        "This simulation report summarizes the development and evaluation of an automated image processing pipeline for detecting GAP pixels within a set of polymeric material images. "
        "The approach employs contrast enhancement, grayscale conversion, and a custom-defined GAP pixel condition to identify regions of interest. "
        "The results provide a foundation for robust quantitative analysis in polymeric materials research and demonstrate the effectiveness of the proposed computational workflow."
    )

    # Introduction
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph(
        "Accurate identification of specific features in polymeric material images is crucial for quantitative material analysis. "
        "In this simulation, a fully automated method was developed to process a series of images with a standardized naming convention (prefix 'Poly_'). "
        "The goal was to enhance image contrast, extract grayscale values, and identify GAP pixels—regions that may correspond to structural or compositional heterogeneities. "
        "Automating this process reduces subjectivity and improves reproducibility compared to manual assessment."
    )

    # Methods
    doc.add_heading('Methods', level=1)
    doc.add_paragraph(
        "The processing pipeline begins by reading all images in the specified input directory with filenames starting with 'Poly_'. "
        "For each image, the following steps are applied:\n"
        "- **CLAHE Enhancement:** Using OpenCV's CLAHE (Contrast Limited Adaptive Histogram Equalization) with a clip limit of 3 and a tile grid size of (10, 10), local contrast is improved, making subtle features more discernible.\n"
        "- **Grayscale Conversion:** Enhanced images are converted to grayscale to simplify pixel intensity analysis.\n"
        "- **GAP Pixel Identification:** Each pixel is labeled as a GAP pixel if its grayscale value is between 1 and 150 (inclusive), and in at least one of the four cardinal directions (up, down, left, right), there are 25 contiguous pixels also within this grayscale range. "
        "This is implemented through pixel-wise iteration and directional checks using NumPy arrays for efficiency.\n"
        "- **Results Saving:** For each image, a CSV file containing per-pixel data (coordinates, grayscale value, GAP flag) is generated. "
        "Additionally, a new binary PNG image is produced, with GAP pixels shown in black and non-GAP pixels in white, facilitating visual inspection.\n"
        "All outputs are saved to a designated directory for downstream analysis."
    )

    # Results
    doc.add_heading('Results', level=1)
    doc.add_paragraph(
        "The automated pipeline successfully processed all provided images. The enhanced images and GAP maps allow for rapid assessment of GAP pixel distribution across the dataset. "
        "Tables of pixel-wise data (not shown here) are available in CSV format for further quantitative analysis. "
        "Below, the generated GAP map images are presented for each processed sample, where black regions correspond to detected GAP pixels."
    )
    for img_path in gap_map_imgs:
        doc.add_paragraph(os.path.basename(img_path), style='Intense Quote').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        add_image(doc, img_path, width_inch=4.0)
        doc.add_paragraph("")  # Add spacing

    # Save the document
    report_path = os.path.join(output_dir, "Simulation_Report_GAP_Pixel_Analysis.docx")
    doc.save(report_path)
    print(f"Report generated and saved to: {report_path}")

if __name__ == "__main__":
    main()
