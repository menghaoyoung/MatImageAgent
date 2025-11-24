import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_simulation_report(results_dir, output_docx_path):
    """
    Generate a detailed 500-word simulation report Word document based on the
    PNG gap-highlighted images in results_dir.
    The document is structured into four sections:
    Abstract, Introduction, Methods, Results (with inserted images).
    """

    # Create new Word document
    doc = Document()

    # Title
    doc_title = "Simulation Report on Pixel GAP Analysis in Grayscale Images"
    title_paragraph = doc.add_heading(doc_title, level=0)
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Abstract (approx 100-120 words)
    abstract_text = (
        "This report presents a detailed simulation analysis of grayscale images "
        "performing per-pixel GAP detection. GAP detection identifies pixels that meet "
        "specific grayscale intensity criteria and adjacency conditions, highlighting "
        "regions of interest within the images. The analysis involved generating "
        "CSV data files and corresponding gap-highlighted images indicating detected "
        "GAP pixels. The resultant outputs facilitate further understanding "
        "of the spatial distribution and characteristics of GAP regions across "
        "various sample images. This report outlines the background, methodology, "
        "and key findings derived from the processed data."
    )
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(abstract_text)

    # Introduction (approx 100-120 words)
    introduction_text = (
        "The purpose of this simulation is to analyze grayscale images at the pixel level "
        "to detect GAP regions, defined by pixels having a grayscale value between 1 and 155 "
        "inclusive, and adjacency to regions of at least 25 contiguous pixels meeting the same "
        "grayscale conditions. This process is crucial for applications requiring the identification "
        "of critical features or anomalous areas in images, such as materials analysis, medical imaging, "
        "and pattern recognition. By automating the detection and visualization of these areas, this "
        "analysis supports enhanced accuracy and efficiency in interpreting image data."
    )
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(introduction_text)

    # Methods (approx 120-150 words)
    methods_text = (
        "The methodology employed involves processing images located in the designated "
        "'Images' directory, focusing on files prefixed with 'Poly_'. Each image "
        "underwent grayscale reading followed by enhancement using CLAHE (Contrast Limited "
        "Adaptive Histogram Equalization) to improve local contrast. The GAP detection algorithm "
        "then identified pixels meeting the grayscale intensity criteria coupled with adjacency checks "
        "for contiguous pixel groups of size 25 or greater within specified thresholds. "
        "The output includes two key files per image: a CSV file detailing per-pixel data "
        "(coordinates, grayscale value, and GAP flag), and a new PNG image highlighting GAP "
        "pixels in black and non-GAP pixels in white. This structured approach ensures "
        "systematic and reproducible analysis across all input samples."
    )
    doc.add_heading("Methods", level=1)
    doc.add_paragraph(methods_text)

    # Results (approx 120-150 words + images)
    results_intro = (
        "The results present the effectiveness of the GAP detection algorithm through visual and "
        "tabular representation. The gap-highlighted PNG images clearly distinguish GAP pixels, "
        "allowing intuitive assessment of detected regions. The accompanying CSV files provide "
        "granular per-pixel detail for in-depth analysis or further processing. "
        "Below are the key processed images generated during the simulation."
    )
    doc.add_heading("Results", level=1)
    doc.add_paragraph(results_intro)

    # Insert all gap-highlighted PNG images in results_dir starting with 'Poly_' and ending with '_gap.png'
    images_inserted = 0
    for filename in sorted(os.listdir(results_dir)):
        if filename.startswith("Poly_") and filename.endswith("_gap.png"):
            img_path = os.path.join(results_dir, filename)
            try:
                doc.add_paragraph(f"Figure: {filename}")
                doc.add_picture(img_path, width=Inches(5))  # Resize to 5 inches width
                images_inserted += 1
            except Exception as e:
                doc.add_paragraph(f"Could not insert image {filename}: {e}")

    if images_inserted == 0:
        doc.add_paragraph("No gap-highlighted PNG images found in the results directory.")

    # Save the document
    doc.save(output_docx_path)
    print(f"Simulation report generated and saved at: {output_docx_path}")

if __name__ == "__main__":
    # Paths must match the output from previous program or fallback
    results_directory = "C:/MatImageAgent/MatImageAgent/Results"
    if not os.path.exists(results_directory) or results_directory.startswith("/share/"):
        results_directory = os.path.join(os.getcwd(), "Results")

    output_report_path = os.path.join(results_directory, "Simulation_Report.docx")

    create_simulation_report(results_directory, output_report_path)
