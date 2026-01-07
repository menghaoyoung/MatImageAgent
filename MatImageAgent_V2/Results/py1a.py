import os
import csv
import cv2
import numpy as np
from collections import defaultdict
import time
import matplotlib.pyplot as plt
from datetime import datetime

# Try to import docx, with graceful fallback
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    print("Warning: python-docx not installed. Will create a text report instead.")
    HAS_DOCX = False

def read_voice_transcript():
    """Read voice transcript if available and merge with task description"""
    merged_description = "Task: Generate per-pixel CSV with coordinates, grayscale values, and GAP flags. Create new images highlighting GAP points."
    
    try:
        if os.path.exists("./Voice_demo.txt"):
            with open("./Voice_demo.txt", "r") as f:
                voice_content = f.read().strip()
                merged_description += f"\nVoice transcript: {voice_content}"
                print("Voice transcript merged with task description.")
        else:
            print("No voice transcript found. Proceeding with original task description.")
    except Exception as e:
        print(f"Error reading voice transcript: {e}")
    
    print("Merged task summary:")
    print(merged_description)
    return merged_description

def generate_report():
    """Generate a Word document report based on the analysis results"""
    results_dir = "./Results"
    if not os.path.exists(results_dir):
        print(f"Results directory {results_dir} not found.")
        return False
    
    # Get all highlighted images
    image_files = [f for f in os.listdir(results_dir) if f.endswith('_gap_highlighted.png')]
    if not image_files:
        print("No highlighted images found in Results directory.")
        return False
    
    # Get all CSV files
    csv_files = [f for f in os.listdir(results_dir) if f.endswith('_gap_analysis.csv')]
    if not csv_files:
        print("No CSV analysis files found in Results directory.")
        return False
    
    # Calculate statistics from CSV files
    stats = {}
    for csv_file in csv_files:
        base_name = csv_file.replace('_gap_analysis.csv', '')
        stats[base_name] = analyze_csv(os.path.join(results_dir, csv_file))
    
    # Create the report
    if HAS_DOCX:
        create_word_report(results_dir, image_files, stats)
    else:
        create_text_report(results_dir, image_files, stats)
    
    return True

def analyze_csv(csv_path):
    """Analyze CSV file to extract statistics"""
    total_pixels = 0
    gap_pixels = 0
    grayscale_values = []
    
    try:
        with open(csv_path, 'r', newline='') as f:
            reader = csv.reader(f)
            next(reader)  # Skip header
            
            for row in reader:
                if len(row) >= 4:
                    total_pixels += 1
                    grayscale = int(row[2])
                    gap_flag = int(row[3])
                    
                    if gap_flag == 1:
                        gap_pixels += 1
                    
                    grayscale_values.append(grayscale)
    except Exception as e:
        print(f"Error analyzing CSV {csv_path}: {e}")
        return {"total_pixels": 0, "gap_pixels": 0, "gap_percentage": 0, "avg_grayscale": 0}
    
    avg_grayscale = sum(grayscale_values) / len(grayscale_values) if grayscale_values else 0
    gap_percentage = (gap_pixels / total_pixels * 100) if total_pixels > 0 else 0
    
    return {
        "total_pixels": total_pixels,
        "gap_pixels": gap_pixels,
        "gap_percentage": gap_percentage,
        "avg_grayscale": avg_grayscale
    }

def create_word_report(results_dir, image_files, stats):
    """Create a Word document report with images and analysis"""
    try:
        doc = Document()
        
        # Title
        title = doc.add_heading('Pixel GAP Analysis Simulation Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        date_paragraph = doc.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_paragraph.add_run(datetime.now().strftime("%B %d, %Y")).italic = True
        
        # Abstract
        doc.add_heading('Abstract', 1)
        abstract = doc.add_paragraph()
        abstract.add_run(
            "This report presents a comprehensive analysis of pixel-level GAP conditions in a series of images. "
            "The analysis identifies pixels meeting specific grayscale value conditions and adjacency patterns, "
            "which are critical for understanding structural integrity in the analyzed materials. "
            "The methodology employs image processing techniques to identify and flag pixels that meet the GAP criteria, "
            "defined as grayscale values between 1-155 with at least one adjacent direction containing 25 contiguous "
            "pixels meeting the same condition. Results are presented as statistical summaries and visual representations "
            "highlighting the spatial distribution of GAP conditions across multiple sample images."
        )
        
        # Introduction
        doc.add_heading('Introduction', 1)
        intro = doc.add_paragraph()
        intro.add_run(
            "The purpose of this analysis is to identify and characterize pixels meeting specific GAP conditions "
            "within a set of images. GAP conditions are defined by two primary criteria: (1) pixels with grayscale "
            "values between 1 and 155 (inclusive), and (2) pixels that have at least one adjacent direction "
            "(up, down, left, or right) containing 25 contiguous pixels that also meet the grayscale criterion. "
            "This analysis is particularly relevant for identifying potential structural weaknesses or anomalies "
            "in materials represented by the analyzed images. By systematically identifying and mapping these GAP "
            "conditions, we can better understand the distribution and patterns of these features across different samples."
        )
        
        # Methods
        doc.add_heading('Methods', 1)
        methods = doc.add_paragraph()
        methods.add_run(
            "The analysis was conducted using Python with OpenCV for image processing. The methodology consisted of "
            "the following steps:\n\n"
            "1. Image Enhancement: Each input image was converted to grayscale (if not already) and enhanced using "
            "Contrast Limited Adaptive Histogram Equalization (CLAHE) to improve feature visibility.\n\n"
            "2. Pixel Analysis: For each pixel in the enhanced image, the grayscale value was extracted and checked "
            "against the first GAP condition (values between 1-155).\n\n"
            "3. Adjacency Analysis: For pixels meeting the grayscale condition, an adjacency check was performed in "
            "four directions (up, down, left, right) to identify if any direction contained 25 contiguous pixels also "
            "meeting the grayscale condition.\n\n"
            "4. Data Recording: The results were recorded in CSV files containing the coordinates, grayscale values, "
            "and GAP flags (1 for meeting conditions, 0 otherwise) for each pixel.\n\n"
            "5. Visualization: New images were generated highlighting pixels meeting GAP conditions in black against "
            "a white background, providing a clear visual representation of the spatial distribution of these features."
        )
        
        # Results
        doc.add_heading('Results', 1)
        results_intro = doc.add_paragraph()
        results_intro.add_run(
            "The analysis was performed on multiple images, with the following results. For each image, "
            "we present the highlighted GAP visualization and key statistics including the percentage of "
            "pixels meeting GAP conditions and average grayscale values."
        )
        
        # Add results for each image
        for image_file in sorted(image_files):
            base_name = image_file.replace('_gap_highlighted.png', '')
            
            doc.add_heading(f"Sample: {base_name}", 2)
            
            # Add statistics
            if base_name in stats:
                stat_data = stats[base_name]
                stat_para = doc.add_paragraph()
                stat_para.add_run(f"Total pixels analyzed: {stat_data['total_pixels']:,}\n")
                stat_para.add_run(f"Pixels meeting GAP conditions: {stat_data['gap_pixels']:,}\n")
                stat_para.add_run(f"Percentage of GAP pixels: {stat_data['gap_percentage']:.2f}%\n")
                stat_para.add_run(f"Average grayscale value: {stat_data['avg_grayscale']:.2f}\n")
            
            # Add image
            image_path = os.path.join(results_dir, image_file)
            if os.path.exists(image_path):
                doc.add_picture(image_path, width=Inches(6))
                image_caption = doc.add_paragraph("Figure: Highlighted GAP analysis results")
                image_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                doc.add_paragraph(f"[Image file not found: {image_file}]")
        
        # Conclusion
        doc.add_heading('Conclusion', 1)
        conclusion = doc.add_paragraph()
        conclusion.add_run(
            "This analysis successfully identified and characterized pixels meeting GAP conditions across multiple "
            "sample images. The results provide valuable insights into the spatial distribution and prevalence of "
            "these features, which can inform further investigation into structural properties and potential anomalies. "
            "The methodology developed for this analysis can be applied to additional samples to build a more "
            "comprehensive understanding of GAP conditions in similar materials."
        )
        
        # Save the document
        report_path = os.path.join(results_dir, "GAP_Analysis_Report.docx")
        doc.save(report_path)
        print(f"Word report generated successfully: {report_path}")
        
    except Exception as e:
        print(f"Error creating Word report: {e}")
        return False
    
    return True

def create_text_report(results_dir, image_files, stats):
    """Create a text-based report as fallback if docx is not available"""
    try:
        report_path = os.path.join(results_dir, "GAP_Analysis_Report.txt")
        
        with open(report_path, 'w') as f:
            f.write("PIXEL GAP ANALYSIS SIMULATION REPORT\n")
            f.write("=" * 40 + "\n\n")
            f.write(f"Date: {datetime.now().strftime('%B %d, %Y')}\n\n")
            
            # Abstract
            f.write("ABSTRACT\n")
            f.write("-" * 8 + "\n")
            f.write(
                "This report presents a comprehensive analysis of pixel-level GAP conditions in a series of images. "
                "The analysis identifies pixels meeting specific grayscale value conditions and adjacency patterns, "
                "which are critical for understanding structural integrity in the analyzed materials. "
                "The methodology employs image processing techniques to identify and flag pixels that meet the GAP criteria, "
                "defined as grayscale values between 1-155 with at least one adjacent direction containing 25 contiguous "
                "pixels meeting the same condition. Results are presented as statistical summaries and visual representations "
                "highlighting the spatial distribution of GAP conditions across multiple sample images.\n\n"
            )
            
            # Introduction
            f.write("INTRODUCTION\n")
            f.write("-" * 12 + "\n")
            f.write(
                "The purpose of this analysis is to identify and characterize pixels meeting specific GAP conditions "
                "within a set of images. GAP conditions are defined by two primary criteria: (1) pixels with grayscale "
                "values between 1 and 155 (inclusive), and (2) pixels that have at least one adjacent direction "
                "(up, down, left, or right) containing 25 contiguous pixels that also meet the grayscale criterion. "
                "This analysis is particularly relevant for identifying potential structural weaknesses or anomalies "
                "in materials represented by the analyzed images. By systematically identifying and mapping these GAP "
                "conditions, we can better understand the distribution and patterns of these features across different samples.\n\n"
            )
            
            # Methods
            f.write("METHODS\n")
            f.write("-" * 8 + "\n")
            f.write(
                "The analysis was conducted using Python with OpenCV for image processing. The methodology consisted of "
                "the following steps:\n\n"
                "1. Image Enhancement: Each input image was converted to grayscale (if not already) and enhanced using "
                "Contrast Limited Adaptive Histogram Equalization (CLAHE) to improve feature visibility.\n\n"
                "2. Pixel Analysis: For each pixel in the enhanced image, the grayscale value was extracted and checked "
                "against the first GAP condition (values between 1-155).\n\n"
                "3. Adjacency Analysis: For pixels meeting the grayscale condition, an adjacency check was performed in "
                "four directions (up, down, left, right) to identify if any direction contained 25 contiguous pixels also "
                "meeting the grayscale condition.\n\n"
                "4. Data Recording: The results were recorded in CSV files containing the coordinates, grayscale values, "
                "and GAP flags (1 for meeting conditions, 0 otherwise) for each pixel.\n\n"
                "5. Visualization: New images were generated highlighting pixels meeting GAP conditions in black against "
                "a white background, providing a clear visual representation of the spatial distribution of these features.\n\n"
            )
            
            # Results
            f.write("RESULTS\n")
            f.write("-" * 8 + "\n")
            f.write(
                "The analysis was performed on multiple images, with the following results. For each image, "
                "we present key statistics including the percentage of pixels meeting GAP conditions and average grayscale values.\n\n"
            )
            
            # Add results for each image
            for image_file in sorted(image_files):
                base_name = image_file.replace('_gap_highlighted.png', '')
                
                f.write(f"Sample: {base_name}\n")
                f.write("=" * (len(base_name) + 8) + "\n")
                
                # Add statistics
                if base_name in stats:
                    stat_data = stats[base_name]
                    f.write(f"Total pixels analyzed: {stat_data['total_pixels']:,}\n")
                    f.write(f"Pixels meeting GAP conditions: {stat_data['gap_pixels']:,}\n")
                    f.write(f"Percentage of GAP pixels: {stat_data['gap_percentage']:.2f}%\n")
                    f.write(f"Average grayscale value: {stat_data['avg_grayscale']:.2f}\n")
                
                f.write(f"Image file: {image_file}\n\n")
            
            # Conclusion
            f.write("CONCLUSION\n")
            f.write("-" * 10 + "\n")
            f.write(
                "This analysis successfully identified and characterized pixels meeting GAP conditions across multiple "
                "sample images. The results provide valuable insights into the spatial distribution and prevalence of "
                "these features, which can inform further investigation into structural properties and potential anomalies. "
                "The methodology developed for this analysis can be applied to additional samples to build a more "
                "comprehensive understanding of GAP conditions in similar materials.\n\n"
            )
        
        print(f"Text report generated successfully: {report_path}")
        
    except Exception as e:
        print(f"Error creating text report: {e}")
        return False
    
    return True

if __name__ == "__main__":
    print("Starting report generation...")
    
    # Verify if the output files from previous step exist
    results_dir = "./Results"
    if not os.path.exists(results_dir):
        print(f"Results directory {results_dir} not found.")
    else:
        csv_files = [f for f in os.listdir(results_dir) if f.endswith('_gap_analysis.csv')]
        image_files = [f for f in os.listdir(results_dir) if f.endswith('_gap_highlighted.png')]
        
        if csv_files and image_files:
            print("Calculation successful")
            print(f"Found {len(csv_files)} CSV files and {len(image_files)} highlighted images.")
            
            # Generate the report
            if generate_report():
                print("Report generation completed successfully.")
            else:
                print("Report generation encountered issues.")
        else:
            print("Required output files not found. Please run py1.py first.")
