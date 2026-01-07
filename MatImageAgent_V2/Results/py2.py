import os
import glob
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import pandas as pd
import numpy as np
from PIL import Image
import time

def verify_outputs():
    """Verify if the output files from py1.py exist"""
    results_dir = "./Results"
    
    if not os.path.exists(results_dir):
        print("Results directory not found.")
        return False
    
    csv_files = glob.glob(os.path.join(results_dir, "*_gap_analysis.csv"))
    png_files = glob.glob(os.path.join(results_dir, "*_gap_highlighted.png"))
    
    if csv_files and png_files:
        print("Calculation successful")
        return True
    else:
        print(f"Missing output files. CSV files: {len(csv_files)}, PNG files: {len(png_files)}")
        return False

def analyze_results():
    """Analyze the CSV files to extract key statistics"""
    results_dir = "./Results"
    csv_files = glob.glob(os.path.join(results_dir, "*_gap_analysis.csv"))
    
    results = {}
    for csv_file in csv_files:
        base_name = os.path.basename(csv_file).replace("_gap_analysis.csv", "")
        df = pd.read_csv(csv_file)
        
        # Calculate statistics
        total_pixels = len(df)
        gap_pixels = df[df['GAP_Flag'] == 1].shape[0]
        gap_percentage = (gap_pixels / total_pixels) * 100 if total_pixels > 0 else 0
        avg_grayscale = df['Grayscale'].mean()
        
        results[base_name] = {
            'total_pixels': total_pixels,
            'gap_pixels': gap_pixels,
            'gap_percentage': gap_percentage,
            'avg_grayscale': avg_grayscale
        }
    
    return results

def create_report(analysis_results):
    """Create a Word document report with the analysis results"""
    doc = docx.Document()
    
    # Set document title
    title = doc.add_heading('Pixel GAP Analysis Simulation Report', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Abstract section
    doc.add_heading('Abstract', level=1)
    abstract = doc.add_paragraph()
    abstract.add_run(
        "This report presents the results of a pixel-level GAP analysis conducted on a series of "
        "images. The analysis identifies pixels meeting specific grayscale and contiguity conditions, "
        "classifying them with a GAP flag. The study quantifies the distribution of GAP-flagged pixels "
        "across multiple images and visualizes their spatial patterns. This information provides insights "
        "into potential structural anomalies or regions of interest within the analyzed images."
    )
    
    # Introduction section
    doc.add_heading('Introduction', level=1)
    intro = doc.add_paragraph()
    intro.add_run(
        "The purpose of this analysis is to identify and quantify pixels meeting specific GAP conditions "
        "across a set of images. GAP conditions are defined as pixels with grayscale values between 1-155 "
        "that have at least one adjacent direction (up, down, left, or right) containing 25 contiguous "
        "pixels also meeting the grayscale condition. This analysis is particularly useful for identifying "
        "structural patterns, potential anomalies, or regions of interest within images.\n\n"
        "Background: Image analysis at the pixel level provides granular insights that might be missed "
        "through more general image processing techniques. By applying specific criteria to identify GAP "
        "pixels, we can detect subtle patterns and features that may have significance for various "
        "applications including material analysis, structural integrity assessment, or pattern recognition."
    )
    
    # Methods section
    doc.add_heading('Methods', level=1)
    methods = doc.add_paragraph()
    methods.add_run(
        "The analysis was implemented using Python with several key libraries including OpenCV for image "
        "processing, NumPy for numerical operations, and Pandas for data manipulation. The methodology "
        "consisted of the following steps:\n\n"
        "1. Image Preprocessing: Images were converted to grayscale and enhanced using Contrast Limited "
        "Adaptive Histogram Equalization (CLAHE) to improve feature visibility.\n\n"
        "2. Pixel-level GAP Analysis: Each pixel was evaluated against two conditions:\n"
        "   - Grayscale value between 1-155 (inclusive)\n"
        "   - At least one adjacent direction (up, down, left, right) containing 25 contiguous pixels "
        "     also meeting the grayscale condition\n\n"
        "3. Data Storage: Results were stored in CSV files containing pixel coordinates, grayscale values, "
        "and GAP flags (1 for pixels meeting conditions, 0 otherwise).\n\n"
        "4. Visualization: New images were generated highlighting GAP pixels in black and non-GAP pixels "
        "in white for visual analysis.\n\n"
        "5. Statistical Analysis: Summary statistics were calculated for each image, including the "
        "percentage of pixels meeting GAP conditions and average grayscale values."
    )
    
    # Results section
    doc.add_heading('Results', level=1)
    results_text = doc.add_paragraph()
    results_text.add_run(
        "The analysis was performed on multiple images, and the results are summarized below. For each "
        "image, we calculated the total number of pixels, the number of pixels meeting GAP conditions, "
        "the percentage of GAP pixels, and the average grayscale value.\n\n"
    )
    
    # Add a table with results
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    # Add header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Image'
    header_cells[1].text = 'Total Pixels'
    header_cells[2].text = 'GAP Pixels'
    header_cells[3].text = 'GAP Percentage (%)'
    header_cells[4].text = 'Avg. Grayscale'
    
    # Add data rows
    for image_name, stats in analysis_results.items():
        row_cells = table.add_row().cells
        row_cells[0].text = image_name
        row_cells[1].text = str(stats['total_pixels'])
        row_cells[2].text = str(stats['gap_pixels'])
        row_cells[3].text = f"{stats['gap_percentage']:.2f}%"
        row_cells[4].text = f"{stats['avg_grayscale']:.2f}"
    
    doc.add_paragraph("\n")
    doc.add_paragraph("Visual representation of the GAP analysis results:")
    
    # Add the generated images to the document
    results_dir = "./Results"
    png_files = glob.glob(os.path.join(results_dir, "*_gap_highlighted.png"))
    
    for png_file in png_files:
        base_name = os.path.basename(png_file).replace("_gap_highlighted.png", "")
        doc.add_paragraph(f"Image: {base_name}")
        doc.add_picture(png_file, width=Inches(6))
        doc.add_paragraph("\n")
    
    # Add conclusion
    doc.add_paragraph(
        "The visual representations clearly show the distribution of GAP pixels across the analyzed "
        "images. Black areas represent pixels meeting the GAP conditions, while white areas represent "
        "non-GAP pixels. The patterns and distributions provide insights into the structural "
        "characteristics of the original images and highlight regions of potential interest for further "
        "investigation."
    )
    
    # Save the document
    doc.save('Simulation_Report.docx')
    print(f"Report saved as 'Simulation_Report.docx'")

if __name__ == "__main__":
    # Verify outputs from py1.py
    if verify_outputs():
        # Analyze results
        analysis_results = analyze_results()
        
        # Create the report
        create_report(analysis_results)
    else:
        print("Error: Required output files not found. Please run py1.py first.")
